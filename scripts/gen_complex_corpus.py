#!/usr/bin/env python3
"""Long, realistic hospital .docx fixtures — where the failure modes are different in KIND.

THE SHORT FIXTURES ANSWER A DIFFERENT QUESTION. gen_sc_corpus.py builds one-defect documents so
that detection and attribution are unambiguous. That is the right shape for "does the rule fire",
and it is the wrong shape for everything below, because the things that break on a customer's
real estate do not appear on a one-page file at all:

  FALSE POSITIVES COMPOUND WITH LENGTH. Four false positives across 32 one-page fixtures reads
  as a good precision number. The same per-paragraph rate on a 25-page benefits handbook is ~40
  spurious findings on one document, and a reviewer who has to dismiss 40 items stops reading
  the list. Precision measured on short fixtures systematically flatters the product.

  THE IMAGE CAPS BAND THE RESULTS, AND NOTHING SAYS SO. _VISION_MAX_IMAGES is 25 and
  ACP_OCR_MAX_IMAGES is 30. A document with 35 images therefore has three regions — images 1-25
  get vision alt text, 26-30 get OCR only, 31+ get neither — and the output looks like a model
  that got tired. `img-cap-35` exists to make that banding visible rather than mysterious.

  JARGON IS AN ADVERSARIAL INPUT FOR LANGUAGE DETECTION. "Administer 5 mg IV q8h PRN; monitor
  CBC, BMP, and INR" is not English by any statistical measure langdetect applies. Clinical
  abbreviation density is the realistic way 3.1.2 produces false positives, and no synthetic
  English paragraph will surface it.

  HEADERS, FOOTERS AND FOOTNOTES ARE SEPARATE PARTS. word/header1.xml is a different part from
  word/document.xml, and a detector that walks only the body silently ignores an undescribed
  logo that appears on all 25 pages. ALT_TARGETS covers headers and footers; footnotes are not
  in that table, which is a claim worth testing rather than assuming.

GROUND TRUTH IS COUNTED AT GENERATION TIME. For a 25-page document nobody can establish after
the fact how many undescribed images it contains — the count is only reliably known at the
moment the generator writes them, so it is recorded there. Expected counts, not just expected
criteria, because "found 1.1.1" on a document with eleven undescribed images is not success.

Run:
    python scripts/gen_complex_corpus.py --out ~/Downloads/acp-docx-eval/complex-corpus
"""
from __future__ import annotations

import argparse
import io
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))
sys.path.insert(0, str(ROOT / "api"))

import corpus_expectations as ce  # noqa: E402
from gen_sc_corpus import (_add_hyperlink, _add_image, _add_table, _doc,  # noqa: E402
                           _run, _solid_png, _text_png)

from docx.enum.text import WD_BREAK  # noqa: E402
from docx.shared import Pt  # noqa: E402

FMT = "docx"

# Real terminology, because the point is to be hostile to a language detector and to any model
# that pattern-matches on plain prose. None of it is real patient data — these are textbook
# abbreviations and public CPT/ICD categories.
CLINICAL = [
    "Prior authorization is required for CPT 70553 (MRI brain w/ and w/o contrast) when "
    "ordered outside the ordering provider's specialty scope.",
    "Administer 5 mg IV q8h PRN; monitor CBC, BMP, and INR at 24h and 72h post-initiation.",
    "Members assigned to Tier 3 formulary status may request an exception via the P&T "
    "committee under UM policy 4.11(c).",
    "Documented contraindications include CKD stage 4 (eGFR < 30 mL/min/1.73m2), active "
    "GI bleed, and known hypersensitivity to NSAIDs.",
    "Claims submitted with ICD-10 Z00.00 as the primary diagnosis are adjudicated under the "
    "preventive services benefit and are not subject to the deductible.",
    "The attending shall document H&P within 24h of admission per CMS CoP 482.24(c)(4)(i).",
]

BENEFITS = [
    "Your deductible is the amount you pay for covered services before the plan begins to pay.",
    "Coinsurance is the percentage of allowed charges you pay after the deductible is met.",
    "Out-of-network providers may bill you for the difference between their charge and the "
    "plan's allowed amount, which does not count toward your out-of-pocket maximum.",
    "Preventive services are covered at one hundred percent when delivered by an in-network "
    "provider and coded as preventive.",
]


def _para(d, text, *, size=None):
    p = d.add_paragraph()
    _run(p, text, size_pt=size)
    return p


def _page_break(d):
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _header_logo(d, *, alt: str | None):
    """An image in word/header1.xml — a SEPARATE part from the body.

    A logo in the header appears on every page. If the detector walks only word/document.xml it
    misses one undescribed image that a screen-reader user meets 25 times.
    """
    hdr = d.sections[0].header
    p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    from docx.shared import Inches
    run = p.add_run()
    run.add_picture(io.BytesIO(_text_png(["UT Health Benefits"], size=(400, 80))),
                    width=Inches(1.6))
    from docx.oxml.ns import qn
    for docpr in run._r.iter(qn("wp:docPr")):
        if alt is None:
            docpr.attrib.pop("descr", None)
        else:
            docpr.set("descr", alt)


# ── documents ─────────────────────────────────────────────────────────────────

def d_benefits_handbook():
    """~25 pages of plausible benefits prose, deliberately MOSTLY CLEAN.

    The seeded defects are few and known. Everything else exists to give false positives room to
    accumulate — which is the measurement. A precision number from this document means something
    a precision number from a one-paragraph fixture does not.
    """
    d = _doc("2026 Employee Benefits Handbook")
    truth = {"undescribed_images": 0, "bad_links": 0, "headerless_tables": 0}

    _header_logo(d, alt="UT Health Benefits")
    d.add_heading("2026 Employee Benefits Handbook", level=1)
    _para(d, "This handbook describes the medical, dental, and vision benefits available to "
             "eligible employees and their dependents for the 2026 plan year.")

    for page in range(24):
        _page_break(d)
        d.add_heading(f"Section {page + 1}", level=2)
        for i in range(4):
            _para(d, BENEFITS[(page + i) % len(BENEFITS)])
        d.add_heading(f"Section {page + 1}.1 — Cost sharing", level=3)
        _add_table(d, [["Service", "In-network", "Out-of-network"],
                       ["Office visit", "$30 copay", "40% coinsurance"],
                       ["Specialist", "$60 copay", "40% coinsurance"]], header=True)
        # One seeded defect every eighth page, so the density resembles a real document rather
        # than a test file. A corpus where every page is broken measures nothing about noise.
        if page % 8 == 3:
            _add_image(d, _solid_png(), alt=None)
            truth["undescribed_images"] += 1
        if page % 8 == 6:
            p = _para(d, "For the full schedule of benefits, ")
            _add_hyperlink(p, f"https://example.org/sob/{page}", "click here")
            truth["bad_links"] += 1

    return d, {"1.1.1": ce.FAIL, "2.4.4": ce.FAIL}, truth, \
        ("~25 pages, mostly compliant. 3 undescribed images and 3 non-descriptive links seeded "
         "at low density. The metric that matters here is PRECISION — how many findings arrive "
         "that were not seeded")


def d_clinical_policy():
    """Dense clinical jargon — adversarial for language detection and for any prose-trained model."""
    d = _doc("Utilization Management Policy 4.11")
    truth = {"undescribed_images": 0}
    d.add_heading("Utilization Management Policy 4.11", level=1)
    for page in range(12):
        if page:
            _page_break(d)
        d.add_heading(f"4.11.{page + 1} Criteria", level=2)
        for i in range(6):
            _para(d, CLINICAL[(page + i) % len(CLINICAL)])
    return d, {}, truth, \
        ("ADVERSARIAL: 12 pages of abbreviation-dense clinical text, NO seeded defects. Any "
         "3.1.2 finding here is a false positive caused by jargon reading as a foreign language "
         "— the realistic way that criterion misfires on a hospital's real documents")


def d_lab_report():
    """Dense numeric tables — the shape that trips the chart heuristic and header detection."""
    d = _doc("Comprehensive Metabolic Panel — Result Summary")
    truth = {"headerless_tables": 0}
    d.add_heading("Comprehensive Metabolic Panel", level=1)
    rows = [["Analyte", "Result", "Units", "Reference range", "Flag"],
            ["Sodium", "138", "mmol/L", "135-145", ""],
            ["Potassium", "5.4", "mmol/L", "3.5-5.1", "H"],
            ["Chloride", "101", "mmol/L", "98-107", ""],
            ["CO2", "24", "mmol/L", "22-29", ""],
            ["BUN", "32", "mg/dL", "7-20", "H"],
            ["Creatinine", "1.9", "mg/dL", "0.6-1.3", "H"],
            ["eGFR", "38", "mL/min/1.73m2", "> 60", "L"],
            ["Glucose", "142", "mg/dL", "70-99", "H"]]
    for panel in range(6):
        if panel:
            _page_break(d)
        d.add_heading(f"Panel {panel + 1}", level=2)
        # Half the panels are correctly marked, half are not — so the document scores BOTH
        # detection and false positives on the same criterion, in the same file.
        marked = panel % 2 == 0
        _add_table(d, rows, header=marked)
        if not marked:
            truth["headerless_tables"] += 1
        _para(d, "Values flagged H or L fall outside the reference range and should be "
                 "correlated clinically.")
    return d, {"1.3.1": ce.FAIL}, truth, \
        ("6 dense numeric tables, 3 correctly marked and 3 not. Numeric density is also what "
         "ocr._looks_like_chart keys on, so this is the realistic collision between a lab "
         "table and the 1.4.5 essential-exception heuristic")


def d_image_cap():
    """35 images — deliberately across both processing caps.

    _VISION_MAX_IMAGES is 25 and ACP_OCR_MAX_IMAGES is 30, so this document has three regions
    and the output should change character twice. Nothing in the product reports which band an
    image fell into; without a fixture that crosses them, the drop-off looks like model quality
    degrading rather than a cap being hit.
    """
    d = _doc("Formulary Update — Illustrated")
    truth = {"undescribed_images": 35, "vision_cap": 25, "ocr_cap": 30}
    d.add_heading("Formulary Update", level=1)
    for i in range(35):
        if i and i % 3 == 0:
            _page_break(d)
        _para(d, f"Figure {i + 1}: tier placement for the agents listed below.")
        # EVERY image must have distinct BYTES. The first version drew `Tier {i % 4 + 1}`, so
        # only four unique PNGs existed; the OOXML package deduplicated 35 references down to 4
        # media parts, and the OCR pass — which walks media parts, not references — saw four
        # images. The fixture reported "4 of 35 reached OCR" and looked exactly like the cap it
        # was built to demonstrate. It was measuring zip deduplication.
        _add_image(d, _text_png([f"Agent {i + 1:02d} formulary tier {i % 4 + 1}",
                                 f"prior authorization required code {7000 + i}",
                                 "step therapy applies to all new starts"]), alt=None)
    return d, {"1.1.1": ce.FAIL, "1.4.5": ce.FAIL}, truth, \
        ("35 undescribed images of text, crossing the vision cap at 25 and the OCR cap at 30. "
         "Expect THREE bands in the remediation output; a scorecard that reports one average "
         "over all 35 hides the cliff")


def d_consent_pack():
    """Forms, footnotes and a header image — parts the body-only walk never reaches."""
    d = _doc("Consent for Treatment and Release of Information")
    truth = {"undescribed_images": 1}
    _header_logo(d, alt=None)                    # header part, NOT word/document.xml
    truth_note = "header logo is undescribed and appears on every page"
    d.add_heading("Consent for Treatment", level=1)
    for page in range(8):
        if page:
            _page_break(d)
        d.add_heading(f"Part {page + 1}", level=2)
        _para(d, CLINICAL[page % len(CLINICAL)])
        _para(d, "I have read and understand the information provided above.")
    return d, {"1.1.1": ce.FAIL}, truth, (
        f"8 pages; the ONLY seeded defect is in word/header1.xml — {truth_note}. A detector "
        f"that walks only the body reports a clean document")


DOCS = [
    ("benefits-handbook-25p", d_benefits_handbook, "scale"),
    ("clinical-policy-jargon", d_clinical_policy, "adversarial"),
    ("lab-report-dense",      d_lab_report,       "mixed"),
    ("img-cap-35",            d_image_cap,        "cap"),
    ("consent-header-only",   d_consent_pack,     "parts"),
]


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--out", type=Path,
                    default=Path.home() / "Downloads" / "acp-docx-eval" / "complex-corpus")
    args = ap.parse_args()
    docs = args.out / "docs"
    docs.mkdir(parents=True, exist_ok=True)

    manifest, problems = [], []
    for name, build, kind in DOCS:
        d, expect, truth, note = build()
        path = docs / f"{name}.docx"
        d.save(path)
        for sc, verdict in expect.items():
            if verdict not in ce.possible_verdicts(sc, FMT):
                problems.append(f"{name}: {sc}={verdict} is unreachable for ({sc}, {FMT})")
        manifest.append({"file": f"docs/{name}.docx", "name": name, "kind": kind,
                         "format": FMT, "expect": expect, "truth": truth, "note": note})
        print(f"  {kind:12} {name:24} {path.stat().st_size // 1024:>4} KB   {truth}")

    if problems:
        print("\nIMPOSSIBLE EXPECTATIONS:", file=sys.stderr)
        for p in problems:
            print(f"  {p}", file=sys.stderr)
        return 1

    (args.out / "manifest.json").write_text(json.dumps(
        {"format": FMT, "fixtures": manifest,
         "lanes": {sc: {"clean_verdict": ce.clean_verdict(sc, FMT),
                        "possible": sorted(ce.possible_verdicts(sc, FMT)),
                        "can_pass": ce.can_ever_pass(sc, FMT)}
                   for sc in sorted(sc for sc, f in
                                    ce.pol.SCOPE_PRESETS["acp-core-17"].items() if FMT in f)}},
        indent=2) + "\n")
    print(f"\n{len(manifest)} documents -> {args.out / 'manifest.json'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
