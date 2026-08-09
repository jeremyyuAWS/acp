#!/usr/bin/env python3
"""Verified Remediation Rate — did the fix actually work, and did it break anything else?

THE CLOSED LOOP, WHICH IS THE WHOLE POINT. Every other measurement of a remediation in this
repo asks a model or a human whether the output looks good. This one asks the DETERMINISTIC
ASSESSOR, twice:

    assess -> remediate -> assess again

A criterion that was FAIL and is no longer FAIL was fixed. One that was fine and is now FAIL is
a regression. Neither judgement involves an opinion, which is why this number can gate a policy
change and a judge score cannot. It is also the only metric here that survives the observation
that ACP's assessment is mostly deterministic: remediation is where the model actually decides
things, so remediation is where model choice can be measured honestly.

VRR IS NOT "DID THE VERDICT IMPROVE". Three things disqualify a fix that a naive before/after
comparison would count as a success:

  CONTENT LOSS. Deleting the image clears 1.1.1. So does deleting the paragraph that contained
  it. The verdict improves and the document is damaged, and on a patient benefits notice that is
  far worse than the finding. Integrity is checked on every document, and a fix that loses a
  paragraph, table, image or section is scored as a FAILURE regardless of what the re-scan says.

  REGRESSION ELSEWHERE. Fixing 1.1.1 must not break 1.3.1. Regressions are computed over EVERY
  criterion, not just the seeded one — a harness that only re-checks what it set out to fix
  cannot see the damage it did on the way past.

  ABSTENTION MISCOUNTED AS FAILURE. ACP deliberately does not auto-apply ungrounded alt text; it
  routes it to `proposals` for a human. That is the honesty split working, and scoring it as
  "unfixed" would rank a model that confabulates above one that declines. Proposals are counted
  as their own outcome and reported separately, never folded into the failure column.

WHY THE IMAGE COUNT IS READ FROM THE ZIP. `Document.inline_shapes` sees inline body images only
— it misses header/footer parts and floating shapes. The consent fixture's only image lives in
word/header1.xml, so an integrity check built on inline_shapes would report 0 images before AND
after and certify the deletion of the one image in the document. Counting word/media entries
sees every part.

Run:
    python scripts/score_remediation.py --corpus ~/Downloads/acp-docx-eval/sc-corpus \
        --models qwen2.5vl:7b,qwen2.5vl:3b
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import statistics
import subprocess
import sys
import tempfile
import time
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "api"))
sys.path.insert(0, str(ROOT / "scripts"))

import corpus_expectations as ce  # noqa: E402
from score_assessment import FMT, scan, verdicts  # noqa: E402

FAIL = ce.FAIL


def remediate(path: Path) -> dict:
    """Run the SHIPPED remediation over a copy, COUNTING every model call it makes.

    THE CALL COUNT IS NOT DIAGNOSTIC DETAIL — IT IS WHAT MAKES THE MODEL COLUMN MEAN ANYTHING.
    "the model made no difference" and "the model never ran" produce byte-identical benchmark
    tables, and the first run of this harness produced the second while reading as the first: a
    tidy 50% VRR, zero regressions, and not one call to a vision or text model. Three separate
    things suppressed them — `infer_decorative` short-circuits on a near-uniform image before
    vision is consulted, the default text model (llama3.2) was not pulled so every text draft was
    skipped, and nothing in the output said either. A sweep over four models would have printed
    four identical rows and the honest conclusion "parameter count does not matter", from an
    experiment in which no parameters were used.

    So the counters are wrapped around ai.describe_image and ai.suggest_fix here, and
    `summarise` refuses to compare models when they are zero.

    `ai` and `remediate_office` are reloaded because the model is chosen from the environment at
    import time; without the reload every model in a sweep runs as whichever was imported first.
    """
    import importlib

    import ai
    import remediate_office
    importlib.reload(ai)
    importlib.reload(remediate_office)

    calls = {"vision": 0, "text": 0, "vision_seconds": 0.0}

    def _count(fn, key):
        def wrapped(*a, **kw):
            t = time.time()
            try:
                return fn(*a, **kw)
            finally:
                calls[key] += 1
                if key == "vision":
                    calls["vision_seconds"] += time.time() - t
        return wrapped

    # EVERY vision entry point, not the obvious one. The remediation lane calls
    # `describe_image_structured`; `describe_image` is what the eval scripts call. Patching only
    # the latter counted 0 vision calls through runs that were visibly spending 3.5s per image in
    # a vision model — a counter that reads zero while the model is plainly running is worse than
    # no counter, because it is the exact signal this harness uses to invalidate a sweep.
    for attr, key in (("describe_image", "vision"),
                      ("describe_image_structured", "vision"),
                      ("suggest_fix", "text"),
                      ("explain_issue", "text")):
        if hasattr(ai, attr):
            setattr(ai, attr, _count(getattr(ai, attr), key))
    # remediate_office may hold its own reference to the module; point it at the patched one.
    remediate_office.ai = ai

    work = Path(tempfile.mkdtemp()) / path.name
    shutil.copy(path, work)
    applied, diffs, proposals = [], [], []
    t0 = time.time()
    try:
        res = remediate_office.remediate_office(work, ai_enabled=True, applied_fixes=applied,
                                                diffs=diffs, proposals=proposals)
    except Exception as e:                                          # noqa: BLE001
        return {"error": f"{e.__class__.__name__}: {e}", "out": work, "seconds": 0.0,
                "edits": 0, "proposals": 0, "calls": calls}
    secs = time.time() - t0
    # remediate_office returns None for the path when it wrote NOTHING, which is the correct
    # answer on a clean document — assuming a path here is what crashed the gate run.
    out = (res[0] if isinstance(res, tuple) else res) or work
    summaries = res[1] if isinstance(res, tuple) and len(res) > 1 else []
    return {"out": Path(out), "seconds": secs, "edits": len(summaries),
            "proposals": len(proposals), "applied": applied, "error": None, "calls": calls}


def preflight(text_model: str, vision_model: str | None) -> list[str]:
    """Models that are configured but not actually pulled.

    ai.py defaults OLLAMA_MODEL to llama3.2. If that is not pulled, every text draft is silently
    skipped and the run still completes, still writes a scorecard, and still looks like a result.
    A one-line warning at import is not enough — it scrolls past. This is checked before any
    measurement, and reported alongside the numbers rather than upstream of them.
    """
    try:
        out = subprocess.run(["ollama", "list"], capture_output=True, text=True,
                             timeout=15).stdout
    except Exception:                                               # noqa: BLE001
        return ["could not run `ollama list` — model availability unverified"]
    have = {ln.split()[0] for ln in out.splitlines()[1:] if ln.split()}
    missing = []
    for m in filter(None, (text_model, vision_model)):
        if m not in have and f"{m}:latest" not in have:
            missing.append(m)
    return [f"NOT PULLED: {m} — its lane will be silently skipped" for m in missing]


def integrity(path: Path) -> dict:
    """Structural census. Reads the PACKAGE, not just the body.

    `Document.inline_shapes` counts inline images in word/document.xml and nothing else, so a
    header logo, a footer image or a floating shape is invisible to it. Media parts are counted
    from the zip instead, which is the only view that sees all of them.
    """
    from docx import Document
    out: dict = {"valid": True}
    try:
        d = Document(str(path))
        out |= {"paragraphs": len(d.paragraphs), "tables": len(d.tables),
                "sections": len(d.sections),
                "characters": sum(len(p.text) for p in d.paragraphs)}
        with zipfile.ZipFile(path) as z:
            names = z.namelist()
            out["media"] = sum(1 for n in names if n.startswith("word/media/"))
            out["parts"] = len(names)
    except Exception as e:                                          # noqa: BLE001
        # A document that no longer opens is the worst possible outcome and must never be
        # mistaken for "no findings". It is reported as invalid, not as clean.
        return {"valid": False, "error": f"{e.__class__.__name__}: {e}"}
    return out


def damage(before: dict, after: dict) -> list[str]:
    """What the remediation destroyed. Characters may move — a fix writes values — but a
    paragraph, table, section or media part disappearing is damage, not remediation."""
    if not after.get("valid"):
        return [f"document no longer opens: {after.get('error')}"]
    bad = []
    for k in ("paragraphs", "tables", "sections", "media"):
        if after.get(k, 0) < before.get(k, 0):
            bad.append(f"{k} {before.get(k)} -> {after.get(k)}")
    # Losing most of the text without losing a paragraph is still damage — a fix that blanks
    # runs in place leaves the structure intact and the document empty.
    if before.get("characters", 0) and after.get("characters", 0) < before["characters"] * 0.9:
        bad.append(f"characters {before['characters']} -> {after['characters']}")
    return bad


def model_size_gb(model: str) -> float | None:
    """Disk size from `ollama list`, for the quality-per-GB column. None when unavailable."""
    try:
        out = subprocess.run(["ollama", "list"], capture_output=True, text=True, timeout=15).stdout
    except Exception:                                               # noqa: BLE001
        return None
    for line in out.splitlines():
        parts = line.split()
        if parts and parts[0] == model:
            for i, tok in enumerate(parts):
                if tok.upper() in ("GB", "MB") and i:
                    try:
                        v = float(parts[i - 1])
                    except ValueError:
                        return None
                    return round(v / 1024, 3) if tok.upper() == "MB" else v
    return None


def run_corpus(corpus: Path, scs: list[str], label: str) -> dict:
    man = json.loads((corpus / "manifest.json").read_text())
    rows, lat = [], []

    for fx in man["fixtures"]:
        path = corpus / fx["file"]
        before_issues, _ = scan(path)
        before = verdicts(before_issues, scs)
        before_int = integrity(path)

        r = remediate(path)
        lat.append(r["seconds"])
        if r["error"]:
            rows.append({"name": fx["name"], "error": r["error"]})
            print(f"  ERR  {fx['name']:24} {r['error']}")
            continue

        after_issues, _ = scan(r["out"])
        after = verdicts(after_issues, scs)
        after_int = integrity(r["out"])

        fixed = sorted(sc for sc in scs if before[sc] == FAIL and after[sc] != FAIL)
        unresolved = sorted(sc for sc in scs if before[sc] == FAIL and after[sc] == FAIL)
        regressed = sorted(sc for sc in scs if before[sc] != FAIL and after[sc] == FAIL)
        dmg = damage(before_int, after_int)

        attempted = bool(before_int.get("valid")) and bool(
            [sc for sc in scs if before[sc] == FAIL])
        # VERIFIED means all three: something got fixed, nothing regressed, nothing was lost.
        verified = bool(fixed) and not regressed and not dmg

        rows.append({"name": fx["name"], "kind": fx.get("kind"),
                     "fixed": fixed, "unresolved": unresolved, "regressed": regressed,
                     "damage": dmg, "proposals": r["proposals"], "edits": r["edits"],
                     "seconds": round(r["seconds"], 1), "attempted": attempted,
                     "verified": verified, "calls": r["calls"]})

        if not attempted:
            mark, note = "----", "nothing to fix"
        elif verified:
            mark, note = "VRR ", f"fixed {','.join(fixed)}"
        elif dmg:
            mark, note = "DMG!", "; ".join(dmg)
        elif regressed:
            mark, note = "REGR", f"broke {','.join(regressed)}"
        elif r["proposals"]:
            mark, note = "prop", f"{r['proposals']} proposal(s) — abstained to a human"
        else:
            mark, note = "miss", f"still failing {','.join(unresolved)}"
        c = r["calls"]
        seen = f"v{c['vision']}/t{c['text']}"
        print(f"  {mark} {fx['name']:24} {r['seconds']:5.1f}s  {seen:8} "
              f"edits={r['edits']:<3} {note}")

    return {"label": label, "rows": rows, "latencies": lat}


def summarise(res: dict, scs: list[str], model_note: str) -> None:
    rows = [r for r in res["rows"] if not r.get("error")]
    att = [r for r in rows if r["attempted"]]
    ver = [r for r in att if r["verified"]]
    dmg = [r for r in att if r["damage"]]
    reg = [r for r in att if r["regressed"]]
    prop = [r for r in att if not r["verified"] and not r["damage"]
            and not r["regressed"] and r["proposals"]]
    miss = [r for r in att if not r["verified"] and not r["damage"]
            and not r["regressed"] and not r["proposals"]]

    n = len(att) or 1
    vision = sum(r["calls"]["vision"] for r in rows)
    text = sum(r["calls"]["text"] for r in rows)

    print(f"\n=== VERIFIED REMEDIATION RATE — {model_note} ===")
    if not (vision or text):
        # Loud, first, and before any number the reader might quote. A scorecard from a run in
        # which no model executed is a measurement of the deterministic engine wearing a model's
        # name, and it is indistinguishable from a real result unless something says so here.
        print("  !! NO MODEL WAS CALLED IN THIS RUN — 0 vision, 0 text invocations.")
        print("  !! Every number below describes the DETERMINISTIC lane only. It is NOT a")
        print("  !! measurement of this model and must not be compared against another model.")
        print("  !! Usual causes: the image is near-uniform so infer_decorative short-circuits")
        print("  !! before vision; or the configured model is not pulled (see preflight above).")
    else:
        print(f"  model calls: {vision} vision, {text} text "
              f"across {len(rows)} document(s)")
    print(f"  VRR                {len(ver):>3} / {len(att)}  ({len(ver) / n * 100:.0f}%)  "
          "fixed something, regressed nothing, lost nothing")
    print(f"  abstained          {len(prop):>3} / {len(att)}  ({len(prop) / n * 100:.0f}%)  "
          "routed to a human proposal — CORRECT behaviour, not a failure")
    print(f"  unresolved         {len(miss):>3} / {len(att)}  ({len(miss) / n * 100:.0f}%)  "
          "no edit, no proposal, still failing")
    print(f"  REGRESSED          {len(reg):>3} / {len(att)}  ({len(reg) / n * 100:.0f}%)  "
          "broke a criterion that was fine before")
    print(f"  INTEGRITY DAMAGE   {len(dmg):>3} / {len(att)}  ({len(dmg) / n * 100:.0f}%)  "
          "content lost — disqualifies the fix whatever the re-scan says")

    # Per-SC, because one number over a corpus hides which criteria are actually automatable and
    # that is the decision this benchmark exists to inform.
    per: dict[str, dict[str, int]] = {}
    for r in att:
        for sc in r["fixed"]:
            per.setdefault(sc, {}).setdefault("fixed", 0)
            per[sc]["fixed"] = per[sc].get("fixed", 0) + 1
        for sc in r["unresolved"]:
            per.setdefault(sc, {})
            per[sc]["unresolved"] = per[sc].get("unresolved", 0) + 1
        for sc in r["regressed"]:
            per.setdefault(sc, {})
            per[sc]["regressed"] = per[sc].get("regressed", 0) + 1
    if per:
        print(f"\n  {'SC':8}{'fixed':>7}{'unresolved':>12}{'regressed':>11}   lane")
        for sc in sorted(per):
            d = per[sc]
            lane = "deterministic" if ce.can_ever_pass(sc, FMT) else "review"
            print(f"  {sc:8}{d.get('fixed', 0):>7}{d.get('unresolved', 0):>12}"
                  f"{d.get('regressed', 0):>11}   {lane}")

    ls = sorted(res["latencies"])
    if ls:
        p95 = ls[min(len(ls) - 1, int(0.95 * len(ls)))]
        total = sum(ls) or 1
        print(f"\n  latency  p50 {statistics.median(ls):.1f}s   p95 {p95:.1f}s   "
              f"{len(ls) / total * 3600:.0f} docs/hour single-threaded")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--corpus", type=Path,
                    default=Path.home() / "Downloads" / "acp-docx-eval" / "sc-corpus")
    ap.add_argument("--models", default="",
                    help="comma-separated vision models to sweep; empty = current environment")
    ap.add_argument("--text-model", default=os.environ.get("OLLAMA_MODEL", "llama3.1:8b"))
    ap.add_argument("--out", type=Path)
    args = ap.parse_args()

    man = json.loads((args.corpus / "manifest.json").read_text())
    scs = sorted(man["lanes"])
    os.environ.setdefault("OLLAMA_MODEL", args.text_model)

    models = [m.strip() for m in args.models.split(",") if m.strip()]

    warnings = preflight(os.environ.get("OLLAMA_MODEL", ""),
                         models[0] if models else os.environ.get("OLLAMA_VISION_MODEL"))
    for w in warnings:
        print(f"  PREFLIGHT  {w}")

    results = []
    for model in (models or [None]):
        if model:
            os.environ["OLLAMA_VISION_MODEL"] = model
        note = model or os.environ.get("OLLAMA_VISION_MODEL", "(environment default)")
        size = model_size_gb(note) if model else None
        print(f"\n### vision={note}"
              + (f"  {size} GB" if size else "")
              + f"  text={os.environ.get('OLLAMA_MODEL')}\n")
        res = run_corpus(args.corpus, scs, note)
        res["model"] = note
        res["size_gb"] = size
        summarise(res, scs, note)
        results.append(res)

    if len(results) > 1:
        # A cross-model table is a claim that the differences between rows are caused by the
        # models. When a run made no model calls that claim is false, so the row is marked rather
        # than quietly averaged in — an unmarked table of identical rows reads as the finding
        # "size does not matter", which is the most expensive wrong conclusion available here.
        print(f"\n=== ACROSS MODELS ===\n{'model':22}{'GB':>6}{'VRR':>8}{'regr':>7}"
              f"{'dmg':>6}{'abst':>7}{'p50':>8}{'calls':>8}")
        inert = []
        for r in results:
            att = [x for x in r["rows"] if x.get("attempted")]
            n = len(att) or 1
            ver = sum(1 for x in att if x["verified"])
            reg = sum(1 for x in att if x["regressed"])
            dmg = sum(1 for x in att if x["damage"])
            ab = sum(1 for x in att if not x["verified"] and not x["damage"]
                     and not x["regressed"] and x["proposals"])
            nc = sum(x["calls"]["vision"] + x["calls"]["text"]
                     for x in r["rows"] if x.get("calls"))
            if not nc:
                inert.append(r["model"])
            p50 = statistics.median(r["latencies"]) if r["latencies"] else 0
            print(f"{r['model']:22}{(r['size_gb'] or 0):>6}{ver / n * 100:>7.0f}%"
                  f"{reg / n * 100:>6.0f}%{dmg / n * 100:>5.0f}%{ab / n * 100:>6.0f}%"
                  f"{p50:>7.1f}s{nc:>8}")
        if inert:
            print(f"\n  !! NOT A MODEL COMPARISON. These made ZERO model calls: "
                  f"{', '.join(inert)}.\n  !! Their rows are the deterministic engine, identical "
                  "by construction. Fix the corpus or\n  !! the model availability before "
                  "reading any difference between rows as a model effect.")
        print("\n  Read VRR against regr and dmg, never alone: a model that edits nothing scores"
              "\n  0% regression and 0% damage, and is useless. A model with a high VRR and any"
              "\n  damage at all is worse than one that abstains.")

    if args.out:
        args.out.write_text(json.dumps(results, indent=2, default=str) + "\n")
        print(f"\nwrote {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
