#!/usr/bin/env python3
"""A DOCX corpus for comparing local models, one issue per file.

WHY NOT gen_demo_fixtures.build_docx. That builds ONE document carrying every issue at once,
which is right for a demo and wrong for a measurement: when a model produces four fixes you
cannot tell which criterion each answered, and a model that fails 2.4.4 while succeeding at 1.1.1
scores the same as the reverse. Isolating the issues is the whole point.

WHAT IT COVERS, and why it is fewer than seventeen. On .docx the Core 17 splits three ways:

    MODEL-DEPENDENT (7)  1.1.1 1.3.2 1.3.3 1.4.5 2.4.4 3.1.2 4.1.2   <- the variables
    DETERMINISTIC   (5)  1.3.1 1.4.3 2.4.2 2.4.6 3.1.1               <- controls
    NO LANE         (3)  1.4.1 1.4.11 2.1.2                          <- known blanks

Only the first group can vary with model choice; the second is rule code and MUST come out
identical for every model. That makes the controls load-bearing rather than filler: if a
deterministic criterion moves between models, the harness is measuring something other than the
model, and the whole comparison is void.

Each file declares its expected criteria in the manifest, so scoring is against a stated
expectation rather than against whatever the first model happened to produce.

Run:
    python scripts/gen_model_eval_corpus.py --out test-corpus/model-eval
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from gen_demo_fixtures import (  # noqa: E402
    _add_hyperlink, _add_unlabeled_checkbox, _png, _png_of_text, _strip_default_language,
)

import io  # noqa: E402

from docx import Document  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402


def _titled(doc, title="Benefits Notice"):
    """Give a document the deterministic basics so a fixture isolates ONE criterion.

    Without this every file also trips 2.4.2 and 3.1.1, every model 'fixes' them, and the
    per-criterion signal is buried under two fixes that had nothing to do with the model.
    """
    doc.core_properties.title = title
    doc.core_properties.language = "en-US"


# Each builder returns the criteria it deliberately breaks. The list IS the expectation the
# harness scores against — a builder that stops producing its issue makes the fixture lie, so
# they are kept minimal and obvious rather than clever.
def f_clean(doc):
    _titled(doc)
    doc.add_heading("Open Enrollment", level=1)
    doc.add_paragraph("Enrollment runs March 1 to March 31, 2026. Contact HR with questions.")
    return []


def f_alt_photo(doc):
    _titled(doc)
    doc.add_paragraph("Our campus:")
    doc.add_picture(io.BytesIO(_png(color=(40, 90, 60))))
    return ["1.1.1"]        # textless image: vision must describe it, OCR cannot ground it


def f_alt_text_image(doc):
    _titled(doc)
    doc.add_paragraph("Notice:")
    doc.add_picture(io.BytesIO(_png_of_text()))
    return ["1.1.1", "1.4.5"]   # image OF text: OCR grounds it, and it is also images-of-text


def f_link_generic(doc):
    _titled(doc)
    p = doc.add_paragraph("For the full accessibility policy, ")
    _add_hyperlink(p, "https://example.com/accessibility-policy", "click here")
    p.add_run(".")
    return ["2.4.4"]


def f_link_raw_url(doc):
    _titled(doc)
    p = doc.add_paragraph("See ")
    _add_hyperlink(p, "https://example.com/benefits/2026/enrollment",
                   "https://example.com/benefits/2026/enrollment")
    p.add_run(" for details.")
    return ["2.4.4"]


def f_sensory(doc):
    _titled(doc)
    doc.add_paragraph("Click the round green button on the right to continue.")
    return ["1.3.3"]


def f_lang_parts(doc):
    _titled(doc)
    # BOTH sentences must clear 12 words on their own — _SEG_SPLIT splits on sentence
    # boundaries, so a 20-word paragraph made of two 10-word sentences contributes nothing.
    # An 11-word English sentence here left French as the only qualifying segment, and the
    # detector needs TWO in different languages before it reports anything.
    doc.add_paragraph("Employees may review their medical, dental, and vision coverage options "
                      "during the annual enrollment window each year without exception.")
    # >= 12 words: textchecks._MIN_SEG_WORDS is 12 because langdetect is unreliable below
    # that, and the detector needs TWO confident segments in different languages. The
    # earlier 10-word passage was never seen, and the fixture scored a working check as a
    # miss on every model.
    doc.add_paragraph("Veuillez consulter la politique d’accessibilité avant le 31 mars 2026 "
                      "pour confirmer votre choix de couverture médicale et dentaire.")
    return ["3.1.2"]


def f_form_unlabeled(doc):
    _titled(doc)
    doc.add_paragraph("Date of birth")
    _add_unlabeled_checkbox(doc.add_paragraph())
    # 3.3.2, NOT 4.1.2. The docx form-field fix reports "Labelled 1 form field(s) from adjacent
    # text · 3.3.2" — Labels or Instructions. Expecting 4.1.2 scored a working fix as a MISS on
    # every model, which reads as a product gap and is a manifest error.
    return ["3.3.2"]


def f_reading_order(doc):
    _titled(doc)
    # 1.3.2 — content whose visual order does not match its logical order: a continuation
    # paragraph placed before the passage it continues.
    doc.add_paragraph("…continued from the section below, which explains the deadline.")
    doc.add_paragraph("Section 1. Enrollment opens March 1.")
    return ["1.3.2"]


# ---- controls: deterministic criteria, identical for every model ----
def f_pseudo_heading(doc):
    _titled(doc)
    r = doc.add_paragraph().add_run("Eligibility")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph("All benefits-eligible employees may enrol.")
    return ["1.3.1"]


def f_table_no_header(doc):
    _titled(doc)
    t = doc.add_table(rows=3, cols=2)
    for i, (a, b) in enumerate([("Plan", "Cost"), ("Basic", "$0"), ("Plus", "$40")]):
        t.rows[i].cells[0].text, t.rows[i].cells[1].text = a, b
    return ["1.3.1"]


def f_low_contrast(doc):
    _titled(doc)
    r = doc.add_paragraph().add_run("Deadline: March 31, 2026.")
    r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    return ["1.4.3"]


def f_no_title_lang(doc):
    doc.add_heading("Enrollment", level=1)
    doc.add_paragraph("Enrollment runs March 1 to March 31.")
    return ["2.4.2", "3.1.1"]       # _titled deliberately NOT called; language stripped after


def f_heading_skip(doc):
    _titled(doc)
    doc.add_heading("Benefits", level=1)
    doc.add_heading("Dental", level=3)          # H1 -> H3, no H2
    doc.add_paragraph("Covered at 80%.")
    return ["2.4.6"]


def f_mixed(doc):
    """One realistic document carrying several issues — the case an estate actually contains.

    Kept LAST and scored separately: a mixed file cannot attribute a failure to one criterion,
    so it measures whether a model degrades when asked for several fixes at once, which the
    isolated fixtures cannot show.
    """
    _titled(doc)
    doc.add_heading("2026 Open Enrollment", level=1)
    p = doc.add_paragraph("Full policy: ")
    _add_hyperlink(p, "https://example.com/policy", "click here")
    doc.add_paragraph("Press the square button at the top to enrol.")
    doc.add_paragraph("Veuillez consulter la politique d’accessibilité avant le 31 mars 2026 "
                      "pour confirmer votre choix de couverture médicale et dentaire.")
    doc.add_picture(io.BytesIO(_png_of_text()))
    return ["1.1.1", "1.4.5", "2.4.4", "1.3.3", "3.1.2"]


BUILDERS = [
    ("clean", f_clean), ("alt-photo", f_alt_photo), ("alt-text-image", f_alt_text_image),
    ("link-generic", f_link_generic), ("link-raw-url", f_link_raw_url),
    ("sensory", f_sensory), ("lang-parts", f_lang_parts),
    ("form-unlabeled", f_form_unlabeled), ("reading-order", f_reading_order),
    ("pseudo-heading", f_pseudo_heading), ("table-no-header", f_table_no_header),
    ("low-contrast", f_low_contrast), ("no-title-lang", f_no_title_lang),
    ("heading-skip", f_heading_skip), ("mixed", f_mixed),
]


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--out", type=Path, default=ROOT / "test-corpus" / "model-eval")
    args = ap.parse_args()
    args.out.mkdir(parents=True, exist_ok=True)

    manifest = []
    for name, fn in BUILDERS:
        doc = Document()
        expects = fn(doc)
        path = args.out / f"{name}.docx"
        doc.save(str(path))
        if name == "no-title-lang":
            _strip_default_language(path)       # python-docx writes a default lang otherwise
        manifest.append({"file": path.name, "expect_criteria": expects,
                         "kb": round(path.stat().st_size / 1024, 1)})
        print(f"  {path.name:24} {', '.join(expects) or '(clean)'}")

    (args.out / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n")
    print(f"\n{len(manifest)} documents + manifest.json in {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
