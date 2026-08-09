#!/usr/bin/env python3
"""A LABELLED .docx corpus across the Core-17 criteria — the ground truth an ACP capability
benchmark needs, as opposed to a draft-quality benchmark.

WHY THIS EXISTS. The judge harness (judge_drafts.py) scores what a model WRITES. It cannot
answer the two questions that actually decide whether ACP ships: does the engine reach the right
verdict, and does a remediation make the violation go away without damaging the document. Both
need fixtures whose correct answer is known before anything runs. That is this file.

THE VOCABULARY IS NOT PASS/FAIL/N-A, AND ASSUMING IT IS WOULD INVALIDATE THE SCORES. ACP answers
four things — PASS, FAIL, REVIEW, NOT_EVALUATED — and which are reachable is a property of the
(criterion, format) pair, not of the document. On .docx, only FOUR of the fifteen Core-17 pairs
can ever return PASS:

    1.3.1  1.4.3  2.4.2  3.1.1        <- the deterministic lane; a full confusion matrix exists

Eight more are REVIEW-lane (1.1.1, 1.3.2, 1.3.3, 1.4.5, 2.4.4, 2.4.6, 3.1.2, 4.1.2): a review
detector does not certify conformance (ADR 0016), so a CLEAN file there resolves to REVIEW or
NOT_EVALUATED and NEVER to PASS. There is no true-negative class on those pairs, which means
"false positive rate" in the usual sense is undefined for them — a clean fixture does not score
as "correctly passed", because the engine is not permitted to say so. Three (1.4.1, 1.4.11,
2.1.2) have no .docx lane at all, and a fixture there exists to prove the engine stays silent
rather than inventing a verdict.

So each fixture declares its expectation per criterion, and EVERY declaration is checked against
corpus_expectations.possible_verdicts() at generation time. A fixture that expects PASS on 1.1.1
fails to build. Without that check the corpus would report a false failure on every run forever,
and it would look like a product defect rather than a manifest defect.

THE EDGE CASES ARE THE POINT. A corpus of obvious violations measures almost nothing: every
model finds a white-on-white heading. What separates them is the near-miss — contrast at 4.49
versus 4.50, a French passage one word under the length floor, a table whose header row is bold
but not marked, an image of text that is a logo. Each `edge` fixture below names the specific
wrong answer it is designed to catch, because a fixture whose failure mode is not written down
gets "fixed" by whoever next sees it fail.

FIXTURES ARE SINGLE-CRITERION ON PURPOSE. A document with six violations cannot tell you which
one a model missed. Attribution accuracy — did it blame the right SC — is only measurable when
the right answer is one criterion.

Run:
    python scripts/gen_sc_corpus.py --out ~/Downloads/acp-docx-eval/sc-corpus
"""
from __future__ import annotations

import argparse
import io
import json
import re
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from gen_demo_fixtures import _add_hyperlink, _add_unlabeled_checkbox  # noqa: E402
from gen_demo_fixtures import _strip_default_language  # noqa: E402

import corpus_expectations as ce  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402
from PIL import Image, ImageDraw  # noqa: E402

FMT = "docx"

# Contrast values COMPUTED from the WCAG relative-luminance formula, then verified against the
# engine. The AA floor is 4.5:1 for body text and 3.0:1 for large text (>=18pt, or >=14pt bold).
#
# The first version of this table was eyeballed and every entry was wrong — #595959 was labelled
# "~4.55:1" and is actually 7.00:1 (that is the AAA grey), #606060 was labelled "~4.05" and is
# 6.29:1. The corpus duly reported a 5.3% FALSE PASS RATE against the product, which is the
# single most alarming number this benchmark can produce, and it was entirely manufactured by
# this dict. Contrast fixtures are the easiest kind to get wrong by eye and the most expensive,
# because a wrong one accuses the engine of certifying an inaccessible document.
_ON_WHITE = {
    "aa_exact":  (0x76, 0x76, 0x76),   # 4.54:1 — clears the body floor by 0.04
    "aa_miss":   (0x7A, 0x7A, 0x7A),   # 4.31:1 — under the body floor
    "large_ok":  (0x94, 0x94, 0x94),   # 3.03:1 — fails as body text, PASSES at 18pt
    "invisible": (0xFF, 0xFF, 0xFF),   # 1.00:1 — white on white
}


# ── primitives ────────────────────────────────────────────────────────────────

def _solid_png(color=(27, 58, 107), size=(240, 120)) -> bytes:
    """A near-uniform block. Deliberately kept for the DECORATIVE fixtures only.

    proposals.infer_decorative treats a near-solid fill as decorative and returns a "Mark as
    decorative?" proposal WITHOUT consulting the vision model — correctly, since a solid block
    conveys nothing. Use `_content_png` for any fixture meant to exercise the vision lane; a
    corpus built on this one measures the decorative heuristic and reports it as a model result.
    """
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    return buf.getvalue()


def _content_png(seed: int = 0, size=(560, 320)) -> bytes:
    """A structured, non-uniform image with NO readable text — a chart, in effect.

    Three properties matter and each was learned by getting it wrong:
      - NOT near-uniform, or infer_decorative short-circuits before vision is called;
      - NOT hairline, tiny, or extreme-aspect, for the same reason (see infer_decorative's
        dimensional signals);
      - NO legible text, so the OCR-grounded path cannot transcribe it either. This is the
        case where a vision model is the ONLY thing that can produce an alt text, which makes
        it the only fixture shape that measures one.
    """
    im = Image.new("RGB", size, (250, 250, 252))
    d = ImageDraw.Draw(im)
    heights = [(i * 37 + seed * 53) % 200 + 40 for i in range(7)]
    for i, h in enumerate(heights):
        x = 40 + i * 70
        d.rectangle([x, size[1] - 40 - h, x + 44, size[1] - 40],
                    fill=(31 + i * 12, 78 + (i * 23) % 120, 160 - i * 9))
    d.line([(30, size[1] - 40), (size[0] - 20, size[1] - 40)], fill=(60, 60, 60), width=3)
    d.line([(30, 20), (30, size[1] - 40)], fill=(60, 60, 60), width=3)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()


def _text_png(lines: list[str], size=(900, 240)) -> bytes:
    """An image with real, OCR-legible text baked in.

    Scaled up after drawing because PIL's default bitmap font is too small for tesseract to
    read reliably at native size — an unreadable image-of-text fixture silently becomes a
    no-text fixture, and 1.4.5 stops firing for a reason that has nothing to do with the rule.
    """
    im = Image.new("RGB", size, "white")
    d = ImageDraw.Draw(im)
    y = 20
    for ln in lines:
        d.text((20, y), ln, fill="black")
        y += 70
    im = im.resize((size[0] * 2, size[1] * 2), Image.LANCZOS)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()


def _doc(title: str = "Patient Benefits Notice", lang: str = "en-US") -> Document:
    """A document that is clean on everything this corpus does not deliberately break.

    Title and language are set because `_ensure` (remediate_office) writes both when either is
    absent — so a fixture missing them collects two unrelated edits, and a remediation scorecard
    reads those as work the model did on a criterion nobody was testing.
    """
    d = Document()
    d.core_properties.title = title
    if lang:
        _set_doc_lang(d, lang)
    return d


def _set_doc_lang(d: Document, lang: str) -> None:
    styles = d.styles.element
    for el in styles.findall(qn("w:docDefaults")):
        rpr = el.find(qn("w:rPrDefault"))
        if rpr is None:
            continue
        inner = rpr.find(qn("w:rPr"))
        if inner is None:
            inner = OxmlElement("w:rPr")
            rpr.append(inner)
        for old in inner.findall(qn("w:lang")):
            inner.remove(old)
        node = OxmlElement("w:lang")
        node.set(qn("w:val"), lang)
        inner.append(node)


def _run(p, text: str, *, color=None, size_pt=None, bold=False, lang=None):
    r = p.add_run(text)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if size_pt:
        r.font.size = Pt(size_pt)
    r.bold = bold
    if lang:
        rpr = r._r.get_or_add_rPr()
        node = OxmlElement("w:lang")
        node.set(qn("w:val"), lang)
        rpr.append(node)
    return r


def _mark_header_row(table) -> None:
    """<w:tblHeader/> on row 0 — the machine-readable header marking.

    Bold text in row 0 is a VISUAL convention and a screen reader cannot announce columns from
    it. Emboldening alone is the classic fixture bug: the document looks compliant and is not.
    """
    trPr = table.rows[0]._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _add_table(d, rows: list[list[str]], *, header: bool):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            para = t.cell(i, j).paragraphs[0]
            _run(para, cell, bold=(i == 0))
    if header:
        _mark_header_row(t)
    return t


def _add_image(d, png: bytes, *, alt: str | None, decorative: bool = False, width_in=3.0):
    from docx.shared import Inches
    stream = io.BytesIO(png)
    run = d.add_paragraph().add_run()
    run.add_picture(stream, width=Inches(width_in))
    pic = run._r.find(qn("w:drawing"))
    if pic is None:
        return
    for docpr in pic.iter(qn("wp:docPr")):
        if decorative:
            # An empty descr ALONE is not a decorative declaration, and assuming it was is what
            # made this fixture report a false positive against the engine. OOXML marks
            # decorative with an <adec:decorative val="1"/> extension (apply_alt._DECORATIVE_MARK
            # writes exactly this); an empty or absent descr without it is indistinguishable
            # from an author who simply never filled the box in. ACP flagging the bare-empty
            # case is the conservative and correct reading — a decorative claim has to be made,
            # not inferred from an omission.
            docpr.set("descr", "")
            docpr.set("title", "")
            from docx.oxml import parse_xml
            docpr.append(parse_xml(
                '<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                '<a:ext uri="{C183D7F6-B498-43B3-948B-1728B52AA6E4}">'
                '<adec:decorative xmlns:adec="http://schemas.microsoft.com/office/drawing/'
                '2017/decorative" val="1"/></a:ext></a:extLst>'))
        elif alt is not None:
            docpr.set("descr", alt)
        else:
            docpr.attrib.pop("descr", None)


def _add_textbox(d, text: str, *, top_pt: int):
    """A FLOATING text box, positioned absolutely.

    1.3.2 on .docx is about content whose visual order differs from its XML order, and only a
    floating (anchored) shape can do that. An earlier reading-order fixture used ordinary
    paragraphs, which are in reading order by construction — so the detector was correct to stay
    silent and the fixture was scored as a miss.
    """
    p = d.add_paragraph()
    r = p.add_run()
    xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                      xmlns:v="urn:schemas-microsoft-com:vml">
      <v:shape style="position:absolute;margin-left:0;margin-top:{top_pt}pt;width:400pt;height:40pt">
        <v:textbox><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:txbxContent></v:textbox>
      </v:shape></w:pict>'''
    r._r.append(OxmlElement("w:rPr"))
    from docx.oxml import parse_xml
    r._r.append(parse_xml(xml))


def _strip_title(path: Path) -> None:
    """Remove dc:title from core.xml — 2.4.2 with no title to find."""
    with zipfile.ZipFile(path) as z:
        parts = {n: z.read(n) for n in z.namelist()}
    if "docProps/core.xml" in parts:
        parts["docProps/core.xml"] = re.sub(
            rb"<dc:title>.*?</dc:title>", b"<dc:title></dc:title>", parts["docProps/core.xml"])
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, data in parts.items():
            z.writestr(n, data)


# ── fixtures ──────────────────────────────────────────────────────────────────
#
# Each builder returns (expectations, note). `expectations` maps SC -> the verdict the ENGINE
# should reach for this document. Only criteria the fixture says something about are listed;
# everything else is checked separately as "must not newly fail".

def f_alt_missing(d):
    d.add_heading("Coverage Summary", level=1)
    d.add_paragraph("Your 2026 plan includes the services shown below.")
    _add_image(d, _content_png(1), alt=None)
    return {"1.1.1": ce.FAIL}, "one image, no descr attribute at all"


def f_alt_junk(d):
    d.add_heading("Coverage Summary", level=1)
    _add_image(d, _content_png(2), alt="image1.png")
    return {"1.1.1": ce.FAIL}, "alt is the FILENAME — present but worthless"


def f_alt_whitespace(d):
    d.add_heading("Coverage Summary", level=1)
    _add_image(d, _content_png(3), alt="   ")
    return {"1.1.1": ce.FAIL}, ("EDGE: alt is whitespace. Must NOT be read as a decorative "
                                "declaration — an empty descr is a deliberate 'this is "
                                "decorative', three spaces is a bug")


def f_alt_decorative(d):
    d.add_heading("Coverage Summary", level=1)
    d.add_paragraph("A divider rule follows; it carries no information.")
    _add_image(d, _solid_png(color=(200, 200, 200), size=(400, 8)), alt=None, decorative=True)
    return {}, ("EDGE/adversarial: correctly marked decorative (descr=\"\"). Any 1.1.1 finding "
                "here is a FALSE POSITIVE — the case that punishes a model for over-flagging")


def f_alt_image_of_text(d):
    d.add_heading("Enrollment Notice", level=1)
    # ocr._MIN_WORDS is 10 for 1.4.5 (AA) and 3 for 1.4.9 (AAA), counted on OCR'd tokens rather
    # than on what was drawn. The first version of this image rendered three short lines and OCR
    # recovered SIX words — dates and a phone number are not words — so 1.4.9 fired, 1.4.5 did
    # not, and the fixture read as an engine bug. Prose, and enough of it to clear the floor with
    # room for the odd OCR merge ("to enroll" came back as "toenroll").
    _add_image(d, _text_png([
        "Open enrollment for the coming plan year",
        "begins on the first of March and closes",
        "on the last day of the same month. Call",
        "the benefits office to change your plan.",
    ], size=(900, 360)), alt=None)
    return {"1.1.1": ce.FAIL, "1.4.5": ce.FAIL}, \
        "image of prose, no alt — must fire on BOTH the AA and the alt-text criterion"


def f_alt_logo(d):
    d.add_heading("Enrollment Notice", level=1)
    _add_image(d, _text_png(["UT Health"], size=(300, 100)), alt="UT Health", width_in=1.2)
    return {}, ("ADVERSARIAL: a small logo with its text as the alt. WCAG 1.4.5 exempts logos "
                "and the alt is correct, so a 1.4.5 or 1.1.1 finding here is a false positive")


def f_table_no_header(d):
    d.add_heading("Plan Comparison", level=1)
    _add_table(d, [["Plan", "Premium", "Deductible"],
                   ["Bronze", "$120", "$3,000"],
                   ["Silver", "$185", "$1,500"]], header=False)
    return {"1.3.1": ce.FAIL}, "row 0 is bold but carries no <w:tblHeader/>"


def f_table_header_ok(d):
    d.add_heading("Plan Comparison", level=1)
    _add_table(d, [["Plan", "Premium", "Deductible"],
                   ["Bronze", "$120", "$3,000"]], header=True)
    return {"1.3.1": ce.PASS}, "properly marked header row — one of the four PASS-capable cells"


def f_table_nested(d):
    d.add_heading("Plan Comparison", level=1)
    outer = _add_table(d, [["Plan", "Detail"], ["Bronze", ""]], header=True)
    cell = outer.cell(1, 1)
    inner = cell.add_table(rows=2, cols=2)
    inner.style = "Table Grid"
    for i, row in enumerate([["Service", "Copay"], ["Office visit", "$30"]]):
        for j, val in enumerate(row):
            _run(inner.cell(i, j).paragraphs[0], val, bold=(i == 0))
    return {"1.3.1": ce.FAIL}, ("EDGE: outer table is correct, NESTED table has no header row. "
                                "Catches a detector that stops at the top level")


def f_contrast_boundary_pass(d):
    d.add_heading("Notice", level=1)
    _run(d.add_paragraph(), "Your coverage begins on March 1.", color=_ON_WHITE["aa_exact"])
    return {"1.4.3": ce.PASS}, "EDGE: ~4.55:1 — just OVER the 4.5 AA floor, must not fire"


def f_contrast_boundary_fail(d):
    d.add_heading("Notice", level=1)
    _run(d.add_paragraph(), "Your coverage begins on March 1.", color=_ON_WHITE["aa_miss"])
    return {"1.4.3": ce.FAIL}, "EDGE: ~4.05:1 — just UNDER the floor. The pair above is the test"


def f_contrast_large_text(d):
    d.add_heading("Notice", level=1)
    # A SENTENCE, not a phrase. The first version set a short standalone 18pt line and ACP
    # flagged 1.3.1 DOCX_PSEUDO_HEADING — correctly: a large, short, isolated line is visually a
    # heading and is not marked as one. The fixture had accidentally built a second, real defect
    # and then scored the engine wrong for finding it.
    _run(d.add_paragraph(),
         "Your coverage begins on the first of March and continues through the end of the "
         "plan year, provided premiums are paid on schedule.",
         color=_ON_WHITE["large_ok"], size_pt=18)
    return {"1.4.3": ce.PASS}, ("EDGE: 3.03:1 at 18pt. Fails the BODY threshold and passes the "
                                "LARGE-TEXT one — catches a checker that applies 4.5 to everything")


def f_contrast_invisible(d):
    d.add_heading("Notice", level=1)
    _run(d.add_paragraph(), "Call 214-555-0134 before March 31.", color=_ON_WHITE["invisible"])
    return {"1.4.3": ce.FAIL}, "white on white, 1:1 — the unmissable case, as a floor check"


def f_no_title(d):
    d.add_heading("Notice", level=1)
    d.add_paragraph("Your coverage begins on March 1.")
    return {"2.4.2": ce.FAIL}, "dc:title emptied post-build"


def f_title_ok(d):
    d.add_heading("Notice", level=1)
    d.add_paragraph("Your coverage begins on March 1.")
    return {"2.4.2": ce.PASS}, "title present — PASS-capable cell"


def f_no_lang(d):
    d.add_heading("Notice", level=1)
    d.add_paragraph("Your coverage begins on March 1.")
    return {"3.1.1": ce.FAIL}, "every <w:lang> stripped post-build"


def f_lang_ok(d):
    d.add_heading("Notice", level=1)
    d.add_paragraph("Your coverage begins on March 1.")
    return {"3.1.1": ce.PASS}, "en-US declared — PASS-capable cell"


def f_lang_parts_long(d):
    d.add_heading("Notice", level=1)
    # BOTH paragraphs must clear textchecks._MIN_SEG_WORDS (12), not just the foreign one. The
    # detector needs two CONFIDENT segments in different languages; the first version gave the
    # French passage 17 words and the English one 8, so only one segment qualified, nothing
    # could be compared, and the silence was correct. This is the third time that 12-word floor
    # has been mistaken for a product defect.
    d.add_paragraph("Please review the accessibility policy for the coming plan year before "
                    "the enrollment deadline at the end of March.")
    d.add_paragraph("Veuillez consulter la politique d'accessibilite de notre etablissement "
                    "avant le trente et un mars prochain sans faute.")
    return {"3.1.2": ce.FAIL}, ("a 17-word French passage next to a 20-word English one — both "
                                "over the 12-word floor, so two confident segments exist")


def f_lang_parts_short(d):
    d.add_heading("Notice", level=1)
    d.add_paragraph("Please review the accessibility policy before the deadline of March 31.")
    d.add_paragraph("Veuillez consulter la politique avant le trente et un mars.")
    return {}, ("EDGE: a 10-word French passage — UNDER the 12-word floor, where langdetect is "
                "unreliable. Silence here is CORRECT. Two earlier fixtures failed on exactly "
                "this and the product was right both times")


def f_lang_parts_marked(d):
    d.add_heading("Notice", level=1)
    d.add_paragraph("Please review the accessibility policy before the deadline.")
    p = d.add_paragraph()
    _run(p, "Veuillez consulter la politique d'accessibilite de notre etablissement "
            "avant le trente et un mars prochain sans faute.", lang="fr-FR")
    return {}, "ADVERSARIAL: the same long French passage, correctly marked fr-FR — silence"


def f_link_click_here(d):
    d.add_heading("Notice", level=1)
    p = d.add_paragraph("For the full accessibility policy, ")
    _add_hyperlink(p, "https://example.org/policy", "click here")
    return {"2.4.4": ce.FAIL}, "the canonical non-descriptive link text"


def f_link_ambiguous_pair(d):
    d.add_heading("Notice", level=1)
    p1 = d.add_paragraph("For the medical plan, ")
    _add_hyperlink(p1, "https://example.org/medical", "click here")
    p2 = d.add_paragraph("For the dental plan, ")
    _add_hyperlink(p2, "https://example.org/dental", "click here")
    return {"2.4.4": ce.FAIL}, ("EDGE: TWO links with identical text and different targets — "
                                "should report both, not collapse to one finding")


def f_link_bare_url(d):
    d.add_heading("Notice", level=1)
    p = d.add_paragraph("Full policy: ")
    _add_hyperlink(p, "https://example.org/accessibility-policy-2026-final-v3",
                   "https://example.org/accessibility-policy-2026-final-v3")
    return {"2.4.4": ce.FAIL}, ("EDGE: a bare URL as link text. Technically unique, unreadable "
                                "aloud — a detector that only greps for 'click here' misses it")


def f_link_descriptive(d):
    d.add_heading("Notice", level=1)
    p = d.add_paragraph("Read the ")
    _add_hyperlink(p, "https://example.org/policy", "2026 accessibility policy")
    return {}, "ADVERSARIAL: good link text — a 2.4.4 finding here is a false positive"


def f_sensory(d):
    d.add_heading("How to enroll", level=1)
    d.add_paragraph("Click the round green button on the right to continue.")
    return {"1.3.3": ce.FAIL}, "shape + colour + position, no text label"


def f_sensory_benign(d):
    d.add_heading("How to enroll", level=1)
    d.add_paragraph("Complete the enrollment form and return it to Member Services.")
    return {}, "ADVERSARIAL: no sensory dependence — silence expected"


def f_heading_empty(d):
    d.add_heading("Notice", level=1)
    d.add_heading("", level=2)
    d.add_paragraph("Your coverage begins on March 1.")
    return {"2.4.6": ce.FAIL}, \
        ("EDGE: an EMPTY H2. This was the one miss in this corpus that belonged to the product "
         "rather than the fixture — the outline walk read pStyle refs and never the heading's "
         "text, so an empty heading was in the outline, broke no level sequence and had no runs "
         "to fail contrast on, and produced ZERO findings. Kept red until DOCX_HEADING_EMPTY "
         "existed rather than relaxed to green, which is how a gap becomes permanent.")


def f_heading_skip(d):
    d.add_heading("Notice", level=1)
    d.add_heading("Deductibles", level=3)
    d.add_paragraph("Your deductible is $1,500.")
    # ACP reports this on BOTH criteria — DOCX-HEAD-001 under 1.3.1 (the outline relationship is
    # broken) and DOCX_HEADING_SKIP under 2.4.6 (the heading set no longer describes the
    # structure). Expecting only 1.3.1 scored the 2.4.6 finding as a false positive. Both
    # readings are defensible and the engine makes both, so the expectation says so.
    return {"1.3.1": ce.FAIL, "2.4.6": ce.FAIL}, \
        ("EDGE: H1 -> H3, no H2. A single defect that legitimately lands on two criteria — the "
         "case that punishes a scorer for assuming one violation means one finding")


def f_reading_order(d):
    d.add_heading("Notice", level=1)
    _add_textbox(d, "SECOND: return the form by March 31.", top_pt=120)
    _add_textbox(d, "FIRST: complete every section of the form.", top_pt=40)
    d.add_paragraph("Follow the steps shown.")
    return {"1.3.2": ce.FAIL}, ("floating text boxes whose VISUAL order (top 40pt before top "
                                "120pt) is the reverse of their XML order")


def f_unlabeled_field(d):
    d.add_heading("Consent", level=1)
    p = d.add_paragraph("I consent to the release of my records ")
    _add_unlabeled_checkbox(p)
    # Measured, not assumed: ACP fires THREE findings here — 3.3.2 (no label), 4.1.2 (no
    # accessible name) and 2.1.2 as REVIEW (an interactive control exists; verify focus can
    # leave it). An earlier note in this corpus claimed ACP reports 3.3.2 "NOT 4.1.2"; that was
    # read off a differently-shaped fixture and was stale. All three are correct, and 2.1.2
    # surfacing as REVIEW is the review lane doing exactly what it is for.
    return {"4.1.2": ce.FAIL, "2.1.2": ce.REVIEW}, \
        ("EDGE: an unlabeled checkbox — one defect, three criteria, across two lanes")


def f_no_docx_lane(d):
    d.add_heading("Notice", level=1)
    d.add_paragraph("Press and hold the slider, then drag it to set your contribution.")
    _run(d.add_paragraph(), "Contribution level", color=(0x11, 0x11, 0x11))
    return {"1.4.1": ce.NOT_EVALUATED, "2.1.2": ce.NOT_EVALUATED,
            "1.4.11": ce.NOT_EVALUATED}, \
        ("CONTROL: content that WOULD implicate 1.4.1 / 1.4.11 / 2.1.2 if a .docx lane existed. "
         "None does. The engine must stay silent — a verdict here is invention, and this is the "
         "fixture that catches an LLM assessor asked to 'score all 17'")


def f_clean(d):
    d.add_heading("Patient Benefits Notice", level=1)
    d.add_paragraph("Your 2026 coverage begins on March 1.")
    d.add_heading("Plan comparison", level=2)
    _add_table(d, [["Plan", "Premium"], ["Bronze", "$120"]], header=True)
    p = d.add_paragraph("Read the ")
    _add_hyperlink(p, "https://example.org/policy", "2026 accessibility policy")
    _add_image(d, _solid_png(), alt="Bar chart comparing Bronze and Silver premiums.")
    return {"1.3.1": ce.PASS, "1.4.3": ce.PASS, "2.4.2": ce.PASS, "3.1.1": ce.PASS}, \
        "the false-positive control: nothing wrong, every PASS-capable cell should certify"


# name -> (builder, post-processing hook, kind)
FIXTURES = [
    ("alt-missing",            f_alt_missing,          None,            "violation"),
    ("alt-junk",               f_alt_junk,             None,            "violation"),
    ("alt-whitespace",         f_alt_whitespace,       None,            "edge"),
    ("alt-decorative-ok",      f_alt_decorative,       None,            "adversarial"),
    ("alt-image-of-text",      f_alt_image_of_text,    None,            "violation"),
    ("alt-logo-ok",            f_alt_logo,             None,            "adversarial"),
    ("table-no-header",        f_table_no_header,      None,            "violation"),
    ("table-header-ok",        f_table_header_ok,      None,            "clean"),
    ("table-nested",           f_table_nested,         None,            "edge"),
    ("contrast-4.55-ok",       f_contrast_boundary_pass, None,          "edge"),
    ("contrast-4.05-fail",     f_contrast_boundary_fail, None,          "edge"),
    ("contrast-large-ok",      f_contrast_large_text,  None,            "edge"),
    ("contrast-invisible",     f_contrast_invisible,   None,            "violation"),
    ("title-missing",          f_no_title,             _strip_title,    "violation"),
    ("title-ok",               f_title_ok,             None,            "clean"),
    ("lang-missing",           f_no_lang,        _strip_default_language, "violation"),
    ("lang-ok",                f_lang_ok,              None,            "clean"),
    ("lang-parts-long",        f_lang_parts_long,      None,            "violation"),
    ("lang-parts-short-ok",    f_lang_parts_short,     None,            "edge"),
    ("lang-parts-marked-ok",   f_lang_parts_marked,    None,            "adversarial"),
    ("link-click-here",        f_link_click_here,      None,            "violation"),
    ("link-ambiguous-pair",    f_link_ambiguous_pair,  None,            "edge"),
    ("link-bare-url",          f_link_bare_url,        None,            "edge"),
    ("link-descriptive-ok",    f_link_descriptive,     None,            "adversarial"),
    ("sensory-shape-colour",   f_sensory,              None,            "violation"),
    ("sensory-benign-ok",      f_sensory_benign,       None,            "adversarial"),
    ("heading-empty",          f_heading_empty,        None,            "edge"),
    ("heading-level-skip",     f_heading_skip,         None,            "edge"),
    ("reading-order",          f_reading_order,        None,            "violation"),
    ("field-unlabeled",        f_unlabeled_field,      None,            "edge"),
    ("no-docx-lane",           f_no_docx_lane,         None,            "control"),
    ("clean",                  f_clean,                None,            "clean"),
]


def _validate(name: str, expectations: dict[str, str]) -> list[str]:
    """Reject an expectation the engine could never produce.

    This is the whole reason corpus_expectations exists. A manifest that expects PASS on a
    REVIEW-lane pair does not describe a product bug; it describes a manifest bug, and it will
    report a false failure on every run until somebody re-derives this by hand.
    """
    problems = []
    for sc, verdict in expectations.items():
        allowed = ce.possible_verdicts(sc, FMT)
        if verdict not in allowed:
            problems.append(
                f"{name}: expects {sc}={verdict} but the engine can only emit "
                f"{','.join(sorted(allowed))} for ({sc}, {FMT})")
    return problems


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--out", type=Path,
                    default=Path.home() / "Downloads" / "acp-docx-eval" / "sc-corpus")
    args = ap.parse_args()
    docs = args.out / "docs"
    docs.mkdir(parents=True, exist_ok=True)

    manifest, problems = [], []
    for name, build, post, kind in FIXTURES:
        d = _doc()
        expectations, note = build(d)
        path = docs / f"{name}.docx"
        d.save(path)
        if post:
            post(path)
        problems += _validate(name, expectations)
        manifest.append({
            "file": f"docs/{name}.docx",
            "name": name,
            "kind": kind,
            "format": FMT,
            "expect": expectations,
            "note": note,
        })
        print(f"  {kind:11} {name:24} {expectations or '(silence expected)'}")

    if problems:
        print("\nIMPOSSIBLE EXPECTATIONS — fix the fixture, not the engine:", file=sys.stderr)
        for p in problems:
            print(f"  {p}", file=sys.stderr)
        return 1

    # Record the reachable vocabulary alongside the fixtures. A scorer reading this file must not
    # have to re-derive which criteria can certify — that is precisely the derivation that goes
    # wrong by hand.
    scs = sorted(sc for sc, fmts in ce.pol.SCOPE_PRESETS["acp-core-17"].items() if FMT in fmts)
    lanes = {sc: {"clean_verdict": ce.clean_verdict(sc, FMT),
                  "possible": sorted(ce.possible_verdicts(sc, FMT)),
                  "can_pass": ce.can_ever_pass(sc, FMT)}
             for sc in scs}

    out = {"format": FMT, "fixtures": manifest, "lanes": lanes}
    (args.out / "manifest.json").write_text(json.dumps(out, indent=2) + "\n")
    kinds: dict[str, int] = {}
    for f in manifest:
        kinds[f["kind"]] = kinds.get(f["kind"], 0) + 1
    print(f"\n{len(manifest)} fixtures: " + ", ".join(f"{v} {k}" for k, v in sorted(kinds.items())))
    print(f"wrote {args.out / 'manifest.json'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
