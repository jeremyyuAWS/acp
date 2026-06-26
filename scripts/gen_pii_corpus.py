"""Generate synthetic-PII test documents to exercise the PII detector (ADR 0006).

All values are SYNTHETIC and clearly fake: standard card-network *test* numbers,
example.com / reserved 555-01xx contacts, and fake-but-valid-format SSNs. No real
person's data is used. Output is a set of HTML/DOCX/PDF files the scanner's PII
dimension will flag (plus one clean control with no PII).

  python3 scripts/gen_pii_corpus.py --out /tmp/acp-pii-corpus

Formats chosen because the scanner only runs detection on HTML/Office/PDF.
"""
from __future__ import annotations
import argparse
from pathlib import Path

# ── synthetic PII (all fake) ──────────────────────────────────────────────────
SSN1, SSN2, SSN3 = "123-45-6789", "078-05-1120", "219-09-9998"
VISA, MC, AMEX = "4111 1111 1111 1111", "5555 5555 5555 4444", "3782 822463 10005"
EMAILS = ["jane.doe@example.com", "john.smith@example.org", "a.patel@example.net"]
PHONES = ["(415) 555-0142", "212-555-0173", "650-555-0118"]


def _html(title: str, body: str) -> str:
    return (f"<!DOCTYPE html><html lang=\"en\"><head><title>{title}</title></head>"
            f"<body><h1>{title}</h1>{body}</body></html>")


def gen_html(out: Path) -> list[Path]:
    files = []
    hr = _html("Employee Onboarding — CONFIDENTIAL", f"""
      <p>New hire record (synthetic test data).</p>
      <ul>
        <li>Name: Jane Doe</li>
        <li>SSN: {SSN1}</li>
        <li>Email: {EMAILS[0]}</li>
        <li>Phone: {PHONES[0]}</li>
        <li>Emergency contact: {EMAILS[1]}, {PHONES[1]}</li>
      </ul>""")
    p = out / "PII-TEST-hr-onboarding.html"; p.write_text(hr); files.append(p)

    fin = _html("Vendor Invoice — Payment Details", f"""
      <p>Card on file for recurring billing (synthetic test cards).</p>
      <table>
        <tr><td>Visa</td><td>{VISA}</td></tr>
        <tr><td>Mastercard</td><td>{MC}</td></tr>
        <tr><td>Amex</td><td>{AMEX}</td></tr>
      </table>
      <p>Billing contact: {EMAILS[2]} / {PHONES[2]}</p>""")
    p = out / "PII-TEST-finance-invoice.html"; p.write_text(fin); files.append(p)

    clean = _html("Quarterly Accessibility Summary", """
      <p>This report contains no personal data. It summarizes remediation progress
      across the document estate for the quarter.</p>
      <h2>Highlights</h2><ul><li>Coverage up 12%</li><li>Zero critical regressions</li></ul>""")
    p = out / "PII-TEST-clean-control.html"; p.write_text(clean); files.append(p)
    return files


def gen_docx(out: Path) -> list[Path]:
    from docx import Document
    doc = Document()
    doc.add_heading("Patient Intake Form — PHI (synthetic)", 0)
    for line in [
        "Patient: John Smith",
        f"SSN: {SSN2}",
        f"Email: {EMAILS[1]}",
        f"Phone: {PHONES[1]}",
        f"Secondary contact SSN: {SSN3}",
        "Insurance: synthetic test record — not real data.",
    ]:
        doc.add_paragraph(line)
    p = out / "PII-TEST-patient-intake.docx"; doc.save(str(p))
    return [p]


def gen_pdf(out: Path) -> list[Path]:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    p = out / "PII-TEST-customer-export.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    y = 740
    rows = [
        "Customer export — CONFIDENTIAL (synthetic test data)",
        "",
        f"1. Jane Doe   SSN {SSN1}   {EMAILS[0]}   {PHONES[0]}   Visa {VISA}",
        f"2. John Smith SSN {SSN2}   {EMAILS[1]}   {PHONES[1]}   MC {MC}",
        f"3. A. Patel   SSN {SSN3}   {EMAILS[2]}   {PHONES[2]}   Amex {AMEX}",
        "",
        "All values above are fabricated for detection testing.",
    ]
    for line in rows:
        c.drawString(54, y, line); y -= 22
    c.showPage(); c.save()
    return [p]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default="/tmp/acp-pii-corpus")
    args = ap.parse_args()
    out = Path(args.out); out.mkdir(parents=True, exist_ok=True)
    made = gen_html(out) + gen_docx(out) + gen_pdf(out)
    print(f"Generated {len(made)} synthetic-PII test documents in {out}:")
    for p in made:
        print(f"  {p.name}  ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
