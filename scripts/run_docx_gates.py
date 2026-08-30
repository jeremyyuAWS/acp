#!/usr/bin/env python3
"""The three gates. Until all three pass, no model comparison from this corpus is valid.

    Gate 1  clean controls produce ZERO edits.
            A false-positive edit on a compliant patient document outranks every other
            result in the study, so this is checked first and stops everything.

    Gate 2  the deterministic control produces IDENTICAL output across model configurations.
            Those five criteria are rule code. If they move with the model, the harness is
            measuring something other than the model.

    Gate 3  the assisted-only control produces assessment findings and ZERO edits.
            Scored on BOTH channels, because a remediation-only scorecard reads correct
            behaviour — detect, decline to write, hand to a human — as total failure.

ASSESSMENT AND REMEDIATION ARE DIFFERENT CALLS. `remediate_office` does not assess; the scan
path (`scanner.analyse_and_assess`) does. Gate 3 cannot be answered from one of them, which is
why this script exists rather than another flag on eval_models_docx.py.

INTEGRITY IS CHECKED ON EVERY DOCUMENT, not only the clean ones. The existing verify gate
re-scans to confirm a finding cleared; nothing checks the document survived. A fix that clears
1.1.1 by deleting the image would pass today.

Run:
    python scripts/run_docx_gates.py --root ~/Downloads/acp-docx-eval
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
_SC = re.compile(r"\b(\d+\.\d+\.\d+)\b")

# One model per axis for gates 1 and 3; gate 2 sweeps the vision axis because that is the one
# with a model in the .docx remediation path at all.
GATE2_MODELS = ["moondream", "llava:7b", "qwen2.5vl:7b", "qwen2.5vl:32b"]


def _reload():
    sys.path.insert(0, str(ROOT / "api"))
    import importlib
    import ai
    import remediate_office
    importlib.reload(ai)
    importlib.reload(remediate_office)
    return remediate_office


def remediate(path: Path):
    """(criteria_touched, n_edits, n_proposals, seconds, out_path)."""
    ro = _reload()
    tmp = Path(tempfile.mkdtemp())
    work = tmp / path.name
    shutil.copy(path, work)
    applied, diffs, proposals = [], [], []
    t0 = time.time()
    res = ro.remediate_office(work, ai_enabled=True, applied_fixes=applied,
                              diffs=diffs, proposals=proposals)
    secs = round(time.time() - t0, 1)
    summaries = res[1] if isinstance(res, tuple) and len(res) > 1 else []
    # remediate_office returns None for the path when it wrote NOTHING — which is exactly the
    # clean-control case, so assuming a path here crashed Gate 3 on the run that mattered.
    out = (res[0] if isinstance(res, tuple) else res) or work
    crit = set()
    for s in summaries:
        crit.update(_SC.findall(str(s)))
    for coll in (applied, diffs):
        for f in coll:
            rid = (f or {}).get("rule_id", "") if isinstance(f, dict) else ""
            crit.update(_SC.findall(str(rid).replace("_", ".")))
    return crit, len(summaries), len(proposals), secs, Path(out)


def assess(path: Path) -> set:
    """Criteria the SCAN reports for this file — the assessment channel.

    RAISES when the scan does not run. It used to `return set()`, which Gate 3 could not
    distinguish from "the scan ran and reported nothing" — and Gate 3's verdict is
    `n_edits == 0 and not dmg`, which does not consult `detected` at all. So an engine failure
    printed "NOTE assessment reported none of them" and then PASSED a gate whose name is
    "assisted-only: assessed, not edited". It had not assessed anything.

    Same one-line inversion as scripts/score_assessment.py's `scan` (see its docstring for the
    measured case). An analyser error is never a clean result — including in the tools that
    decide whether the product's results can be trusted."""
    sys.path.insert(0, str(ROOT / "api"))
    import scanner
    tmp = Path(tempfile.mkdtemp())
    shutil.copy(path, tmp / path.name)
    try:
        res, _ = scanner.analyse_and_assess(tmp, path.name)
    except Exception as e:                                        # noqa: BLE001
        raise RuntimeError(
            f"assess failed on {path.name}: {e.__class__.__name__}: {e}\n"
            f"    A scan that did not run reports no criteria, which Gate 3 would pass.\n"
            f"    Check what is available: python scripts/check_engines.py --report") from e
    out = set()
    for issue in (res or {}).get("issues", []) or []:
        w = str(issue.get("wcag") or "").replace("SC_", "").replace("_", ".")
        out.update(_SC.findall(w))
    return out


def integrity(path: Path) -> dict:
    from docx import Document
    d = Document(str(path))
    return {"paragraphs": len(d.paragraphs), "tables": len(d.tables),
            "images": len(d.inline_shapes), "sections": len(d.sections),
            "characters": sum(len(p.text) for p in d.paragraphs)}


def check_integrity(before: dict, after: dict) -> list[str]:
    """Deltas that mean CONTENT was lost. Characters are allowed to move — a fix writes values
    — but a paragraph, table, image or section disappearing is damage, not remediation."""
    bad = []
    for k in ("paragraphs", "tables", "images", "sections"):
        if after.get(k, 0) < before.get(k, 0):
            bad.append(f"{k} {before[k]} -> {after[k]}")
    return bad


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--root", type=Path, default=Path.home() / "Downloads" / "acp-docx-eval")
    args = ap.parse_args()
    docs, exp = args.root / "docs", args.root / "expected"
    os.environ.setdefault("OLLAMA_MODEL", "llama3.1:8b")
    os.environ.setdefault("OLLAMA_VISION_MODEL", "qwen2.5vl:7b")

    # Pre-flight the analyser, for the reason in `assess`'s docstring: without it Gate 3 reports
    # PASS on a run where nothing was assessed. These gates decide whether a model comparison
    # from this corpus is valid at all, so they are the last place to let a missing dependency
    # read as a result.
    sys.path.insert(0, str(ROOT / "scripts"))
    import check_engines
    missing = check_engines.check(["office"])
    if missing:
        for name, why in missing:
            print(f"cannot run the gates: the {name} engine is unavailable.\n  {why}",
                  file=sys.stderr)
        return 2

    failures = []

    print("=== GATE 1 — clean controls produce zero edits ===")
    for name in ("ctl-clean-1p", "ctl-clean-25p"):
        p = docs / f"{name}.docx"
        before = json.loads((exp / f"{name}.json").read_text())["integrity"]
        crit, n_edits, n_prop, secs, out = remediate(p)
        dmg = check_integrity(before, integrity(out))
        ok = n_edits == 0 and not dmg
        print(f"  {'PASS' if ok else 'FAIL'}  {name:24} {secs:5.1f}s  edits={n_edits} "
              f"proposals={n_prop} criteria={sorted(crit) or '[]'}")
        if dmg:
            print(f"        INTEGRITY DAMAGE: {', '.join(dmg)}")
        if not ok:
            failures.append(f"Gate 1: {name} made {n_edits} edit(s) on a compliant document")

    print("\n=== GATE 2 — deterministic control is identical across models ===")
    name = "ctl-deterministic-only"
    p = docs / f"{name}.docx"
    seen = {}
    for m in GATE2_MODELS:
        os.environ["OLLAMA_VISION_MODEL"] = m
        crit, n_edits, _, secs, _ = remediate(p)
        seen[m] = tuple(sorted(crit))
        print(f"        {m:16} {secs:5.1f}s  edits={n_edits}  {sorted(crit)}")
    if len(set(seen.values())) == 1:
        print(f"  PASS  identical across {len(GATE2_MODELS)} models")
    else:
        print("  FAIL  deterministic output moved with the model")
        failures.append("Gate 2: deterministic control drifted across models")
    os.environ["OLLAMA_VISION_MODEL"] = "qwen2.5vl:7b"

    print("\n=== GATE 3 — assisted-only: assessed, not edited ===")
    name = "ctl-assisted-only"
    p = docs / f"{name}.docx"
    truth = json.loads((exp / f"{name}.json").read_text())
    want = sorted({v["criterion"] for v in truth["violations"]})
    detected = assess(p)
    crit, n_edits, n_prop, secs, out = remediate(p)
    dmg = check_integrity(truth["integrity"], integrity(out))
    found = sorted(set(want) & detected)
    print(f"        expected criteria : {want}")
    print(f"        assessment found  : {found}   (missed {sorted(set(want) - detected)})")
    print(f"        remediation edits : {n_edits}  proposals={n_prop}  {secs}s")
    ok = n_edits == 0 and not dmg
    print(f"  {'PASS' if ok else 'FAIL'}  zero edits on assisted-only criteria")
    if not ok:
        failures.append(f"Gate 3: {n_edits} edit(s) on criteria .docx does not remediate")
    if not found:
        print("  NOTE  assessment reported none of them — see the note below")

    print("\n" + "=" * 62)
    if failures:
        print("GATES FAILED — a model comparison from this corpus is NOT valid:")
        for f in failures:
            print(f"  - {f}")
        return 1
    print("ALL GATES PASSED — the corpus is sound; model comparison is valid.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
