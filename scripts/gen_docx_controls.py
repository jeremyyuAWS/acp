#!/usr/bin/env python3
"""The four gate controls for the .docx evaluation, plus located ground truth.

These are not fixtures among fixtures. Every other measurement in the study is contingent on
them, so each exists to make one specific failure impossible to miss:

    ctl-clean-1p            a compliant document. ANY edit is a false positive.
    ctl-clean-25p           the same, at scale — catches false positives that only appear
                            once there is enough content to misread.
    ctl-deterministic-only  the five rule-code criteria. Must be IDENTICAL for every model.
    ctl-assisted-only       criteria ACP assesses but does not remediate on .docx. Must
                            produce assessment findings and ZERO edits.

The clean documents are the fiddly ones. `_ensure` (remediate_office) sets a document title and
language when either is absent, so a "clean" file missing them collects two edits and Gate 1
fails for a reason that has nothing to do with the product. Both are set explicitly, and the
heading outline, contrast, table headers and link text are all deliberately correct.

Ground truth is written per document as LOCATED violations — `expected/<name>.json` — so a file
with five broken links cannot score as correct because one was found. Integrity counts are
captured at generation time, which is the only moment they are known to be right.

Run:
    python scripts/gen_docx_controls.py --out ~/Downloads/acp-docx-eval
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from gen_demo_fixtures import _add_hyperlink, _strip_default_language  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402


def _mark_header_row(table):
    """Actually MARK row 0 as a header — <w:tblHeader/> — not merely embolden it.

    Gate 1 failed on the first run because the clean control's table had bold text in row 0 and
    no tblHeader element. ACP flagged 1.3.1 and it was RIGHT: bold text is a visual convention,
    and a screen reader cannot announce column headers from it. The generator was asserting a
    header row existed while producing a table that has none.

    Worth stating plainly, because it is the same class of error the study exists to catch: a
    document that LOOKS compliant and is not. Had this been left in, Gate 1 would have reported
    a false positive against the product for a defect in the fixture.
    """
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)

LOREM = ("Employees may review coverage options during the annual enrollment window. "
         "Elections remain in effect for the plan year unless a qualifying life event occurs. "
         "Contact the benefits office for assistance with enrollment questions.")


def _compliant_base(doc, title):
    """Everything a clean .docx needs so remediation finds nothing to do.

    Each line here corresponds to a criterion that would otherwise fire: title (2.4.2),
    language (3.1.1), a real heading rather than bold text (1.3.1), and default-coloured body
    text (1.4.3). Omitting any one of them turns Gate 1 into a test of this generator.
    """
    doc.core_properties.title = title
    doc.core_properties.language = "en-US"
    doc.add_heading(title, level=1)


def build_clean(doc, pages: int):
    _compliant_base(doc, "Annual Benefits Enrollment")
    for i in range(pages):
        doc.add_heading(f"Section {i + 1}", level=2)          # H1 -> H2, no skip (2.4.6)
        for _ in range(6 if pages > 1 else 3):
            doc.add_paragraph(LOREM)                          # default colour (1.4.3)
        t = doc.add_table(rows=3, cols=2)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text = "Plan", "Monthly cost"     # a real header row (1.3.1)
        _mark_header_row(t)                                   # <w:tblHeader/> (1.3.1)
        for r, (a, b) in enumerate([("Basic", "$0"), ("Plus", "$40")], start=1):
            t.rows[r].cells[0].text, t.rows[r].cells[1].text = a, b
        p = doc.add_paragraph("See the ")
        _add_hyperlink(p, "https://example.com/benefits-guide", "2026 Benefits Guide")
        p.add_run(" for full details.")                        # descriptive link text (2.4.4)
        if i < pages - 1:
            doc.add_page_break()
    return []                                                  # NO violations, by construction


def build_deterministic(doc):
    """The five rule-code criteria, and nothing a model could touch."""
    doc.add_heading("Enrollment", level=1)                     # no title/lang set -> 2.4.2, 3.1.1
    doc.add_heading("Dental", level=3)                         # H1 -> H3 -> 2.4.6
    r = doc.add_paragraph().add_run("Eligibility")
    r.bold = True
    r.font.size = Pt(16)                                       # pseudo-heading -> 1.3.1
    doc.add_paragraph("All benefits-eligible employees may enrol during the window.")
    t = doc.add_table(rows=3, cols=2)                          # no header row -> 1.3.1
    for i, (a, b) in enumerate([("Plan", "Cost"), ("Basic", "$0"), ("Plus", "$40")]):
        t.rows[i].cells[0].text, t.rows[i].cells[1].text = a, b
    low = doc.add_paragraph().add_run("Deadline: March 31, 2026.")
    low.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)            # -> 1.4.3
    return [
        {"criterion": "2.4.2", "location": "docProps.core.title",
         "expected_assessment": "detected", "expected_remediation": "fixed"},
        {"criterion": "3.1.1", "location": "docProps.core.language",
         "expected_assessment": "detected", "expected_remediation": "fixed"},
        {"criterion": "2.4.6", "location": "body.heading2",
         "expected_assessment": "detected", "expected_remediation": "fixed"},
        {"criterion": "1.3.1", "location": "body.p3.pseudo-heading",
         "expected_assessment": "detected", "expected_remediation": "fixed"},
        {"criterion": "1.3.1", "location": "body.table1.header-row",
         "expected_assessment": "detected", "expected_remediation": "fixed"},
        {"criterion": "1.4.3", "location": "body.p5.run1",
         "expected_assessment": "detected", "expected_remediation": "fixed"},
    ]


def build_assisted(doc):
    """Criteria ACP ASSESSES on .docx but does not remediate.

    Three violations of 2.4.4 rather than one, deliberately: a criterion-level scorecard would
    call a single detection a full pass, and located ground truth is what makes that visible.
    """
    _compliant_base(doc, "Accessibility Policy Summary")
    p = doc.add_paragraph("For the full policy, ")
    _add_hyperlink(p, "https://example.com/policy", "click here")
    p.add_run(".")
    p2 = doc.add_paragraph("For the appeals process, ")
    _add_hyperlink(p2, "https://example.com/appeals", "click here")
    p2.add_run(".")
    p3 = doc.add_paragraph("More: ")
    _add_hyperlink(p3, "https://example.com/benefits/2026/enrollment/details",
                   "https://example.com/benefits/2026/enrollment/details")
    doc.add_paragraph("Click the round green button on the right to continue.")
    doc.add_paragraph("Veuillez consulter la politique d’accessibilité avant le 31 mars.")
    return [
        {"id": "v001", "criterion": "2.4.4", "location": "body.p2.hyperlink1",
         "expected_assessment": "detected", "expected_remediation": "none"},
        {"id": "v002", "criterion": "2.4.4", "location": "body.p3.hyperlink1",
         "expected_assessment": "detected", "expected_remediation": "none"},
        {"id": "v003", "criterion": "2.4.4", "location": "body.p4.hyperlink1",
         "expected_assessment": "detected", "expected_remediation": "none"},
        {"id": "v004", "criterion": "1.3.3", "location": "body.p5",
         "expected_assessment": "detected", "expected_remediation": "none"},
        {"id": "v005", "criterion": "3.1.2", "location": "body.p6",
         "expected_assessment": "detected", "expected_remediation": "none"},
    ]


def _integrity(path: Path) -> dict:
    """Structural counts, captured now — the only moment they are known to be correct."""
    d = Document(str(path))
    return {"paragraphs": len(d.paragraphs), "tables": len(d.tables),
            "images": len(d.inline_shapes), "sections": len(d.sections),
            "characters": sum(len(p.text) for p in d.paragraphs)}


CONTROLS = [
    ("ctl-clean-1p", lambda d: build_clean(d, 1), 1),
    ("ctl-clean-25p", lambda d: build_clean(d, 25), 25),
    ("ctl-deterministic-only", build_deterministic, 1),
    ("ctl-assisted-only", build_assisted, 1),
]


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--out", type=Path, default=Path.home() / "Downloads" / "acp-docx-eval")
    args = ap.parse_args()
    docs, expected = args.out / "docs", args.out / "expected"
    docs.mkdir(parents=True, exist_ok=True)
    expected.mkdir(parents=True, exist_ok=True)

    for name, fn, pages in CONTROLS:
        doc = Document()
        violations = fn(doc)
        path = docs / f"{name}.docx"
        doc.save(str(path))
        if name == "ctl-deterministic-only":
            _strip_default_language(path)     # python-docx writes a default lang otherwise
        payload = {"file": path.name, "pages": pages, "violations": violations,
                   "integrity": _integrity(path)}
        (expected / f"{name}.json").write_text(json.dumps(payload, indent=2) + "\n")
        crits = sorted({v["criterion"] for v in violations})
        print(f"  {path.name:26} {len(violations)} violation(s)  {crits or '(clean)'}")

    print(f"\n4 controls in {docs}\nground truth in {expected}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
