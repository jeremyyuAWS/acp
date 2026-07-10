#!/usr/bin/env python3
"""Generate the Word + Excel "accessibility demo" fixtures.

Each fixture is a realistic-looking business document deliberately seeded with one
instance of every in-scope WCAG failure ACP can ACTION on that format — so a live
scan → remediate → re-scan visibly drives the file from many findings to a clean
certificate, exercising all three remediation lanes:

  * AUTO      — cleared deterministically on re-scan, no human touch
                (title, language, heading skip/outline, table headers, justified
                 text, low-contrast text, ambiguous links, hidden content).
  * PROPOSED  — a prefilled one-click HITL card the reviewer approves
                (image alt text 1.1.1, images-of-text 1.4.5, language-of-parts
                 3.1.2, sensory rewrite 1.3.3).
  * (Reading level 3.1.5 is surfaced as guidance.)

Self-contained: builds valid OOXML from scratch with python-docx / openpyxl (so
the vendored .NET OpenXML analysers parse them), then injects the few raw-XML
conditions those libraries can't express (content-control form field, hyperlinks).

Usage:
    python scripts/gen_demo_fixtures.py [OUT_DIR]     # default: demo-fixtures/
"""
from __future__ import annotations

import io
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font, PatternFill
from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, TwoCellAnchor
from openpyxl.worksheet.table import Table
from PIL import Image, ImageDraw

# A long, deliberately convoluted sentence (many clauses / long words) so the
# Flesch–Kincaid reading-level check (3.1.5) fires on both formats.
COMPLEX = (
    "Notwithstanding the aforementioned contractual stipulations, the counterparties "
    "hereby irrevocably acknowledge that the indemnification obligations enumerated "
    "throughout the preceding paragraphs shall remain operative in perpetuity, "
    "insofar as such obligations are not subsequently extinguished by a superseding "
    "instrument executed with commensurate formality and corresponding authority."
)
# French passage inside an English document → language-of-parts (3.1.2).
FRENCH = ("Le rapport trimestriel montre une croissance significative des revenus "
          "dans toutes les régions européennes cette année.")
# A sensory-only instruction → 1.3.3.
SENSORY = ("To approve the request, click the large round green button located in the "
           "bottom-right corner of the screen.")
# ~230 words of dense, high-grade prose so the Flesch–Kincaid reading-level check
# (3.1.5 — needs >= 200 words at grade >= 15) fires where a single sentence can't.
DENSE_BLOCK = COMPLEX * 4


def _png(color=(27, 58, 107), size=(240, 120)) -> bytes:
    """A plain solid image (no text) — an alt-text gap (1.1.1) with no OCR content."""
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    return buf.getvalue()


def _png_of_text() -> bytes:
    """An image that BAKES IN a full sentence of real text → images-of-text (1.4.5)
    plus an alt gap (1.1.1). Rendered large enough for tesseract to read >= 10 words."""
    im = Image.new("RGB", (900, 240), "white")
    d = ImageDraw.Draw(im)
    lines = ["Quarterly Revenue Report 2026",
             "Total revenue increased fourteen percent",
             "across every regional business unit"]
    y = 20
    for ln in lines:
        d.text((20, y), ln, fill="black")
        y += 70
    # scale up so the default bitmap font is large enough to OCR reliably
    im = im.resize((1800, 480), Image.LANCZOS)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()


# ── raw-XML helpers for the pieces python-docx can't express ──────────────────

def _add_hyperlink(paragraph, url: str, text: str):
    """A real external hyperlink (relationship + <w:hyperlink>). Two of these sharing
    the visible text 'click here' but pointing at different URLs → ambiguous link
    purpose (2.4.9)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run = paragraph._p.makeelement(qn("w:r"), {})
    rpr = paragraph._p.makeelement(qn("w:rPr"), {})
    t = paragraph._p.makeelement(qn("w:t"), {})
    t.text = text
    run.append(rpr)
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_unlabeled_checkbox(paragraph):
    """A checkbox content control with no <w:alias> → an unlabeled form field (3.3.2)."""
    sdt = paragraph._p.makeelement(qn("w:sdt"), {})
    sdtpr = paragraph._p.makeelement(qn("w:sdtPr"), {})
    # <w14:checkbox> is the modern checkbox; the detector matches <w:checkbox> too —
    # use the plain w: form the detector keys on.
    checkbox = paragraph._p.makeelement(qn("w:checkbox"), {})
    sdtpr.append(checkbox)
    sdtcontent = paragraph._p.makeelement(qn("w:sdtContent"), {})
    run = paragraph._p.makeelement(qn("w:r"), {})
    t = paragraph._p.makeelement(qn("w:t"), {})
    t.text = "☐"
    run.append(t)
    sdtcontent.append(run)
    sdt.append(sdtpr)
    sdt.append(sdtcontent)
    paragraph._p.append(sdt)


def _strip_default_language(path: Path) -> None:
    """python-docx writes a default <w:lang> into word/styles.xml, which the (correct)
    .NET DocumentLanguageRule accepts as the document's declared language — so 3.1.1
    would pass. Remove every <w:lang .../> so the document declares NO language and the
    rule fires, giving the demo a language finding to auto-fix."""
    with zipfile.ZipFile(path) as z:
        parts = {n: z.read(n) for n in z.namelist()}
    for n in ("word/styles.xml", "word/settings.xml", "word/document.xml"):
        if n in parts:
            parts[n] = re.sub(rb"<w:lang\b[^>]*/>", b"", parts[n])
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, data in parts.items():
            z.writestr(n, data)


def build_docx(path: Path) -> None:
    doc = Document()
    # 2.4.2 / 3.1.1 — leave title blank and declare no language.
    doc.core_properties.title = ""
    doc.core_properties.language = ""

    # 2.4.6 — heading outline that skips a level (H1 → H3, no H2).
    doc.add_heading("AI Platform Evaluation — Executive Summary", level=1)
    doc.add_heading("Scoring Methodology", level=3)

    # 1.3.1 — a line that reads as a heading (bold, 16pt) but is left Normal-styled.
    p = doc.add_paragraph()
    r = p.add_run("Regional Performance Breakdown")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "The following sections summarize the evaluation across every business unit, "
        "highlighting where the platform met, exceeded, or fell short of expectations."
    )

    # 1.4.3 — body text set in a light grey that fails contrast on white.
    p = doc.add_paragraph()
    r = p.add_run("Confidential — internal distribution only.")
    r.font.color.rgb = RGBColor(0xC8, 0xC8, 0xC8)

    # 1.3.1 (tables) — a data table with no header row marked.
    table = doc.add_table(rows=4, cols=3)
    data = [["North", "1,240", "up"], ["South", "980", "flat"],
            ["East", "1,510", "up"], ["West", "760", "down"]]
    for row, vals in zip(table.rows, data):
        for cell, v in zip(row.cells, vals):
            cell.text = v

    # 2.4.9 — two links with identical visible text but different destinations.
    p = doc.add_paragraph("For the full methodology ")
    _add_hyperlink(p, "https://example.com/methodology", "click here")
    p.add_run(", and for the raw scores ")
    _add_hyperlink(p, "https://example.com/scores", "click here")
    p.add_run(".")

    # 1.4.8 — three justified paragraphs (both-margin alignment).
    for _ in range(3):
        para = doc.add_paragraph(
            "Each criterion was weighted according to its impact on end users, then "
            "normalized so that no single dimension could dominate the composite score "
            "reported to the steering committee for their quarterly review."
        )
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3.1.5 — a deliberately dense, hard-to-read sentence.
    doc.add_paragraph(COMPLEX)

    # 3.1.2 — a French passage inside the English document.
    doc.add_paragraph(FRENCH)

    # 1.3.3 — a sensory-only instruction.
    doc.add_paragraph(SENSORY)

    # 3.3.2 — an unlabeled checkbox form field.
    p = doc.add_paragraph("I accept the terms: ")
    _add_unlabeled_checkbox(p)

    # 1.1.1 — an image with no alt text.
    doc.add_picture(io.BytesIO(_png()))

    # 1.4.5 (+ 1.1.1) — an image that bakes in real text.
    doc.add_picture(io.BytesIO(_png_of_text()))

    doc.save(str(path))
    _strip_default_language(path)     # so 3.1.1 (language of page) fires


def build_xlsx(path: Path) -> None:
    wb = Workbook()
    # 2.4.2 / 3.1.1 — no title, no language.
    wb.properties.title = None
    wb.properties.language = None

    ws = wb.active
    ws.title = "Q4 Results"

    # 1.3.1 — a defined Table with NO header row (headerRowCount=0), which the
    # analyser's TableHeaderRule flags. A plain range of cells is not a table, so it
    # must be a real ListObject to be seen as a headerless table.
    rows = [["North", 1240, "up"], ["South", 980, "flat"],
            ["East", 1510, "up"], ["West", 760, "down"]]
    for row in rows:
        ws.append(row)
    tbl = Table(displayName="RegionData", ref="A1:C4")
    tbl.headerRowCount = 0
    ws.add_table(tbl)

    # 1.4.3 / 1.4.6 — a cell whose light-grey text fails contrast on the white fill.
    ws["A7"] = "Confidential — internal only"
    ws["A7"].font = Font(color="C8C8C8")
    ws["A7"].fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    # 3.1.5 — a dense block (>= 200 words, high grade); 3.1.2 — a French passage;
    # 1.3.3 — a sensory instruction.
    ws["A9"] = DENSE_BLOCK
    ws["A10"] = FRENCH
    ws["A11"] = SENSORY

    # 1.3.2 — hidden rows/columns that still hold data (HiddenContentRule). The rule
    # reads a cell's value via CellValue.Text, which is populated for numbers but not
    # for inline strings, so the hidden data is numeric (a hidden reconciliation figure).
    ws["A13"] = 4821955
    ws.row_dimensions[13].hidden = True
    ws["E1"] = 990211
    ws.column_dimensions["E"].hidden = True

    # 1.1.1 — images with no alt. The xlsx AltTextRule inspects TwoCellAnchor drawings,
    # so anchor both images across a cell range (openpyxl's default single-cell anchor
    # is not examined). One plain image, one that bakes in text (also 1.4.5).
    def _two_cell(col, row):
        return TwoCellAnchor(editAs="oneCell",
                             _from=AnchorMarker(col=col, row=row),
                             to=AnchorMarker(col=col + 4, row=row + 6))

    img1 = XLImage(io.BytesIO(_png()))
    img1.anchor = _two_cell(7, 1)          # H2:L8
    ws.add_image(img1)
    img2 = XLImage(io.BytesIO(_png_of_text()))
    img2.anchor = _two_cell(7, 11)         # H12:L18
    ws.add_image(img2)

    wb.save(str(path))


def main() -> None:
    out_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("demo-fixtures")
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "word-accessibility-demo.docx"
    xlsx_path = out_dir / "excel-accessibility-demo.xlsx"
    build_docx(docx_path)
    build_xlsx(xlsx_path)
    print(f"wrote {docx_path}")
    print(f"wrote {xlsx_path}")


if __name__ == "__main__":
    main()
