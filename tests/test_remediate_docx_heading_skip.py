"""Round-trip for the docx 2.4.6 heading-skip remediator (office_structure detector), and
for the levels the pseudo-heading promoter assigns in the first place.

DOCX_HEADING_SKIP (2.4.6) is declared fix_mode='auto' and detected by
office_structure.docx_checks, but had no remediator — so it never cleared and was
routed to the HITL human lane. This builds a doc with a skipped heading level
(Heading 1 → Heading 3), remediates it, and re-scans with the SAME python detector,
asserting 0 residual. Needs only python-docx (no .NET CLI), since this SC is a
python-side check.

The promoter's levels matter beyond the skip detector, which only ever proves the outline
has no GAPS: it says nothing about whether the levels describe the document. Ranking
pseudo-headings by absolute font size across the whole document let one oversized
paragraph take Heading 1 and scattered the real sections underneath it, and both docx
detectors still passed — so it shipped unattended (this path writes w:pStyle without
review). Those levels now come from the font hierarchy in document order
(proposals.heading_level_sequence).
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

import pytest

pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

ACP = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ACP / "api"))

import office_structure as osx      # noqa: E402
import remediate_office as R        # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
BODY = "Ordinary body prose with enough words in it to read as a real paragraph."


def _heading_skip_findings(path: Path):
    return [f for f in osx.docx_checks(path) if f["ruleId"] == "DOCX_HEADING_SKIP"]


def _outline(path) -> list[tuple[str, int]]:
    """[(text, heading level)] for every heading-styled paragraph, in document order. Reads
    the style id off the XML rather than via `p.style`, which resolves against the styles
    part; the promoter writes "Heading2" where python-docx wrote "Heading 2", and both
    spellings are levels as far as office_structure and Word are concerned."""
    out = []
    for para in Document(str(path)).paragraphs:
        pPr = para._p.find(f"{W}pPr")
        style = pPr.find(f"{W}pStyle") if pPr is not None else None
        val = style.get(f"{W}val") if style is not None else None
        level = re.fullmatch(r"Heading\s?([1-6])", val or "")
        if level:
            out.append((para.text, int(level.group(1))))
    return out


def _pseudo(doc, text: str, pt: int) -> None:
    """A body-styled paragraph that LOOKS like a heading: large and bold, no Heading style.
    14pt (office_structure.PSEUDO_HEADING_MIN_HALF_PT) is the floor for the detector."""
    run = doc.add_paragraph().add_run(text)
    run.bold = True
    run.font.size = Pt(pt)


def test_docx_heading_skip_clears_2_4_6(tmp_path):
    d = Document()
    d.add_heading("Overview", level=1)
    d.add_paragraph("Some intro text with enough words to be a real paragraph.")
    d.add_heading("Deep detail", level=3)          # skips Heading 2 → 2.4.6 failure
    d.add_paragraph("More body text under the deep detail heading here.")
    src = tmp_path / "skip.docx"
    d.save(src)

    assert _heading_skip_findings(src), "fixture should trip DOCX_HEADING_SKIP"

    fixed, applied, _skipped = R.remediate_office(src)
    assert fixed is not None
    assert any("2.4.6" in a for a in applied), f"expected a 2.4.6 change: {applied}"
    assert _heading_skip_findings(Path(fixed)) == [], "2.4.6 should clear on re-scan"


def test_promoted_levels_nest_in_document_order(tmp_path):
    """A report whose pull figure is set larger than its own title. Absolute size rank made
    "$4.2B" the document's ONLY Heading 1, demoted the real title to Heading 2, and left the
    four peer sections split across Heading 2/3 with the subsection at Heading 4 — an outline
    with one H1 and no gaps, so nothing downstream objected."""
    doc = Document()
    for text, pt in [("Annual Accessibility Report", 30), ("Executive Summary", 18),
                     ("$4.2B", 44), ("Scope", 18), ("Findings", 18),
                     ("Findings by region", 15), ("Appendix", 18)]:
        _pseudo(doc, text, pt)
        doc.add_paragraph(BODY)
    src = tmp_path / "report.docx"
    doc.save(src)

    fixed, applied, _skipped = R.remediate_office(src)
    assert any("pseudo-heading" in a for a in applied)
    levels = dict(_outline(fixed))
    assert levels["Annual Accessibility Report"] == 1        # the title keeps the top level
    sections = ["Executive Summary", "Scope", "Findings", "Appendix"]
    assert len({levels[s] for s in sections}) == 1           # peers share one level…
    assert levels["Findings by region"] == levels["Findings"] + 1        # …with one below
    ordered = [lvl for _t, lvl in _outline(fixed)]
    assert all(b - a <= 1 for a, b in zip(ordered, ordered[1:]))


def test_a_pull_figure_is_not_promoted_at_all(tmp_path):
    """The shared furniture predicate (office_structure.looks_like_heading_furniture) keeps
    "$4.2B" out of the heading structure entirely rather than merely off Heading 1. It has to
    move the DETECTOR too, or the promoter skipping the figure would leave
    DOCX_PSEUDO_HEADING firing and the fix would stop clearing the re-scan — so this asserts
    both ends: nothing about the figure is promoted, and the document still re-scans clean."""
    doc = Document()
    _pseudo(doc, "Annual Report", 30)
    doc.add_paragraph(BODY)
    _pseudo(doc, "$4.2B", 44)                               # a pull figure, not a section
    doc.add_paragraph(BODY)
    _pseudo(doc, "“We grew faster than the market.”", 20)   # a pull quote, not a section
    doc.add_paragraph(BODY)
    _pseudo(doc, "Outlook", 18)
    doc.add_paragraph(BODY)
    src = tmp_path / "figures.docx"
    doc.save(src)

    fixed, _applied, _skipped = R.remediate_office(src)
    promoted = [text for text, _lvl in _outline(fixed)]
    assert promoted == ["Annual Report", "Outlook"]
    assert [f["ruleId"] for f in osx.docx_checks(Path(fixed))
            if f["ruleId"] == "DOCX_PSEUDO_HEADING"] == []


def test_promotion_leaves_an_existing_styled_heading_1_as_the_title(tmp_path):
    """The promoter's levels feed the 1.3.1 "exactly one Heading 1" pass, which demotes the
    LATER competing H1. Ordering the promoted levels means the first pseudo-heading claims
    Heading 1, so a document that already styled its own title correctly must still keep it."""
    doc = Document()
    doc.add_heading("Benefits Policy", level=1)             # a real, correctly styled title
    doc.add_paragraph(BODY)
    for text in ["Eligibility", "How to enrol", "Appeals"]:
        _pseudo(doc, text, 18)
        doc.add_paragraph(BODY)
    src = tmp_path / "policy.docx"
    doc.save(src)

    fixed, applied, _skipped = R.remediate_office(src)
    assert any("pseudo-heading" in a for a in applied)
    outline = _outline(fixed)
    assert outline[0] == ("Benefits Policy", 1)
    assert [lvl for _t, lvl in outline].count(1) == 1        # only one Heading 1 survives
    assert {lvl for t, lvl in outline if t != "Benefits Policy"} == {2}
    assert _heading_skip_findings(Path(fixed)) == []


def _sized_body(doc, pt: int) -> None:
    """A real body paragraph at an explicit size — too long to be heading-like, so it is never a
    candidate, but it is what the baseline is measured from."""
    doc.add_paragraph().add_run(BODY).font.size = Pt(pt)


def _large_print_policy(tmp_path, body_pt: int) -> Path:
    """The same document twice, at two body sizes. Everything else is identical: one clearly
    larger title, three short bold lines set at 14pt, body between them.

    At 11pt body those three lines are 3pt larger than their surroundings and are headings. At
    14pt body they are exactly the size of the text around them and could be anything — a
    section name, or a bold lead-in. The bytes of those three paragraphs do not change; only the
    document they sit in does.
    """
    doc = Document()
    _pseudo(doc, "Benefits Policy", 20)               # clearly larger at either body size
    _sized_body(doc, body_pt)
    for text in ["Eligibility", "How to enrol", "Appeals"]:
        _pseudo(doc, text, 14)                        # exactly PSEUDO_HEADING_MIN_HALF_PT
        _sized_body(doc, body_pt)
        _sized_body(doc, body_pt)
    src = tmp_path / f"policy-{body_pt}pt.docx"
    doc.save(src)
    return src


def test_an_ordinary_document_still_promotes_all_of_its_headings(tmp_path):
    """The control, and the half that makes the other test mean anything. Without it, a gate
    that simply refused to promote 14pt lines would pass the large-print test and be a
    regression."""
    fixed, _applied, _skipped = R.remediate_office(_large_print_policy(tmp_path, 11))
    promoted = [text for text, _lvl in _outline(fixed)]
    assert promoted == ["Benefits Policy", "Eligibility", "How to enrol", "Appeals"]


def test_a_large_print_document_is_not_restyled_out_of_its_own_body_text(tmp_path):
    """14pt body. The three 14pt lines clear the ABSOLUTE floor and are indistinguishable from
    the text around them, so promoting them writes a heading outline out of body copy —
    unattended, on precisely the documents most likely to have been remediated for low vision
    already."""
    src = _large_print_policy(tmp_path, 14)
    fixed, _applied, skipped = R.remediate_office(src)

    promoted = [text for text, _lvl in _outline(fixed)]
    assert promoted == ["Benefits Policy"], (
        "a 14pt line in a 14pt document was restyled as a heading")

    # Deferred, not dropped. The finding stands — silence here would read as "nothing to fix".
    assert any("not larger than this document's body text" in s for s in skipped), skipped
    assert any("3 paragraph(s)" in s for s in skipped), skipped


def test_the_deferred_paragraphs_keep_their_finding(tmp_path):
    """The honest consequence, asserted rather than discovered: declining to guess means
    DOCX_PSEUDO_HEADING still fires on re-scan. A fix that cleared the detector by leaving the
    document alone would be the worst of both — no change, and no finding either."""
    fixed, _applied, _skipped = R.remediate_office(_large_print_policy(tmp_path, 14))
    assert [f["ruleId"] for f in osx.docx_checks(Path(fixed))
            if f["ruleId"] == "DOCX_PSEUDO_HEADING"], "the deferred finding disappeared"


def test_docx_valid_outline_untouched(tmp_path):
    d = Document()
    d.add_heading("Overview", level=1)
    d.add_heading("Section", level=2)
    d.add_heading("Subsection", level=3)
    src = tmp_path / "ok.docx"
    d.save(src)

    assert not _heading_skip_findings(src)
    fixed, applied, _skipped = R.remediate_office(src)
    # A gap-free outline must not draw a spurious 2.4.6 change …
    assert not any("2.4.6" in a for a in applied)
    # … and must still re-scan clean.
    assert _heading_skip_findings(Path(fixed or src)) == []
