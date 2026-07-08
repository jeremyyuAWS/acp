"""Round-trip regression for deterministic Office structural remediation (ADR 0012 program).

For each format we build a deliberately-broken file, scan it with the owned analyser CLI,
run it through remediate_office, re-scan the output, and assert the deterministic findings
are gone (0 residual). Self-skips when the .NET CLI / python libs are unavailable.
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path

import pytest

pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.shared import RGBColor  # noqa: E402

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "api"))
import remediate_office as R  # noqa: E402

_ACP = Path(__file__).resolve().parents[1]
_DOTNET = Path(os.environ.get("ACP_DOTNET") or os.path.expanduser("~/.dotnet/dotnet"))
_CLI = Path(os.environ.get("ACP_OFFICE_CLI")
            or _ACP / "spike/dotnet/AcpScan.Cli/bin/Release/net10.0/AcpScan.Cli.dll")

pytestmark = pytest.mark.skipif(
    not _DOTNET.exists() or not _CLI.exists(),
    reason="office analyser CLI or dotnet runtime not built/available",
)


def _scan_one(path: Path, out: Path) -> list[dict]:
    """Scan a single file (in an isolated dir) and return its issues."""
    d = out.parent / (out.stem + "_dir")
    d.mkdir(parents=True, exist_ok=True)
    tgt = d / path.name
    tgt.write_bytes(path.read_bytes())
    env = {**os.environ, "DOTNET_ROOT": os.path.expanduser("~/.dotnet")}
    subprocess.run([str(_DOTNET), str(_CLI), str(d), str(out)],
                   capture_output=True, text=True, env=env, check=True)
    results = json.loads(out.read_text())
    return results[0].get("issues", []) if results else []


def _round_trip(tmp_path: Path, src: Path, wcags: set[str]) -> tuple[list[dict], list[dict]]:
    pre = _scan_one(src, tmp_path / "pre.json")
    fixed, _applied, _skipped = R.remediate_office(src)
    assert fixed is not None, "remediation produced no output"
    post = _scan_one(Path(fixed), tmp_path / "post.json")
    pre_hit = [i for i in pre if i["wcag"] in wcags]
    post_hit = [i for i in post if i["wcag"] in wcags]
    return pre_hit, post_hit


def test_docx_tables_and_headings_clear_1_3_1(tmp_path):
    d = Document()
    d.add_heading("First section", level=2)          # no Heading 1 anywhere
    d.add_paragraph("Body text with several words of content under the first section.")
    t = d.add_table(rows=3, cols=3)                  # multi-row table, no header row
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.text = f"cell {i}"
    d.add_heading("Second section", level=2)
    src = tmp_path / "broken.docx"
    d.save(src)

    pre, post = _round_trip(tmp_path, src, {"SC_1_3_1"})
    assert len(pre) >= 2, f"fixture should trip table-header + no-H1: {pre}"
    assert post == [], f"1.3.1 findings should clear: {post}"


def test_docx_contrast_clears_1_4_3(tmp_path):
    d = Document()
    run = d.add_paragraph().add_run(
        "This grey text has insufficient contrast against the white background here.")
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)   # ~2.8:1 on white
    src = tmp_path / "contrast.docx"
    d.save(src)

    pre, post = _round_trip(tmp_path, src, {"SC_1_4_3"})
    assert len(pre) >= 1, f"fixture should trip contrast: {pre}"
    assert post == [], f"1.4.3 should clear after recolour: {post}"


def test_xlsx_table_header_and_hidden_content_clear(tmp_path):
    openpyxl = pytest.importorskip("openpyxl")
    from openpyxl.worksheet.table import Table

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Name", "Amount"])
    for r in (["Alice", 10], ["Bob", 20], ["Carol", 30]):
        ws.append(r)
    tab = Table(displayName="T1", ref="A1:B4")
    tab.headerRowCount = 0                             # 1.3.1 violation
    ws.add_table(tab)
    ws.append(["HiddenSecret", 99])
    ws.row_dimensions[5].hidden = True                # 1.3.2 hidden row with data
    src = tmp_path / "sheet.xlsx"
    wb.save(src)

    pre, post = _round_trip(tmp_path, src, {"SC_1_3_1", "SC_1_3_2"})
    assert len(pre) >= 2, f"fixture should trip table-header + hidden row: {pre}"
    assert post == [], f"1.3.1/1.3.2 should clear: {post}"
