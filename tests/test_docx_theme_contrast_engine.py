"""Regression test for the owned Office analyser's DOCX contrast rule reading
w:color/@themeColor (+ themeTint/themeShade), not only a direct w:color/@val hex.

Before this fix, `Docx/Rules/ColourContrastRule.cs` read only `w:color/@val` and
`continue`d past any run with no direct hex — a run coloured via theme (the common
case in professionally-authored documents) was never evaluated for contrast at all,
silently. See engine/office-analysers/DigitalA11y.Analysers.DotNet/Docx/Helpers/
ThemeColorHelper.cs for the resolution (theme scheme lookup + ECMA-376 HSL tint/shade).

The test shells out to the built .NET CLI, so it self-skips when the toolchain
(python-docx, dotnet, or the built AcpScan.Cli.dll) is not present — same posture as
tests/test_office_language_rules.py.
"""
from __future__ import annotations

import json
import os
import subprocess
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

_ACP = Path(__file__).resolve().parents[1]
_DOTNET = Path(os.environ.get("ACP_DOTNET") or os.path.expanduser("~/.dotnet/dotnet"))
_CLI = Path(os.environ.get("ACP_OFFICE_CLI")
            or _ACP / "spike/dotnet/AcpScan.Cli/bin/Release/net10.0/AcpScan.Cli.dll")

pytestmark = pytest.mark.skipif(
    not _DOTNET.exists() or not _CLI.exists(),
    reason="office analyser CLI or dotnet runtime not built/available",
)


def _run_with_theme_color(paragraph, text, *, theme_color, tint=None, shade=None):
    run = paragraph.add_run(text)
    rpr = run._element.get_or_add_rPr()
    color = OxmlElement("w:color")
    color.set(qn("w:themeColor"), theme_color)
    if tint is not None:
        color.set(qn("w:themeTint"), tint)
    if shade is not None:
        color.set(qn("w:themeShade"), shade)
    rpr.append(color)
    return run


def _build_fixtures(d: Path) -> None:
    # accent3 (A5A5A5, mid grey) at full strength on the default white background —
    # true WCAG ratio ~2.5:1, well below the 4.5:1 AA threshold. Must FLAG.
    gray = Document()
    _run_with_theme_color(gray.add_paragraph(), "Themed grey body text.", theme_color="accent3")
    gray.save(d / "theme_gray_fail.docx")

    # accent1 (4472C4, blue) tinted heavily toward white (themeTint=D0 ~ 0.82) — pushes
    # an already-borderline blue into clearly-failing territory. Must FLAG.
    tinted = Document()
    _run_with_theme_color(tinted.add_paragraph(), "Themed tinted body text.",
                           theme_color="accent1", tint="D0")
    tinted.save(d / "theme_tinted_fail.docx")

    # text1 (maps to dk1 = black) at full strength on white — 21:1, unambiguously
    # passing. Proves theme resolution isn't just always-flagging themed runs.
    black = Document()
    _run_with_theme_color(black.add_paragraph(), "Themed black body text.", theme_color="text1")
    black.save(d / "theme_black_pass.docx")


def _run_cli(indir: Path, tmp_path: Path) -> dict[str, list[dict]]:
    out = tmp_path / "out.json"
    env = {**os.environ, "DOTNET_ROOT": os.path.expanduser("~/.dotnet"),
           "DOTNET_CLI_TELEMETRY_OPTOUT": "1", "DOTNET_NOLOGO": "1"}
    subprocess.run([str(_DOTNET), str(_CLI), str(indir), str(out)],
                   capture_output=True, text=True, env=env, check=True)
    return {item["file"]: item.get("issues", []) for item in json.loads(out.read_text())}


def _has_contrast_finding(issues: list[dict]) -> bool:
    return any(i.get("wcag") == "SC_1_4_3" for i in issues)


def test_theme_colored_text_is_evaluated_for_contrast(tmp_path):
    _build_fixtures(tmp_path)
    res = _run_cli(tmp_path, tmp_path)

    assert _has_contrast_finding(res["theme_gray_fail.docx"])
    assert _has_contrast_finding(res["theme_tinted_fail.docx"])
    assert not _has_contrast_finding(res["theme_black_pass.docx"])
