"""Metamorphic properties for .docx assessment — confidence without ground truth.

WHY THIS EXISTS ALONGSIDE THE LABELLED CORPUS. `test_docx_corpus_regression_gate` asserts that
32 documents whose correct answer is KNOWN get that answer. Its reach is bounded by how many
answers somebody wrote down: 21 seeded violations, which by the rule of three bounds a
false-rate claim at ~14%. Labelling more is the honest fix and it is slow.

Metamorphic testing needs no labels at all. Instead of "what is the right answer for this
document", it asserts RELATIONSHIPS between two scans:

    changing something irrelevant must not change the verdict
    adding a clean page must not disturb the criteria already there
    fixing criterion X must not move criterion Y
    remediating twice must be the same as remediating once

Every one of those can be checked on a document nobody has labelled — including, eventually, a
customer's own estate. And they attack the failure mode that actually bit repeatedly today:
not "the detector is wrong about this document" but "the detector is right in isolation and
wrong in company". A criterion whose verdict moves because an unrelated paragraph was added is
a bug no single-document fixture can express, because it is a property of the PAIR.

WHAT THEY CANNOT DO. A detector that is uniformly wrong satisfies every relation here — if
1.4.3 never fires at all, appending a page will certainly not change its verdict. These are a
complement to the labelled corpus, never a replacement, and the corpus gate is what stops the
uniform-silence case. Both are in CI for that reason.
"""
import io
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "api"))
sys.path.insert(0, str(ROOT / "scripts"))

from engines import OFFICE_OK  # noqa: E402

pytestmark = pytest.mark.skipif(
    not OFFICE_OK,
    reason="Office .NET CLI not built — every relation below would hold trivially over a silent "
           "engine, which is the one way these tests can pass while meaning nothing")

import assessment_policy as pol  # noqa: E402
import corpus_expectations as ce  # noqa: E402

# PASS/FAIL are plain strings; only REVIEW and NOT_EVALUATED are module attributes on
# assessment_policy. corpus_expectations already normalises all four, so use that rather than
# reaching for FAIL — which does not exist and fails at the assertion, not at import.
FAIL, PASS = ce.FAIL, ce.PASS

CORE17 = sorted(sc for sc, f in pol.SCOPE_PRESETS["acp-core-17"].items() if "docx" in f)


def _verdicts(path: Path) -> dict[str, str]:
    """Certified per-criterion outcome, the way `store` computes it."""
    import scanner
    tmp = Path(tempfile.mkdtemp())
    shutil.copy(path, tmp / path.name)
    res, _ = scanner.analyse_and_assess(tmp, path.name)
    issues = list((res or {}).get("issues") or [])
    fail, review = pol._split_sc_counts(issues)
    return {sc: pol._rule_outcome(sc, "docx", fail.get(sc, 0), review.get(sc, 0), "AA", None)
            for sc in CORE17}


def _doc(title="Patient Benefits Notice"):
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    d = Document()
    d.core_properties.title = title
    for el in d.styles.element.findall(qn("w:docDefaults")):
        rpr = el.find(qn("w:rPrDefault"))
        if rpr is None:
            continue
        inner = rpr.find(qn("w:rPr"))
        if inner is None:              # `or` truth-tests the element, which lxml deprecates
            inner = OxmlElement("w:rPr")
            rpr.append(inner)
        node = OxmlElement("w:lang")
        node.set(qn("w:val"), "en-US")
        inner.append(node)
    return d


def _clean_page(d):
    """Content that introduces no finding on any criterion."""
    d.add_heading("Coverage", level=1)
    d.add_paragraph("Your 2026 coverage begins on the first of March.")
    t = d.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for i, row in enumerate((["Plan", "Premium"], ["Bronze", "$120"])):
        for j, v in enumerate(row):
            t.cell(i, j).paragraphs[0].add_run(v)
    trPr = t.rows[0]._tr.get_or_add_trPr()          # a REAL header row, not merely bold
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)
    return d


def _with_defect(d):
    """One seeded defect: a table whose header row is not marked (1.3.1)."""
    d.add_heading("Plan comparison", level=1)
    t = d.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    for i, row in enumerate((["Plan", "Premium", "Deductible"], ["Bronze", "$120", "$3,000"])):
        for j, v in enumerate(row):
            r = t.cell(i, j).paragraphs[0].add_run(v)
            r.bold = (i == 0)                        # bold, but no <w:tblHeader/>
    return d


def _save(d, path: Path) -> Path:
    d.save(path)
    return path


# ── relation 1: irrelevant change, identical verdict ──────────────────────────

def test_renaming_the_file_changes_nothing(tmp_path):
    """The cheapest relation, and it catches a real class: anything that keys off a filename.
    `_scoped_for_scoring` resolves a format from the name, and a detector that read more than
    the extension would show up here."""
    a = _save(_with_defect(_clean_page(_doc())), tmp_path / "benefits.docx")
    b = tmp_path / "a-completely-different-name-2026-final-v3.docx"
    shutil.copy(a, b)
    assert _verdicts(a) == _verdicts(b)


def test_a_clean_page_does_not_disturb_the_criteria_already_there(tmp_path):
    """THE relation this file exists for. A criterion whose verdict moves because unrelated
    compliant content was added is wrong in a way no single-document fixture can express."""
    base = _save(_with_defect(_doc()), tmp_path / "base.docx")
    before = _verdicts(base)

    grown = _with_defect(_doc())
    _clean_page(grown)
    _clean_page(grown)
    after = _verdicts(_save(grown, tmp_path / "grown.docx"))

    moved = {sc: (before[sc], after[sc]) for sc in CORE17 if before[sc] != after[sc]}
    assert not moved, f"adding compliant content moved verdicts: {moved}"


def test_duplicating_the_body_does_not_change_any_verdict(tmp_path):
    """Twice as many instances of the same defect is still the same set of criteria. Counts may
    double; verdicts must not move — including the clean ones, which is where a detector that
    trips on volume rather than content would betray itself."""
    single = _save(_with_defect(_clean_page(_doc())), tmp_path / "one.docx")
    doubled = _doc()
    for _ in range(2):
        _clean_page(doubled)
        _with_defect(doubled)
    assert _verdicts(single) == _verdicts(_save(doubled, tmp_path / "two.docx"))


# ── relation 2: remediation is well-behaved ───────────────────────────────────

def _remediate(path: Path) -> Path:
    import remediate_office
    work = Path(tempfile.mkdtemp()) / path.name
    shutil.copy(path, work)
    res = remediate_office.remediate_office(work, ai_enabled=False, applied_fixes=[],
                                            proposals=[], diffs=[])
    out = (res[0] if isinstance(res, tuple) else res) or work
    return Path(out)


def test_remediation_is_idempotent(tmp_path):
    """Remediating an already-remediated document must be a no-op. If the second pass still
    changes verdicts, the first one did not converge — and the re-scan verify gate would credit
    a fix that is still moving."""
    once = _remediate(_save(_with_defect(_clean_page(_doc())), tmp_path / "src.docx"))
    v1 = _verdicts(once)
    twice = _remediate(once)
    v2 = _verdicts(twice)
    moved = {sc: (v1[sc], v2[sc]) for sc in CORE17 if v1[sc] != v2[sc]}
    assert not moved, f"a second remediation pass changed verdicts: {moved}"


def test_remediation_never_moves_a_criterion_backwards(tmp_path):
    """No criterion may get WORSE. The VRR harness checks this per document; here it is checked
    per criterion, which is finer — a fix that clears 1.3.1 while breaking 2.4.6 nets out as
    'improved' on a document-level count and is a regression on the axis that matters."""
    src = _save(_with_defect(_clean_page(_doc())), tmp_path / "src.docx")
    before, after = _verdicts(src), _verdicts(_remediate(src))
    broke = [sc for sc in CORE17 if before[sc] != FAIL and after[sc] == FAIL]
    assert not broke, f"remediation introduced failures on {broke}"


def test_remediating_a_clean_document_changes_nothing(tmp_path):
    """The false-positive relation. A document with no findings must come back with the same
    verdicts — a remediator that rewrites something on every file would show up nowhere else,
    because its output still scans clean."""
    clean = _save(_clean_page(_doc()), tmp_path / "clean.docx")
    assert _verdicts(clean) == _verdicts(_remediate(clean))


# ── the guard against a uniformly-silent engine ───────────────────────────────

def test_the_relations_are_not_holding_vacuously(tmp_path):
    """Every relation above is satisfied by an engine that reports nothing at all. This asserts
    the fixtures actually produce a verdict worth preserving — the same lesson as the corpus
    gate's own anti-vacuous guard, which passed for a silent engine until it counted detections
    rather than tallies."""
    v = _verdicts(_save(_with_defect(_clean_page(_doc())), tmp_path / "x.docx"))
    assert v.get("1.3.1") == FAIL, f"the seeded 1.3.1 defect did not fire: {v.get('1.3.1')}"
    assert any(x == PASS for x in v.values()), "no criterion certified — the engine is silent"
