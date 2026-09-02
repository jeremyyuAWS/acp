"""An image OCR could not READ is not an image with no text in it.

THE DEFECT, and how it surfaced. `ocr_text()` returns "" on any failure — including a
`pytesseract` timeout — and `images_of_text()` counted the words in that "" and moved on. Zero
words is below the 1.4.5 floor, so the criterion reported nothing, and *nothing* is what a
document with no images of text reports too. A reading that never happened and a reading that
found nothing were the same value.

It is not hypothetical. On 2026-09-02, #1190's CI shard 1 failed with

    word-accessibility-demo.docx: python detectors did not fire for ['1.4.5']

on a PR that touches no detector. The same run's stage log shows `analyse.ocr` taking **38.976s**
on that docx and 0.11-0.42s on the three other fixtures — which embed the byte-identical image.
`_OCR_TIMEOUT_S` is 30. So the first reading hit the timeout, the 1.4.5 pass concluded "no text",
and the 1.4.9 pass then re-read the SAME image successfully (the failure is not cached) and
flagged it. One image, one document, two contradictory conclusions, and the AA one was silent.

TWO THINGS ARE WRONG THERE AND THIS FILE PINS BOTH.

1. A timeout is a statement about the machine, not about the image — and the AAA pass proved the
   read succeeds moments later. So a failed reading is RETRIED, once per document (once, because
   a document where reads keep failing is resource-starved and retrying each one only makes the
   scan someone is waiting on worse).
2. If it still cannot be read, say so. `OCR_IMAGE_UNREAD`, REVIEW — the same shape as
   `OCR_IMAGE_CAP_REACHED` and for the same reason: the honest claim is "this criterion was not
   fully checked on this document", not "these images fail" (nobody read them) and not silence
   (which asserts a conformance nobody measured).
"""
import io
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "api"))

import ocr  # noqa: E402

pytestmark = pytest.mark.skipif(not ocr.is_available(), reason="tesseract not installed")


def _png_of_text() -> bytes:
    """One image carrying well over the 1.4.5 floor of 10 real words."""
    from PIL import Image, ImageDraw

    im = Image.new("RGB", (900, 240), "white")
    d = ImageDraw.Draw(im)
    for i, ln in enumerate(["Quarterly Revenue Report 2026",
                            "Total revenue increased fourteen percent",
                            "across every regional business unit"]):
        d.text((20, 20 + i * 70), ln, fill="black")
    buf = io.BytesIO()
    im.resize((1800, 480), Image.LANCZOS).save(buf, format="PNG")
    return buf.getvalue()


def _docx_with_one_image_of_text(path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.core_properties.title = "Unread"
    doc.add_paragraph().add_run().add_picture(io.BytesIO(_png_of_text()), width=Inches(4))
    doc.save(path)
    return path


@pytest.fixture
def docx(tmp_path):
    return _docx_with_one_image_of_text(tmp_path / "d.docx")


@pytest.fixture(autouse=True)
def _cold_cache():
    """Every test here decides what a READING does, so none may be served an earlier one."""
    ocr._OCR_CACHE.clear()
    yield
    ocr._OCR_CACHE.clear()


def _fail_first_n(monkeypatch, n: int) -> dict:
    """Make the first `n` tesseract calls raise what a timeout raises, then behave normally."""
    import pytesseract

    real = pytesseract.image_to_string
    seen = {"calls": 0}

    def flaky(im, **kw):
        seen["calls"] += 1
        if seen["calls"] <= n:
            raise RuntimeError("Tesseract process timeout")
        return real(im, **kw)

    monkeypatch.setattr(pytesseract, "image_to_string", flaky)
    return seen


def _rules(findings):
    return [f["ruleId"] for f in findings]


# ── the premise ───────────────────────────────────────────────────────────────

def test_the_image_really_is_an_image_of_text(docx):
    """Without this the rest of the file could pass on a fixture that fires nothing."""
    assert _rules(ocr.images_of_text(docx, ".docx")) == ["OCR_IMAGE_OF_TEXT"]


# ── 1. a failed reading is retried, because the machine failed and not the image ──

def test_a_timed_out_reading_is_retried_and_the_finding_survives(docx, monkeypatch):
    """The exact #1190 shape: one transient failure, and 1.4.5 must NOT go quiet."""
    seen = _fail_first_n(monkeypatch, 1)
    out = ocr.images_of_text(docx, ".docx")
    assert seen["calls"] == 2, f"expected one retry after the failure, saw {seen['calls']} call(s)"
    assert _rules(out) == ["OCR_IMAGE_OF_TEXT"], (
        f"a transient timeout silenced the criterion: {_rules(out)}")


def test_the_retry_is_spent_once_per_document_not_once_per_image(tmp_path, monkeypatch):
    """A starved machine must not turn a 30s timeout into 30 of them. Two unreadable images
    cost two attempts and one retry between them — not one retry each."""
    import zipfile

    from docx import Document
    from docx.shared import Inches
    from PIL import Image, ImageDraw

    def png(word: str) -> bytes:
        # DISTINCT bytes per image, or the OOXML package deduplicates the two pictures into one
        # media part — the OCR pass walks parts, not references, so the second image this test
        # is entirely about would not exist. (test_ocr_cap_is_reported.py was bitten by this.)
        im = Image.new("RGB", (900, 240), "white")
        ImageDraw.Draw(im).text((20, 60), f"{word} revenue increased fourteen percent this year",
                                fill="black")
        buf = io.BytesIO()
        im.resize((1800, 480), Image.LANCZOS).save(buf, format="PNG")
        return buf.getvalue()

    doc = Document()
    for word in ("Quarterly", "Regional"):
        doc.add_paragraph().add_run().add_picture(io.BytesIO(png(word)), width=Inches(4))
    p = tmp_path / "two.docx"
    doc.save(p)
    with zipfile.ZipFile(p) as z:
        n_media = len([n for n in z.namelist() if n.startswith("word/media/")])
    assert n_media == 2, f"the fixture holds {n_media} media part(s), so it measures nothing"

    seen = _fail_first_n(monkeypatch, 99)          # nothing will ever read
    ocr.images_of_text(p, ".docx")
    assert seen["calls"] == 3, (
        f"2 images with ONE retry for the document should cost 3 attempts, not {seen['calls']}")


# ── 2. an image that still cannot be read is REPORTED, never passed over ──────

def test_an_unreadable_image_is_reported_not_silent(docx, monkeypatch):
    _fail_first_n(monkeypatch, 99)
    out = ocr.images_of_text(docx, ".docx")
    assert _rules(out) == ["OCR_IMAGE_UNREAD"], (
        f"an image nobody could read produced {_rules(out) or 'silence'} — which is what a "
        f"document with no images of text produces")


def test_the_unread_notice_is_advisory_and_says_what_was_not_looked_at(docx, monkeypatch):
    _fail_first_n(monkeypatch, 99)
    (f,) = ocr.images_of_text(docx, ".docx")
    assert f["severity"] == "REVIEW", (
        "blocking would assert a defect nobody has evidence for; the evidence is that nobody "
        "looked")
    assert f["wcag"].startswith("1.4.5"), f["wcag"]
    detail = f["detail"]
    assert "1 " in detail and "could not be read" in detail, detail
    assert "ACP_OCR_TIMEOUT_S" in detail, (
        "an operator who cannot act on a limit is only being made anxious by it")


def test_a_readable_document_carries_no_unread_notice(docx):
    """The other direction. If this notice appeared on a clean read it would be noise on every
    scan, and an operator would learn to ignore the line that matters."""
    assert "OCR_IMAGE_UNREAD" not in _rules(ocr.images_of_text(docx, ".docx"))


# ── the contract underneath both ──────────────────────────────────────────────

def test_read_image_text_separates_read_nothing_from_read_nothing_at_all(docx, monkeypatch):
    """`""` = tesseract ran and this image holds no text. `None` = no reading happened. The
    whole defect above is these two having been the same value."""
    img = next(iter(ocr._embedded_images(docx, ".docx")))
    assert ocr.read_image_text(img, min_pixels=ocr._MIN_PIXELS) is not None

    _fail_first_n(monkeypatch, 99)
    ocr._OCR_CACHE.clear()
    assert ocr.read_image_text(img, min_pixels=ocr._MIN_PIXELS) is None


def test_ocr_text_still_hands_its_callers_a_string(docx, monkeypatch):
    """`ocr_text` grounds alt-text proposals and is used in string context throughout; the new
    contract must not reach those callers as a None."""
    img = next(iter(ocr._embedded_images(docx, ".docx")))
    _fail_first_n(monkeypatch, 99)
    assert ocr.ocr_text(img) == ""


def test_a_failed_reading_is_never_cached_as_an_answer(docx, monkeypatch):
    """Caching "" for a failure would poison every later pass over the same image — which is how
    the AA and AAA passes came to disagree in the first place, only worse: they would agree, on
    the wrong answer, for the rest of the process's life."""
    seen = _fail_first_n(monkeypatch, 1)
    img = next(iter(ocr._embedded_images(docx, ".docx")))
    assert ocr.read_image_text(img, min_pixels=ocr._MIN_PIXELS) is None
    assert seen["calls"] == 1
    text = ocr.read_image_text(img, min_pixels=ocr._MIN_PIXELS)
    assert text and "Quarterly" in text, f"the failure was cached as an answer: {text!r}"
