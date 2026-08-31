"""The 4.1.2 docx field-name lane, proved end to end (WCAG 4.1.2 Name, Role, Value).

Same bar as the lanes before it: the original document trips the finding, an approval changes
the saved document, a REAL re-scan verifies it, unrelated content survives, and a broken engine
earns no credit. Nothing but the blob store is patched.

WHAT THIS LANE WRITES. A Word content control's Title — `<w:alias w:val="…">` inside
`<w:sdtPr>`. That one attribute is simultaneously the accessible NAME Word exposes to assistive
technology (4.1.2) and the visible LABEL the field shows (3.3.2), which is why the detector
raises both criteria on the same condition and why the write clears both at once.
`test_the_same_write_also_clears_3_3_2` asserts that rather than leaving it as a happy accident:
the lane's `scs_to_clear` is {4.1.2} alone, so 3.3.2 clearing is a property of the DOCUMENT the
write produced, not of the lane's bookkeeping — and it is the kind of thing that goes quietly
wrong when a writer is retargeted.

WHY THE FIXTURE IS ASSEMBLED BY HAND. python-docx has no API for content controls at all, so
the `<w:sdt>` is injected into `word/document.xml` directly. That is also how the gallery type
matters: BOTH the detector and the writer gate on `_SDT_INPUT_TYPE` —
checkbox|date|dropDownList|comboBox|picture — because `w:sdt` wraps a great deal of non-form
Word content (tables of contents, citations, building blocks) that legitimately has no alias.
`w:text` and `w:richText` are excluded on purpose, since Word uses them for template
placeholders too.

That exclusion is not a detail. The first fixture written for this file used `<w:text/>`, and
the lane looked broken from both ends at once — the scan reported no 4.1.2 and
`apply_docx_field_name` returned the locator as unresolved. Nothing was wrong with the lane; the
document simply had no form field in it. `test_a_text_content_control_is_not_a_form_field`
pins that below, so the next reader meets the exclusion as an assertion rather than as a
mysterious empty result.

WHAT THIS CLAIMS: the control now carries an accessible name, the document survived, and ACP's
own criteria stop firing. NOT that the name is the right one for the field — that is the
reviewer's judgement, which is why this lane never auto-applies.
"""
from __future__ import annotations

import io
import sys
import tempfile
import zipfile
from pathlib import Path

import pytest

ACP = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ACP / "api"))

pytest.importorskip("docx")

FILE = "consent-form.docx"
SID = "rv-docx-412"
DOC = "word/document.xml"

HEADING = "Consent form"
LEAD = "Tick the box to agree to the terms below."
TAIL = "Unrelated closing paragraph that must survive the write."
APPROVED = "I agree to the terms"
SDT_ID = "101"
LOCATOR = f"docx:sdt:{SDT_ID}"


def _form(*, alias: str | None = None, gallery: str = "checkbox") -> bytes:
    """A Word document with one content control of `gallery` type, optionally already named.

    Assembled by editing `word/document.xml` because python-docx has no content-control API.
    """
    from docx import Document

    d = Document()
    d.add_heading(HEADING, level=1)
    d.add_paragraph(LEAD)
    d.add_paragraph(TAIL)
    base = Path(tempfile.mkdtemp()) / FILE
    d.save(base)

    alias_tag = f'<w:alias w:val="{alias}"/>' if alias else ""
    sdt = ("<w:p><w:sdt><w:sdtPr>" + alias_tag +
           f'<w:id w:val="{SDT_ID}"/><w:{gallery}/></w:sdtPr>'
           "<w:sdtContent><w:r><w:t>[ ]</w:t></w:r></w:sdtContent></w:sdt></w:p>")

    out = io.BytesIO()
    with zipfile.ZipFile(base) as zin, zipfile.ZipFile(out, "w") as zout:
        for name in zin.namelist():
            payload = zin.read(name)
            if name == DOC:
                payload = payload.replace(b"</w:body>", sdt.encode("utf-8") + b"</w:body>")
            zout.writestr(name, payload)
    return out.getvalue()


def _assess(data: bytes) -> set[str]:
    from assessment_policy import _extract_sc
    from scanner import analyse_and_assess
    with tempfile.TemporaryDirectory() as d:
        (Path(d) / FILE).write_bytes(data)
        fd, _ = analyse_and_assess(Path(d), FILE, detect_pii=False)
    return {sc for i in (fd or {}).get("issues", []) if (sc := _extract_sc(i.get("wcag", "")))}


def _spill(data: bytes) -> Path:
    p = Path(tempfile.mkdtemp()) / FILE
    p.write_bytes(data)
    return p


def _document_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read(DOC).decode("utf-8")


class _Blob:
    def __init__(self, data: bytes):
        self.data, self.uploads = data, []

    def download_remediated(self, owner, sid, f):
        return self.data

    def upload_remediated(self, owner, sid, f, data, mime):
        self.data = data
        self.uploads.append((f, mime))
        return "http://b/2"


@pytest.fixture()
def store(monkeypatch):
    import store as store_mod
    monkeypatch.setattr(store_mod, "_SQLITE_PATH", Path(tempfile.mkdtemp()) / "rv.db")
    return store_mod.Store()


def _seed(store, values: dict[str, str]) -> str:
    store.init_scan_run(SID, "drive", 1, "2026-08-31T00:00:00Z", "rubric", "hash")
    store.save_file_result(SID, {
        "file": FILE, "engine": "office", "status": "pass", "score": 60, "compliant": 0,
        "skipped_rules": 0, "drive_file_id": "d1",
        "issues": [{"ruleId": "DOCX_FORM_FIELD_NO_NAME", "wcag": "4.1.2 Name, Role, Value",
                    "severity": "SERIOUS", "locator": loc} for loc in values],
    }, "2026-08-31T00:00:00Z")
    store.record_remediation(SID, FILE, drive_write_url="http://d/1", blob_url="http://b/1")
    item_id = store.enqueue_proposals(SID, FILE, "4.1.2", [
        {"locator": loc, "before": "(no accessible name)", "proposed_value": "", "rationale": "r",
         "source": "reviewer"} for loc in values], rule_name="Name, Role, Value")
    store.update_hitl_item(item_id, "approved", None, None)
    store.approve_proposal_values(item_id, list(values.values()))
    return item_id


def _run_lane(monkeypatch, store, blob):
    import core
    import handlers
    monkeypatch.setattr(core, "store", store)
    monkeypatch.setitem(sys.modules, "blob", blob)
    handlers._apply_approved_values({"scan_id": SID, "file": FILE}, {})


@pytest.fixture(scope="module")
def unnamed() -> bytes:
    return _form()


# ── 1. the finding ────────────────────────────────────────────────────────────

def test_a_real_assessment_reports_4_1_2_on_an_unnamed_control(unnamed):
    assert "4.1.2" in _assess(unnamed)


def test_the_same_control_with_a_title_is_not_flagged():
    """The control. Without it, a detector that fired on every content control would satisfy
    the test above and nothing would notice. It also establishes independently of the writer
    that what clears the criterion is the alias."""
    assert "4.1.2" not in _assess(_form(alias="Agree to terms"))


def test_a_text_content_control_is_not_a_form_field():
    """`w:sdt` wraps tables of contents, citations and building blocks as well as form fields,
    so both the detector and the writer gate on the unambiguous input galleries only
    (`_SDT_INPUT_TYPE`: checkbox|date|dropDownList|comboBox|picture). `w:text` is excluded on
    purpose — Word uses it for template placeholders that are not fields a user fills in.

    Pinned here because the first fixture for this file used `w:text` and the lane read as
    broken from both ends: no finding from the scan, and an unresolved locator from the writer.
    """
    from apply_field_name import apply_docx_field_name
    text_control = _form(gallery="text")
    assert "4.1.2" not in _assess(text_control)
    _, applied, unresolved = apply_docx_field_name(text_control, {LOCATOR: APPROVED})
    assert applied == [] and unresolved == [LOCATOR]


# ── 2. approval → write → re-scan → credit, through the real path ─────────────

@pytest.fixture()
def applied(store, monkeypatch, unnamed):
    blob = _Blob(unnamed)
    _seed(store, {LOCATOR: APPROVED})
    _run_lane(monkeypatch, store, blob)
    return blob, store


def test_the_control_now_carries_the_approved_name(applied):
    blob, _ = applied
    assert f'<w:alias w:val="{APPROVED}"/>' in _document_xml(blob.data)


def test_the_same_write_also_clears_3_3_2(applied):
    """One attribute, two criteria. `w:alias` is both the accessible name Word exposes to
    assistive technology (4.1.2) and the visible label the field shows (3.3.2), and the
    detector raises both on the same condition.

    This is a property of the DOCUMENT rather than of the lane's bookkeeping — the lane's
    scs_to_clear is {4.1.2} alone — so it is asserted rather than assumed. A writer retargeted
    at some other attribute could still clear 4.1.2 and quietly leave 3.3.2 failing.
    """
    blob, _ = applied
    after = _assess(blob.data)
    assert "3.3.2" not in after and "4.1.2" not in after, after


def test_unrelated_content_survives(applied):
    blob, _ = applied
    xml = _document_xml(blob.data)
    for fragment in (HEADING, LEAD, TAIL):
        assert fragment in xml, f"{fragment!r} did not survive the write"


def test_the_saved_document_still_opens(applied):
    """Through python-docx, which had no part in writing the change."""
    from docx import Document
    blob, _ = applied
    path = _spill(blob.data)
    assert zipfile.ZipFile(path).testzip() is None
    assert any(HEADING in p.text for p in Document(str(path)).paragraphs)


def test_a_second_real_assessment_no_longer_reports_4_1_2(applied):
    """THE claim: a fresh assessment of the SAVED bytes, not the writer's return value."""
    blob, _ = applied
    assert "4.1.2" not in _assess(blob.data)


def test_the_row_is_credited_and_the_copy_is_stored(applied):
    blob, store = applied
    assert store.count_unapplied_approved_values(SID, FILE) == 0
    assert blob.uploads


# ── 3. where the lane must NOT credit or overwrite ────────────────────────────

def test_a_control_the_author_already_named_is_never_overwritten(store, monkeypatch):
    """The author's own Title wins. A reviewer approving a value for a field that has since
    been named must not silently replace it — the approval is stale, and the field is fine."""
    from apply_field_name import apply_docx_field_name
    named = _form(alias="Agree to terms")

    _, applied, unresolved = apply_docx_field_name(named, {LOCATOR: APPROVED})
    assert applied == [], "an author-supplied accessible name was overwritten"
    assert unresolved == [LOCATOR]

    blob = _Blob(named)
    _seed(store, {LOCATOR: APPROVED})
    _run_lane(monkeypatch, store, blob)
    assert blob.data == named and not blob.uploads
    assert 'w:val="Agree to terms"' in _document_xml(blob.data)


def test_an_approval_aimed_at_a_control_that_is_not_there_is_not_credited(store, monkeypatch,
                                                                          unnamed):
    """Nothing resolves, so nothing is written, nothing is credited, and nothing is
    published — never a guess at a different field."""
    blob = _Blob(unnamed)
    _seed(store, {"docx:sdt:999": APPROVED})
    _run_lane(monkeypatch, store, blob)

    assert blob.data == unnamed
    assert store.count_unapplied_approved_values(SID, FILE) == 1
    assert not blob.uploads


# ── 4. a broken engine earns nothing ──────────────────────────────────────────

@pytest.mark.parametrize("name,script,timeout", [
    ("cannot be launched", None, None),
    ("exits non-zero", "#!/bin/sh\necho boom >&2\nexit 9\n", None),
    ("hangs past the timeout", "#!/bin/sh\nsleep 30\n", "2"),
])
def test_a_broken_office_analyser_never_credits_this_lane_either(monkeypatch, unnamed,
                                                                 name, script, timeout):
    """Re-asserted per lane rather than assumed to inherit: the fail-open #1058 closed lived in
    ONE shared seam, so a regression there takes every lane at once. The 4.1.2 signal here is
    first-party (office_structure's content-control walk), so the residual is a real set even
    with no analyser at all."""
    import stat as _stat

    import scanner
    if script is None:
        monkeypatch.setattr(scanner, "DOTNET", "/nonexistent/dotnet", raising=False)
    else:
        fake = Path(tempfile.mkdtemp()) / "dotnet"
        fake.write_text(script)
        fake.chmod(fake.stat().st_mode | _stat.S_IEXEC)
        monkeypatch.setattr(scanner, "DOTNET", str(fake), raising=False)
    if timeout:
        monkeypatch.setenv("ACP_OFFICE_CLI_TIMEOUT", timeout)

    from proposals import verify_residual_scs
    residual = verify_residual_scs(unnamed, FILE)
    assert residual is not None, (
        f"an office CLI that {name} made the re-scan return None — every approved value on this "
        f"lane would be credited on a scan that never happened")
    assert "4.1.2" in residual
