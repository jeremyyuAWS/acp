"""1.4.1 and 1.4.11 on .docx: declared detectors, and a lane that carries the measurement.

WHAT WAS WRONG. Both detectors have shipped since ADR 0023 Phase 1b and their findings already
reached users through `office_structure.checks_for`. Neither pair was DECLARED: no registry
entry, so the capability matrix could not say what technique reached them, and no
`remediation_capability` lane, so the fix axis read "No Remediation" — an absence, when what
actually existed was work with nowhere to record it.

THE LANE IS `assisted`, AND GETTING THERE TOOK A CORRECTION. It shipped first as `human`, on the
reasoning that neither criterion has a value to prefill — 1.4.1 is editorial and recolouring a
shape is a design decision. That was too conservative, and the distinction it missed is between
the CRITERION and the SIGNAL: 1.4.1 in general is editorial, but what ACP detects is specifically
an explicitly removed underline, and putting it back is exact. Likewise the shade that brings an
outline to 3:1 against a known fill is computed, not guessed.

So both emit a prefilled card the human ELECTS — the 1.4.8 shape, where the change is exact and
the judgement is whether to make it. The card carries the measurement so nobody re-measures:
1.4.11 states the ratio found, the 3:1 needed, and a specific colour that reaches it.
"""
import io
import re
import sys
import zipfile
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "api"))
sys.path.insert(0, str(ROOT / "scripts"))

import formats.docx  # noqa: E402,F401  — importing the package is what runs register()
import assessment_policy as pol  # noqa: E402
import remediation_capability as cap  # noqa: E402


# ── the declarations ──────────────────────────────────────────────────────────

@pytest.mark.parametrize("sc", ["1.4.1", "1.4.11"])
def test_the_pair_is_registered(sc):
    reg = pol._registry_for(sc, "docx")
    assert reg is not None, f"{sc} on docx has a shipping detector but no registry declaration"
    assert reg.detector is not None


@pytest.mark.parametrize("sc", ["1.4.1", "1.4.11"])
def test_coverage_is_partial_not_full(sc):
    """FULL would assert we had looked at every way the criterion can fail. We have not: 1.4.1
    reads one signal (a link's removed underline) and 1.4.11 reads solid-outline-on-solid-fill
    shapes. Declaring FULL is how a matrix cell becomes a claim nobody checked."""
    from assessment import Coverage
    assert pol._registry_for(sc, "docx").coverage is Coverage.PARTIAL


@pytest.mark.parametrize("sc", ["1.4.1", "1.4.11"])
def test_the_reason_names_what_is_NOT_examined(sc):
    """A PARTIAL coverage whose reason only says what it does check is indistinguishable from
    FULL to anyone reading the matrix. The boundary is the point of the field."""
    reason = pol._registry_for(sc, "docx").reason
    assert "not examined" in reason, f"{sc}'s reason does not state its boundary: {reason!r}"


@pytest.mark.parametrize("sc", ["1.4.1", "1.4.11"])
def test_the_lane_is_assisted(sc):
    """`assisted` promises a prefilled card a human elects. Both criteria now have one, and the
    drift guard (test_remediation_capability) re-derives that against the real proposer — so
    this asserts the declaration and that one asserts it is true."""
    assert cap.REMEDIATION["docx"][sc] == "assisted"


@pytest.mark.parametrize("sc", ["1.4.1", "1.4.11"])
def test_the_criterion_still_cannot_pass_on_docx(sc):
    """Declaring a detector must not turn into certifying conformance. Both are review-lane: a
    signal reads REVIEW, silence reads NOT_EVALUATED, and neither is ever PASS."""
    assert pol._rule_outcome(sc, "docx", 0, 1, "AA", None) == pol.REVIEW
    assert pol._rule_outcome(sc, "docx", 0, 0, "AA", None) != "PASS"


# ── the findings still fire, and carry their numbers ──────────────────────────

def _docx_with_underlineless_link(path: Path) -> Path:
    import gen_demo_fixtures as gen
    from docx import Document
    d = Document()
    d.core_properties.title = "Colour only"
    p = d.add_paragraph("Read the ")
    gen._add_hyperlink(p, "https://example.org/policy", "accessibility policy")
    d.save(path)
    with zipfile.ZipFile(path) as z:
        parts = {n: z.read(n) for n in z.namelist()}
    parts["word/document.xml"] = parts["word/document.xml"].decode().replace(
        "<w:rPr/>", '<w:rPr><w:u w:val="none"/></w:rPr>').encode()
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in parts.items():
            z.writestr(n, b)
    return path


def test_the_registered_141_detector_still_fires(tmp_path):
    """The registration must point at the working implementation, not a stub. A declaration whose
    detector finds nothing is worse than no declaration: the matrix gains a claim and the product
    loses none of its blind spot."""
    from formats.docx.detectors import use_of_color
    out = use_of_color.detect(_docx_with_underlineless_link(tmp_path / "c.docx"))
    assert [f["ruleId"] for f in out] == ["DOCX_COLOR_ONLY_LINK"]
    assert out[0]["severity"] == "REVIEW"


def test_the_141_finding_says_how_many_links(tmp_path):
    """Carrying the count is what makes a human lane useful rather than a shrug."""
    from formats.docx.detectors import use_of_color
    out = use_of_color.detect(_docx_with_underlineless_link(tmp_path / "c.docx"))
    assert re.search(r"\b1 hyperlink", out[0]["detail"])
    assert out[0].get("evidence", {}).get("value") == 1


def test_a_normal_document_is_silent(tmp_path):
    """A link that kept its underline is not a 1.4.1 finding. Without this the detector could
    fire on everything and every assertion above would still pass."""
    import gen_demo_fixtures as gen
    from docx import Document
    from formats.docx.detectors import use_of_color
    d = Document()
    d.core_properties.title = "Fine"
    p = d.add_paragraph("Read the ")
    gen._add_hyperlink(p, "https://example.org/policy", "accessibility policy")
    path = tmp_path / "ok.docx"
    d.save(path)
    assert use_of_color.detect(path) == []


def test_the_registered_1411_detector_is_the_shipping_one():
    """Wrapper, not a reimplementation — the risk in this change is a declaration drifting from
    the code it describes, and a second copy of the measurement is how that starts."""
    import inspect

    from formats.docx.detectors import nontext_contrast
    src = inspect.getsource(nontext_contrast.detect)
    assert "docx_nontext_contrast_checks" in src


def test_the_1411_finding_carries_the_ratio_and_the_target():
    """The measurement is the whole value of a human lane here: a reviewer should never have to
    re-measure a contrast ratio ACP already computed. Asserted on the detector's own contract
    rather than by authoring a shape, which python-docx cannot express."""
    import inspect

    import office_structure as osx
    src = inspect.getsource(osx.docx_nontext_contrast_checks)
    assert '"required": 3.0' in src, "the 3:1 target must travel with the finding"
    assert "needs 3:1" in src, "the detail must state the target a reviewer is judging against"
    assert '"metric": "Contrast"' in src


# ── the proposers that make the lane `assisted` ───────────────────────────────
#
# `assisted` in this codebase means a proposer emits a PREFILLED card a human elects with one
# click — the 1.4.8 shape, where the change is exact and the judgement is whether to make it.
# The first version of this change shipped both criteria as `human` on the reasoning that
# neither had a value to prefill. That was too conservative, and the distinction is worth
# stating: 1.4.1 in GENERAL is editorial, but the SIGNAL detected is specifically an explicitly
# removed underline, and putting it back is exact. Same for 1.4.11: the shade that reaches 3:1
# against a known fill is computed, not guessed.

def test_the_141_proposer_emits_a_prefilled_card(tmp_path):
    import proposals
    out = proposals.propose_underline_restore(
        _docx_with_underlineless_link(tmp_path / "c.docx"), ".docx")
    assert len(out) == 1
    assert out[0]["sc"] == "1.4.1"
    assert "Restore the underline" in out[0]["proposed_value"]


def test_the_141_card_says_declining_it_is_legitimate(tmp_path):
    """Any non-colour cue satisfies 1.4.1 — an icon, bold weight, a "(link)" suffix. A card that
    presented the underline as THE fix would be overstating a worked example as a requirement."""
    import proposals
    out = proposals.propose_underline_restore(
        _docx_with_underlineless_link(tmp_path / "c.docx"), ".docx")
    assert "decline" in out[0]["rationale"].lower()


def test_the_141_proposer_is_silent_on_a_clean_document(tmp_path):
    import gen_demo_fixtures as gen
    import proposals
    from docx import Document
    d = Document()
    d.core_properties.title = "Fine"
    p = d.add_paragraph("Read the ")
    gen._add_hyperlink(p, "https://example.org/policy", "accessibility policy")
    path = tmp_path / "ok.docx"
    d.save(path)
    assert proposals.propose_underline_restore(path, ".docx") == []


def test_the_1411_proposer_names_a_shade_that_actually_reaches_3to1():
    """The whole value of the card. A suggestion that does not clear the threshold would send a
    reviewer to make a change that leaves the finding standing."""
    import office_structure as osx
    import proposals
    # Mid-grey outline on a light fill: below 3:1, and darkening reaches it.
    got = proposals._outline_reaching_3to1("BBBBBB", "FFFFFF")
    assert got is not None
    assert osx._contrast_ratio(got, "FFFFFF") >= 3.0


def test_the_1411_suggestion_stays_near_the_authors_colour():
    """It walks outward from the author's own colour and stops at the first shade that clears —
    a card proposing pure black on every faint outline would be correct and ignored."""
    import proposals
    got = proposals._outline_reaching_3to1("BBBBBB", "FFFFFF")
    assert got != "000000", "jumped straight to black instead of the nearest passing shade"


def test_the_search_always_finds_a_shade():
    """Written to prove the opposite, and the code was right. There is NO fill where neither
    black nor white reaches 3:1 — both failing would need its luminance below 0.10 and above
    0.30 at once — so the search cannot come back empty on well-formed input. Checked
    exhaustively across the greyscale range rather than argued."""
    import proposals
    for v in range(0, 256, 8):
        fill = f"{v:02X}" * 3
        assert proposals._outline_reaching_3to1("808080", fill) is not None, fill


def test_malformed_input_returns_nothing_rather_than_raising():
    """The only way None is reachable. A proposer that raised here would take down the whole
    scan's proposal pass for one unparseable colour."""
    import proposals
    assert proposals._outline_reaching_3to1("zzz", "FFFFFF") is None
    assert proposals._outline_reaching_3to1("", "FFFFFF") is None


@pytest.mark.parametrize("fn", ["propose_underline_restore", "propose_outline_contrast"])
def test_the_proposers_ignore_other_formats(fn, tmp_path):
    import proposals
    path = tmp_path / "x.pptx"
    path.write_bytes(b"not a docx")
    assert getattr(proposals, fn)(path, ".pptx") == []
