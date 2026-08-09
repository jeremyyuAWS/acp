"""The progress line, and what a review card tells the reviewer.

Two features, one file, because they answer the same complaint from opposite ends: ACP was silent
about what it was doing while it worked, and silent about WHY a decision reached a human when it
finished.

The assertions worth having here are the negative ones. A progress line is decoration; a progress
line that can raise is a regression that fails scans in production for the sake of a status bar.
So the contract under test is mostly "this cannot hurt anything": no scan_id is a no-op, a broken
store is swallowed, and the remediation path still completes when the activity channel is dead.
"""
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "api"))

import activity  # noqa: E402
import proposals  # noqa: E402


# ── the line itself ───────────────────────────────────────────────────────────

def test_line_names_file_criterion_and_action_in_that_order():
    """File first: the question a reader is actually asking is 'is it on MY document yet?'."""
    out = activity.line(file="benefits.docx", sc="1.1.1", action="describing images")
    assert out == "benefits.docx · 1.1.1 Non-text Content · describing images"


def test_line_expands_the_criterion_number_to_a_name():
    assert "Link Purpose (In Context)" in activity.line(sc="2.4.4", action="drafting")


def test_line_falls_back_to_the_bare_number_for_an_unknown_criterion():
    """A lookup miss must degrade to something readable, not to 'None' or a KeyError."""
    assert activity.line(sc="9.9.9", action="checking") == "9.9.9 · checking"


def test_line_survives_every_part_being_absent():
    assert activity.line() == ""
    assert activity.line(detail="warming up the model") == "warming up the model"


def test_detail_is_appended_after_a_dash():
    out = activity.line(file="a.docx", action="drafting", detail="nothing is written yet")
    assert out == "a.docx · drafting — nothing is written yet"


# ── the "this cannot hurt anything" contract ──────────────────────────────────

def test_record_without_a_scan_id_is_a_no_op_and_touches_no_store(monkeypatch):
    """Every offline caller — the benchmarks, the CLI, these tests — passes scan_id=None.

    None of them should pay for a store round-trip, and none should fail because no store is
    configured. Asserted by making any store access explode.
    """
    import core
    def boom(*a, **kw):
        raise AssertionError("record() must not touch the store without a scan_id")
    monkeypatch.setattr(core, "set_job", boom)
    activity.record(None, file="a.docx", action="describing images")   # must not raise


def test_record_swallows_a_broken_store(monkeypatch):
    """A progress line must never be able to fail the work it describes."""
    import core
    monkeypatch.setattr(core, "set_job", lambda *a, **kw: (_ for _ in ()).throw(RuntimeError("redis down")))
    activity.record("scan123", file="a.docx", action="describing images")   # must not raise


def test_current_returns_none_when_the_store_is_broken(monkeypatch):
    import core
    monkeypatch.setattr(core, "get_job_state",
                        lambda *a, **kw: (_ for _ in ()).throw(RuntimeError("redis down")))
    assert activity.current("scan123") is None


def test_rapid_calls_are_rate_limited(monkeypatch):
    """Per-finding publishing is the point; per-finding Redis round-trips are not.

    A vision call takes seconds and deserves a line. A junk-descr skip takes microseconds, and
    the same loop that publishes 35 useful lines over two minutes would publish 400 useless ones
    over 40ms — each a store write on the remediation hot path.
    """
    writes: list = []
    import core
    monkeypatch.setattr(core, "set_job", lambda k, v: writes.append(v))
    activity._last.clear()
    for i in range(200):
        activity.record("s-fast", file="a.docx", sc="1.1.1", action=f"describing image {i}")
    assert len(writes) == 1, f"expected the rate limit to collapse 200 calls, got {len(writes)}"


def test_force_bypasses_the_rate_limit(monkeypatch):
    """Stage transitions must never be dropped — a skipped one leaves the line stale on a
    criterion the run has already moved past, which is worse than no line at all."""
    writes: list = []
    import core
    monkeypatch.setattr(core, "set_job", lambda k, v: writes.append(v))
    activity._last.clear()
    activity.record("s-force", file="a.docx", sc="1.1.1", action="describing images")
    for sc in ("1.3.1", "2.4.4", "3.1.2"):
        activity.record("s-force", file="a.docx", sc=sc, action="stage", force=True)
    assert [w["sc"] for w in writes] == ["1.1.1", "1.3.1", "2.4.4", "3.1.2"]


def test_the_rate_limit_is_per_scan(monkeypatch):
    """Two documents remediating concurrently must not silence each other's line."""
    writes: list = []
    import core
    monkeypatch.setattr(core, "set_job", lambda k, v: writes.append(v))
    activity._last.clear()
    activity.record("scan-a", file="a.docx", action="x")
    activity.record("scan-b", file="b.docx", action="x")
    assert [w["file"] for w in writes] == ["a.docx", "b.docx"]


def test_record_then_current_round_trips(monkeypatch):
    store: dict = {}
    import core
    monkeypatch.setattr(core, "set_job", lambda k, v: store.__setitem__(k, v))
    monkeypatch.setattr(core, "get_job_state", lambda k: store.get(k))

    activity.record("s1", file="benefits.docx", sc="1.1.1", action="describing images",
                    phase="remediating")
    got = activity.current("s1")
    assert got["file"] == "benefits.docx"
    assert got["sc"] == "1.1.1"
    assert got["sc_name"] == "Non-text Content"
    assert got["phase"] == "remediating"
    assert "describing images" in got["text"]
    # Keyed by scan, not by job: assessment and remediation are different jobs in different
    # processes and must land on the same line.
    assert "activity:s1" in store


# ── the reviewer hand-off ─────────────────────────────────────────────────────

def test_proposal_carries_why_review_and_context():
    p = proposals.proposal(locator="word/document.xml#rId4", before="(no alt text)",
                           proposed_value="A bar chart comparing premiums.",
                           rationale="AI vision description", source="AI vision model (x)",
                           sc="1.1.1", why_review="nothing corroborates it",
                           context="Figure 3: premiums by plan")
    assert p["why_review"] == "nothing corroborates it"
    assert p["context"] == "Figure 3: premiums by plan"
    assert p["sc"] == "1.1.1"


def test_proposal_omits_the_handoff_keys_when_absent():
    """Optional, so every existing proposer keeps working unchanged and no card grows an
    empty 'why is this mine?' section it cannot fill."""
    p = proposals.proposal(locator="l", before="b", proposed_value="v", rationale="r", source="s")
    assert "why_review" not in p
    assert "context" not in p


@pytest.mark.parametrize("field", ["why_review", "context"])
def test_falsy_handoff_values_are_not_stored(field):
    p = proposals.proposal(locator="l", before="b", proposed_value="v", rationale="r",
                           source="s", **{field: ""})
    assert field not in p


def _chart_png() -> bytes:
    """Non-uniform and textless — the one image shape that reaches the vision lane.

    A near-uniform fill is intercepted by proposals.infer_decorative before vision is consulted,
    and an image with readable text is transcribed by OCR and auto-applied as GROUNDED. Only an
    image that is neither produces the ungrounded proposal this test is about. Getting this wrong
    is how a benchmark run measured four vision models as identical without calling any of them.
    """
    import io

    from PIL import Image, ImageDraw
    im = Image.new("RGB", (480, 280), (250, 250, 252))
    d = ImageDraw.Draw(im)
    for i in range(6):
        d.rectangle([30 + i * 70, 240 - (i * 31 % 170) - 30, 74 + i * 70, 240],
                    fill=(31 + i * 14, 78 + (i * 23) % 120, 160 - i * 9))
    d.line([(20, 240), (460, 240)], fill=(60, 60, 60), width=3)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()


def test_ungrounded_vision_proposal_explains_why_it_reached_a_human(tmp_path, monkeypatch):
    """The end-to-end assertion, through the PUBLIC path: an image nothing can corroborate
    produces a card that says so.

    This is the whole feature. The reviewer previously got a draft, a model name and a picture,
    and had to re-derive the entire assessment to decide whether to trust it — which is the
    low-value human work the review lane exists to remove.

    Driven through `remediate_office()` rather than the private `_vision_alt`, because that
    function takes a live regex match over the part's XML; constructing one by hand pins this
    test to an internal signature instead of to the behaviour anyone cares about.
    """
    import io

    import remediate_office
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.core_properties.title = "Premiums"
    doc.add_paragraph("Figure 3: premiums by plan")
    doc.add_paragraph().add_run().add_picture(io.BytesIO(_chart_png()), width=Inches(4))
    path = tmp_path / "premiums.docx"
    doc.save(path)

    import ai
    monkeypatch.setattr(ai, "vision_is_available", lambda *a, **kw: True)
    monkeypatch.setattr(ai, "describe_image_structured",
                        lambda *a, **kw: {"alt": "A bar chart comparing premiums by plan.",
                                          "grounded": False, "model": "test-vlm",
                                          "evidence": "no readable text in the image"})
    monkeypatch.setattr(ai, "_ALT_VALIDATOR_MODEL", "", raising=False)

    proposals_out: list = []
    remediate_office.remediate_office(path, ai_enabled=True, applied_fixes=[],
                                      proposals=proposals_out, diffs=[])

    vision = [p for p in proposals_out if "test-vlm" in str(p.get("source", ""))]
    assert vision, f"expected an ungrounded vision proposal, got {proposals_out}"
    p = vision[0]
    assert p["sc"] == "1.1.1"
    assert "corroborate" in p["why_review"], p["why_review"]
    assert p["proposed_value"] == "A bar chart comparing premiums by plan."
