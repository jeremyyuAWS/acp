"""The 1.1.1 docx alt-text lane, proved end to end (WCAG 1.1.1 Non-text Content).

The sixth lane to meet the REMEDIATION-VERIFIED bar, asserting the same five properties: the
original document trips the finding, an approval changes the document, a REAL re-scan verifies
it, unrelated content survives, and a broken engine earns no credit. Nothing but the blob store
is patched — `handlers._apply_approved_values` runs the production seam through
`proposals.verify_residual_scs` to `scanner.analyse_and_assess`.

WHAT IS DIFFERENT FROM THE 1.1.1 pptx PROOF, which is otherwise the same criterion and the same
writer. Two things, and both are docx-specific:

  * The locator resolves against `wp:docPr` in `word/document.xml` rather than a slide's
    `p:cNvPr` — a different part, a different tag, one applier. `test_the_finding_carries_a_
    locator_the_writer_can_resolve` pins the docx spelling, because a locator the writer cannot
    resolve makes the lane unreachable however good detection and writing are on their own.

  * THE DECORATIVE RESOLUTION IS EXERCISED HERE. It is the lane's second half and it writes no
    text at all: a reviewer who says "this image carries no information" is answered with the
    OOXML `adec:decorative` marker, not with a suppressed finding. That path reaches
    `_apply_one_value_kind` through `extra_work=True` — the one branch that lets a lane run with
    an EMPTY value map — and a file whose only 1.1.1 decision was "decorative" short-circuits
    without it. No end-to-end proof covered it before this file.

WHY THE APPROVED DESCRIPTION IS AUTHORED RATHER THAN PROPOSED, as in the pptx proof: 1.1.1's
proposer is an AI vision model, so a proposer-driven test would skip on a host without one, and
a lane that can only be proved where a model happens to be installed is not proved. A reviewer
writing alt text from scratch is also the real workflow for any image a model cannot describe;
`handlers` reads the approved value identically either way.

WHAT THIS CLAIMS, and it is narrower than "the image is now accessible": a `descr` a screen
reader will announce is present, the document survived, and ACP's own criterion stops firing.
Whether the sentence DESCRIBES the picture is a human judgement no detector makes — which is
why 1.1.1 resolves to REVIEW and never to PASS (ADR 0016).
"""
from __future__ import annotations

import io
import sys
import tempfile
import zipfile
from pathlib import Path

import pytest

ACP = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ACP / "api"))

pytest.importorskip("docx")

FILE = "accessibility-policy.docx"
SID = "rv-docx-111"
DOC = "word/document.xml"

HEADING = "Accessibility policy"
LEAD = "The board reviewed the programme. "
BOLD = "Progress was material"
TAIL = "Unrelated closing paragraph that must survive the write."
FIRST = "A bar chart of Q3 findings, grouped by severity."
SECOND = "A photograph of the accessibility working group."

# A 1x1 PNG — the smallest thing that is genuinely a picture part, so the document exercises a
# real image relationship rather than a shape pretending to be one.
_PNG = (b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00"
        b"\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4"
        b"\x00\x00\x00\x00IEND\xaeB`\x82")


def _doc(*, pictures: int = 2, alts: tuple[str | None, ...] = (None, None)) -> bytes:
    """A heading, a paragraph with a bold run, a table, `pictures` images, a closing paragraph.

    TWO pictures is not padding: it is what makes
    `test_a_partial_write_is_not_credited_because_the_criterion_still_fails` able to fail. A
    lane that credited on "the writer wrote something" rather than on the re-scan would pass a
    one-image document and be wrong about every real one. The heading, bold run and table are
    what would catch a writer that round-tripped the document through a lossy rebuild.
    """
    from docx import Document
    from docx.shared import Inches

    img = Path(tempfile.mkdtemp()) / "i.png"
    img.write_bytes(_PNG)

    d = Document()
    d.add_heading(HEADING, level=1)
    p = d.add_paragraph(LEAD)
    p.add_run(BOLD).bold = True
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "Quarter"
    t.cell(0, 1).text = "128"
    for n in range(pictures):
        run = d.add_paragraph().add_run()
        run.add_picture(str(img), width=Inches(1))
        alt = alts[n] if n < len(alts) else None
        if alt:
            run._element.xpath(".//*[local-name()='docPr']")[0].set("descr", alt)
    d.add_paragraph(TAIL)

    out = Path(tempfile.mkdtemp()) / FILE
    d.save(out)
    return out.read_bytes()


def _assess(data: bytes) -> set[str]:
    """The SCs a REAL assessment reports — the same call the production re-verification makes."""
    from assessment_policy import _extract_sc
    from scanner import analyse_and_assess
    with tempfile.TemporaryDirectory() as d:
        (Path(d) / FILE).write_bytes(data)
        fd, _ = analyse_and_assess(Path(d), FILE, detect_pii=False)
    return {sc for i in (fd or {}).get("issues", []) if (sc := _extract_sc(i.get("wcag", "")))}


def _spill(data: bytes) -> Path:
    p = Path(tempfile.mkdtemp()) / FILE
    p.write_bytes(data)
    return p


def _locators(data: bytes) -> list[str]:
    """The locators the FIRST-PARTY detector emits, read off the detector rather than off a full
    scan — and that distinction is the whole reason this helper exists.

    A full scan is not a stable source for these. `scanner._collapse_duplicate_alt` deliberately
    DROPS the first-party 1.1.1 findings whenever the .NET partner engine also reported 1.1.1
    ("the engine is the richer detector, so it wins where it ran"), and the engine's finding
    carries `page`/`location` rather than the `locator` an applier resolves. So on a host with
    no Office CLI the scan yields DOCX_IMAGE_NO_ALT with its locator, and on CI — where the
    engine IS built — it yields the engine's finding with none.

    The pptx proof learned this the expensive way: its first draft read locators off
    `analyse_and_assess`, passed locally for exactly that reason, and failed CI with `assert []`.
    Asking the detector directly is also the RIGHT question — the applier's locator contract is
    with the first-party detector, since that is the module that mints `part#name`.
    """
    from formats.docx.detectors import non_text_content
    return [f["locator"] for f in non_text_content.detect(_spill(data)) if f.get("locator")]


def _document_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read(DOC).decode("utf-8")


class _Blob:
    """The only thing patched in this module. Stores bytes verbatim; decides nothing."""

    def __init__(self, data: bytes):
        self.data, self.uploads = data, []

    def download_remediated(self, owner, sid, f):
        return self.data

    def upload_remediated(self, owner, sid, f, data, mime):
        self.data = data
        self.uploads.append((f, mime))
        return "http://b/2"


@pytest.fixture()
def store(monkeypatch):
    import store as store_mod
    monkeypatch.setattr(store_mod, "_SQLITE_PATH", Path(tempfile.mkdtemp()) / "rv.db")
    return store_mod.Store()


def _card(store, locators: list[str]) -> str:
    """A scanned + remediated document with one 1.1.1 card covering `locators`."""
    store.init_scan_run(SID, "drive", 1, "2026-08-31T00:00:00Z", "rubric", "hash")
    store.save_file_result(SID, {
        "file": FILE, "engine": "office", "status": "pass", "score": 60, "compliant": 0,
        "skipped_rules": 0, "drive_file_id": "d1",
        "issues": [{"ruleId": "DOCX_IMAGE_NO_ALT", "wcag": "1.1.1 Non-text Content",
                    "severity": "CRITICAL", "locator": loc} for loc in locators],
    }, "2026-08-31T00:00:00Z")
    store.record_remediation(SID, FILE, drive_write_url="http://d/1", blob_url="http://b/1")
    return store.enqueue_proposals(SID, FILE, "1.1.1", [
        {"locator": loc, "before": "(no alt text)", "proposed_value": "", "rationale": "r",
         "source": "reviewer"} for loc in locators], rule_name="Non-text Content")


def _seed_described(store, values: dict[str, str]) -> str:
    """The reviewer wrote a description for each image."""
    item_id = _card(store, list(values))
    store.update_hitl_item(item_id, "approved", None, None)
    store.approve_proposal_values(item_id, list(values.values()))
    return item_id


def _seed_decorative(store, locators: list[str]) -> str:
    """The reviewer judged the images decorative — a resolution, not a value."""
    item_id = _card(store, locators)
    store.update_hitl_item(item_id, "approved", None, None, resolution="decorative")
    return item_id


def _run_lane(monkeypatch, store, blob):
    """The production handler, with the re-scan UNPATCHED."""
    import core
    import handlers
    monkeypatch.setattr(core, "store", store)
    monkeypatch.setitem(sys.modules, "blob", blob)
    handlers._apply_approved_values({"scan_id": SID, "file": FILE}, {})


@pytest.fixture(scope="module")
def undescribed() -> bytes:
    return _doc(pictures=2, alts=(None, None))


# ── 1. the finding ────────────────────────────────────────────────────────────

def test_a_real_assessment_reports_1_1_1_on_undescribed_images(undescribed):
    assert "1.1.1" in _assess(undescribed), (
        "the fixture is supposed to fail 1.1.1 before remediation; without that every "
        "'the fix cleared it' assertion below is vacuous")


def test_the_same_document_with_both_images_described_is_not_flagged():
    """The control. Without it, a detector that flagged every image would satisfy the test
    above and the whole file would be measuring nothing."""
    assert "1.1.1" not in _assess(_doc(pictures=2, alts=(FIRST, SECOND)))


def test_the_finding_carries_a_locator_the_writer_can_resolve(undescribed):
    """The join between detection and writing — different modules, different parsers. A
    locator the writer cannot resolve makes the lane unreachable however good both halves are.
    """
    from apply_alt import parse_locator
    locs = _locators(undescribed)
    assert len(locs) == 2, f"expected two undescribed images, got {locs}"
    assert [parse_locator(loc) for loc in locs] == [(DOC, "Picture 1"), (DOC, "Picture 2")], locs


# ── 2. approval → write → re-scan → credit, through the real path ─────────────

@pytest.fixture()
def applied(store, monkeypatch, undescribed):
    locs = _locators(undescribed)
    blob = _Blob(undescribed)
    _seed_described(store, {locs[0]: FIRST, locs[1]: SECOND})
    _run_lane(monkeypatch, store, blob)
    return blob, store


def test_the_saved_document_carries_both_approved_descriptions(applied):
    blob, _ = applied
    xml = _document_xml(blob.data)
    assert f'descr="{FIRST}"' in xml and f'descr="{SECOND}"' in xml


def test_unrelated_content_and_formatting_survive(applied):
    """The assertion that would catch a writer that round-tripped the document through a lossy
    rebuild rather than editing the part in place."""
    blob, _ = applied
    xml = _document_xml(blob.data)
    for fragment in (HEADING, LEAD.strip(), BOLD, TAIL, "Quarter", "128"):
        assert fragment in xml, f"{fragment!r} did not survive the write"
    assert "<w:b/>" in xml, "the bold run property was dropped"


def test_the_saved_document_still_opens(applied):
    """Through python-docx, which had no part in writing the change, and as a valid zip."""
    from docx import Document
    blob, _ = applied
    path = _spill(blob.data)
    assert zipfile.ZipFile(path).testzip() is None
    doc = Document(str(path))
    assert any(HEADING in p.text for p in doc.paragraphs)
    assert any(TAIL in p.text for p in doc.paragraphs)


def test_a_second_real_assessment_no_longer_reports_1_1_1(applied):
    """THE claim: a fresh assessment of the SAVED bytes, not the writer's return value."""
    blob, _ = applied
    assert "1.1.1" not in _assess(blob.data)


def test_the_row_is_credited_and_the_copy_is_stored(applied):
    blob, store = applied
    assert store.count_unapplied_approved_values(SID, FILE) == 0
    assert blob.uploads, "the corrected copy was never stored"


# ── 3. the decorative resolution — the half of this lane that writes no text ───

def test_marking_both_images_decorative_also_clears_the_criterion(store, monkeypatch,
                                                                  undescribed):
    """A reviewer says these images carry no information. WCAG asks for that to be MARKED, not
    left blank — an unmarked decorative image is indistinguishable from a forgotten one, so
    every future scan re-raises the finding and every other tool reading the file still sees a
    picture missing its alt text.

    This path reaches `_apply_one_value_kind` with an EMPTY value map, through `extra_work=True`
    — the branch that exists so a file whose only 1.1.1 decision was "decorative" does not
    short-circuit before the marker is written. It is the one place in the lane where the write,
    the re-scan and the credit all happen with nothing in `approved_alt_values` at all.
    """
    locs = _locators(undescribed)
    blob = _Blob(undescribed)
    _seed_decorative(store, locs)
    _run_lane(monkeypatch, store, blob)

    xml = _document_xml(blob.data)
    assert "decorative" in xml, "no decorative marker reached the document"
    assert f'descr="{FIRST}"' not in xml, "a description was invented for a decorative image"
    assert "1.1.1" not in _assess(blob.data), (
        "the marker did not clear the criterion, so the reviewer's decision is recorded "
        "nowhere the document itself can carry")
    assert blob.uploads


# ── 4. where the lane must NOT credit ─────────────────────────────────────────

def test_a_partial_write_is_not_credited_because_the_criterion_still_fails(store, monkeypatch,
                                                                           undescribed):
    """Two undescribed images, one approved. The write succeeds and the document genuinely
    improves — and 1.1.1 still fails, because the other image is still silent.

    This is the control that separates "the writer wrote something" from "the criterion
    cleared". A lane crediting on the write would mark the file compliant here and publish a
    document that still fails the criterion it was certified against; under the
    `residual=set()` stub every other apply test supplies, it would pass by construction.
    """
    from apply_alt import apply_alt_text
    locs = _locators(undescribed)

    # What the write PRODUCES, established by calling the writer directly. The lane itself
    # discards these bytes — `_apply_one_value_kind` returns the ORIGINAL `working` when the
    # re-scan still reports the criterion, by design, so an unverified copy never becomes the
    # corrected one. Without this half, "no credit" below could mean the writer simply failed.
    written, applied, _ = apply_alt_text(undescribed, {locs[0]: FIRST})
    assert applied, "the writer refused the value, so this control is not about crediting"
    assert f'descr="{FIRST}"' in _document_xml(written)
    assert "1.1.1" in _assess(written), (
        "the fixture no longer fails after describing one of two images — this control cannot "
        "distinguish a withheld credit from a cleared one")

    blob = _Blob(undescribed)
    _seed_described(store, {locs[0]: FIRST})
    _run_lane(monkeypatch, store, blob)

    assert store.count_unapplied_approved_values(SID, FILE) == 1, (
        "the value was credited even though the criterion still fails on re-scan")
    assert store.mark_file_compliant_if_reviewed(SID, FILE) is False
    assert not blob.uploads, "an uncleared write was published as the corrected copy"
    assert blob.data == undescribed, "an unverified copy replaced the stored one"


def test_an_approved_value_aimed_at_an_image_the_document_does_not_have_is_not_credited(
        store, monkeypatch, undescribed):
    """The reviewer's approval names a third picture. Nothing resolves, so nothing is written,
    nothing is credited, and nothing is published — never a guess at different content."""
    blob = _Blob(undescribed)
    _seed_described(store, {f"{DOC}#Picture 9": FIRST})
    _run_lane(monkeypatch, store, blob)

    assert blob.data == undescribed
    assert store.count_unapplied_approved_values(SID, FILE) == 1
    assert not blob.uploads


def test_an_already_described_document_is_left_byte_identical(store, monkeypatch):
    """'We did not make it worse' is a claim a customer is entitled to."""
    clean = _doc(pictures=2, alts=(FIRST, SECOND))
    assert _locators(clean) == []
    blob = _Blob(clean)
    _card(store, [])
    _run_lane(monkeypatch, store, blob)
    assert blob.data == clean and not blob.uploads


# ── 5. a broken engine earns nothing ──────────────────────────────────────────

@pytest.mark.parametrize("name,script,timeout", [
    ("cannot be launched", None, None),
    ("exits non-zero", "#!/bin/sh\necho boom >&2\nexit 9\n", None),
    ("writes unparseable output", '#!/bin/sh\nprintf "not json" > "$3"\nexit 3\n', None),
    ("hangs past the timeout", "#!/bin/sh\nsleep 30\n", "2"),
])
def test_a_broken_office_analyser_never_credits_this_lane_either(monkeypatch, undescribed,
                                                                 name, script, timeout):
    """Re-asserted per lane rather than assumed to inherit: the fail-open #1058 closed lived in
    ONE shared seam (`scanner._analyse_office` → `verify_residual_scs` → `cleared = residual is
    None or …`), so a regression there takes every lane at once.

    The re-scan must return a real answer and never None, which `_apply_one_value_kind` reads
    as "cleared". It can, because `formats/docx/detectors/non_text_content` is pure Python and
    runs after the .NET call — the detector that module's docstring says was written precisely
    so an Office 1.1.1 credit is not granted on a criterion no first-party check observes.
    """
    import stat as _stat

    import scanner
    if script is None:
        monkeypatch.setattr(scanner, "DOTNET", "/nonexistent/dotnet", raising=False)
    else:
        fake = Path(tempfile.mkdtemp()) / "dotnet"
        fake.write_text(script)
        fake.chmod(fake.stat().st_mode | _stat.S_IEXEC)
        monkeypatch.setattr(scanner, "DOTNET", str(fake), raising=False)
    if timeout:
        monkeypatch.setenv("ACP_OFFICE_CLI_TIMEOUT", timeout)

    from proposals import verify_residual_scs
    residual = verify_residual_scs(undescribed, FILE)
    assert residual is not None, (
        f"an office CLI that {name} made the re-scan return None — every approved value on this "
        f"lane would be credited on a scan that never happened")
    assert "1.1.1" in residual
