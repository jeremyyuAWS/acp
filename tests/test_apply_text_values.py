"""Writing approved text-span values into a Word document (WCAG 1.3.3 / 3.1.2).

Mirrors test_apply_link_text.py's guarantees for the two lanes that had no writer at all:
the value a human signed their name to lands on the runs carrying that passage, a locator
that no longer matches is reported rather than approximated, and a document we could not
change comes back byte-identical.

The run-splitting cases are the ones worth having. A proposer's locator is a 60-character
prose prefix, and real prose does not respect run boundaries -- a sentence routinely starts
mid-run and ends in a different one. A writer that swallowed whole runs would quietly rewrite
text nobody approved, and that is invisible in the output unless a test looks for it.
"""
from __future__ import annotations
import io
import sys
import zipfile
from pathlib import Path

ACP = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ACP / "api"))

from apply_text_values import apply_sensory_rewrite, apply_language_parts  # noqa: E402

_DOC = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>{body}</w:body></w:document>"""


def _para(*runs: str) -> str:
    return "<w:p>" + "".join(runs) + "</w:p>"


def _run(text: str, rpr: str = "") -> str:
    return f'<w:r>{rpr}<w:t xml:space="preserve">{text}</w:t></w:r>'


def _pkg(body: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", "<Types/>")
        z.writestr("word/document.xml", _DOC.format(body=body))
    return buf.getvalue()


def _doc(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read("word/document.xml").decode("utf-8")


def _texts(data: bytes) -> list[str]:
    import re
    return re.findall(r"<w:t[^>]*>([^<]*)</w:t>", _doc(data))


# ── 1.3.3 sensory rewrite ─────────────────────────────────────────────────────

def test_sensory_rewrite_replaces_the_sentence():
    pkg = _pkg(_para(_run("Click the round green button to continue. Then sign out.")))
    out, applied, unresolved = apply_sensory_rewrite(
        pkg, "docx", {"Click the round green button to continue.": "Click Continue."})
    assert not unresolved
    assert len(applied) == 1
    joined = "".join(_texts(out))
    assert "Click Continue." in joined
    assert "round green button" not in joined
    # The following sentence is a different sentence and must survive untouched.
    assert "Then sign out." in joined


def test_sensory_rewrite_spanning_runs_keeps_the_rest_of_each_run():
    """The sentence starts mid-run and ends mid-run -- neither neighbour may be swallowed."""
    pkg = _pkg(_para(
        _run("Intro text. Press the red "),
        _run("square icon"),
        _run(" to save. Trailing text."),
    ))
    out, applied, _ = apply_sensory_rewrite(
        pkg, "docx", {"Press the red square icon to save.": "Press Save."})
    assert len(applied) == 1
    joined = "".join(_texts(out))
    assert joined.startswith("Intro text. ")
    assert "Press Save." in joined
    assert joined.endswith(" Trailing text.")
    assert "square icon" not in joined


def test_sensory_rewrite_preserves_formatting_of_the_untouched_half():
    bold = "<w:rPr><w:b/></w:rPr>"
    pkg = _pkg(_para(_run("Keep me bold. Use the blue tab.", bold)))
    out, _, _ = apply_sensory_rewrite(pkg, "docx", {"Use the blue tab.": "Use the Reports tab."})
    doc = _doc(out)
    # Both halves of the split run keep the original run properties.
    assert doc.count("<w:b/>") == 2
    joined = "".join(_texts(out))
    assert "Keep me bold." in joined and "Use the Reports tab." in joined


def test_sensory_locator_is_matched_across_a_run_boundary():
    """The proposer collapsed whitespace before truncating; the document did not."""
    pkg = _pkg(_para(_run("Select the large\n   round"), _run(" control now.")))
    out, applied, unresolved = apply_sensory_rewrite(
        pkg, "docx", {"Select the large round control now.": "Select Refresh."})
    assert applied and not unresolved
    assert "Select Refresh." in "".join(_texts(out))


# ── 3.1.2 language of parts ───────────────────────────────────────────────────

def test_language_marks_the_runs_without_changing_text():
    pkg = _pkg(_para(_run("Bonjour tout le monde")))
    out, applied, unresolved = apply_language_parts(pkg, "docx", {"Bonjour tout le monde": "fr"})
    assert not unresolved and len(applied) == 1
    assert '<w:lang w:val="fr"/>' in _doc(out)
    assert "".join(_texts(out)) == "Bonjour tout le monde"


def test_language_marks_only_the_foreign_sentence():
    """Neighbouring sentences must not be dragged into the mark.

    The passage is delimited the way textchecks._SEG_SPLIT delimits it -- on sentence
    punctuation -- so an adjacent English sentence is a different passage and stays unmarked.
    """
    pkg = _pkg(_para(_run("English lead in. "), _run("Bonjour le monde."), _run(" English tail.")))
    out, _, _ = apply_language_parts(pkg, "docx", {"Bonjour le monde.": "fr"})
    doc = _doc(out)
    assert doc.count('<w:lang w:val="fr"/>') == 1
    assert "".join(_texts(out)) == "English lead in. Bonjour le monde. English tail."
    # The marked run carries the French and nothing else.
    import re
    marked = re.findall(r'<w:r><w:rPr><w:lang w:val="fr"/></w:rPr>'
                        r'<w:t[^>]*>([^<]*)</w:t></w:r>', doc)
    assert marked == ["Bonjour le monde."]


def test_language_mark_covers_the_whole_passage_not_just_the_locator():
    """Regression: the locator is a 60-char lookup key, not the unit of work.

    Marking only the locator's span left a long passage tagged for 60 characters and
    untagged (mid-word) for the rest -- so textchecks still saw an unmarked foreign passage,
    3.1.2 kept firing, and the approved value could never earn credit.
    """
    fr = ("Nous avons publie le rapport trimestriel detaille pour toutes les regions "
          "concernees par cette reorganisation importante.")
    assert len(fr) > 60                                  # the bug needs a passage past the cut
    pkg = _pkg(_para(_run("English lead in. "), _run(fr)))
    out, _, _ = apply_language_parts(pkg, "docx", {fr[:60]: "fr"})
    import re
    marked = "".join(re.findall(r'<w:rPr><w:lang w:val="fr"/></w:rPr>'
                                r'<w:t[^>]*>([^<]*)</w:t>', _doc(out)))
    assert marked == fr, f"marked only {len(marked)} of {len(fr)} chars"


def test_language_splits_a_run_that_only_partly_covers_the_span():
    pkg = _pkg(_para(_run("Intro. Bonjour le monde. Outro.")))
    out, _, _ = apply_language_parts(pkg, "docx", {"Bonjour le monde": "fr"})
    doc = _doc(out)
    assert doc.count('<w:lang w:val="fr"/>') == 1
    # Text is preserved exactly -- only the middle run is marked.
    assert "".join(_texts(out)) == "Intro. Bonjour le monde. Outro."
    assert "Intro. " in _texts(out) and " . Outro." not in _texts(out)


def test_language_retags_an_existing_lang_rather_than_duplicating():
    pkg = _pkg(_para(_run("Bonjour le monde", '<w:rPr><w:lang w:val="en-US"/></w:rPr>')))
    out, _, _ = apply_language_parts(pkg, "docx", {"Bonjour le monde": "fr"})
    doc = _doc(out)
    assert '<w:lang w:val="fr"/>' in doc
    assert "en-US" not in doc
    assert doc.count("<w:lang") == 1


def test_language_rejects_a_value_that_is_not_a_language_code():
    """w:val is read by assistive tech as a language -- prose there is worse than nothing."""
    pkg = _pkg(_para(_run("Bonjour le monde")))
    out, applied, unresolved = apply_language_parts(
        pkg, "docx", {"Bonjour le monde": "probably French?"})
    assert not applied
    assert unresolved == ["Bonjour le monde"]
    assert out == pkg


# ── shared contract (matches apply_alt / apply_link_text) ─────────────────────

def test_unresolved_locator_is_reported_and_bytes_unchanged():
    pkg = _pkg(_para(_run("Nothing here matches.")))
    out, applied, unresolved = apply_sensory_rewrite(
        pkg, "docx", {"A sentence this document does not contain.": "Rewritten."})
    assert not applied
    assert unresolved == ["A sentence this document does not contain."]
    assert out == pkg


def test_partial_resolution_reports_only_the_missing_locator():
    pkg = _pkg(_para(_run("Use the green tab.")))
    _, applied, unresolved = apply_sensory_rewrite(
        pkg, "docx", {"Use the green tab.": "Use the Reports tab.", "Absent sentence.": "X"})
    assert [a["locator"] for a in applied] == ["Use the green tab."]
    assert unresolved == ["Absent sentence."]


def test_unsupported_format_resolves_nothing():
    """xlsx, specifically: SpreadsheetML has no per-run language element, so a language lane
    there could never clear its criterion. Unsupported by construction, not by omission."""
    pkg = _pkg(_para(_run("Bonjour le monde")))
    out, applied, unresolved = apply_language_parts(pkg, "xlsx", {"Bonjour le monde": "fr"})
    assert not applied and unresolved == ["Bonjour le monde"]
    assert out == pkg


# ── pptx: same primitive, different dialect ───────────────────────────────────

_SLIDE = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
<p:cSld><p:spTree>{body}</p:spTree></p:cSld></p:sld>"""


def _apara(*runs: str) -> str:
    return "<a:p>" + "".join(runs) + "</a:p>"


def _arun(text: str, rpr: str = "") -> str:
    return f"<a:r>{rpr}<a:t>{text}</a:t></a:r>"


def _ppkg(*slides: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", "<Types/>")
        for i, body in enumerate(slides, start=1):
            z.writestr(f"ppt/slides/slide{i}.xml", _SLIDE.format(body=body))
    return buf.getvalue()


def _slide(data: bytes, n: int = 1) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read(f"ppt/slides/slide{n}.xml").decode("utf-8")


def test_pptx_language_is_written_as_an_attribute_on_arpr():
    """PowerPoint's language is an ATTRIBUTE on a:rPr, not a child element like Word's."""
    pkg = _ppkg(_apara(_arun("Bonjour le monde.")))
    out, applied, unresolved = apply_language_parts(pkg, "pptx", {"Bonjour le monde.": "fr"})
    assert applied and not unresolved
    assert '<a:rPr lang="fr"/>' in _slide(out)


def test_pptx_replaces_the_default_language_rather_than_adding_a_second():
    """PowerPoint stamps lang="en-US" on nearly every run; the mark must retag, not duplicate."""
    pkg = _ppkg(_apara(_arun("Bonjour le monde.", '<a:rPr lang="en-US" dirty="0"/>')))
    out, _, _ = apply_language_parts(pkg, "pptx", {"Bonjour le monde.": "fr"})
    slide = _slide(out)
    assert slide.count("lang=") == 1
    assert 'lang="fr"' in slide and "en-US" not in slide
    assert 'dirty="0"' in slide                       # other run properties survive


def test_pptx_sensory_rewrite_replaces_the_sentence():
    pkg = _ppkg(_apara(_arun("Press the round green button to submit. Then close.")))
    out, applied, _ = apply_sensory_rewrite(
        pkg, "pptx", {"Press the round green button to submit.": "Press Submit."})
    assert len(applied) == 1
    slide = _slide(out)
    assert "Press Submit." in slide and "round green button" not in slide
    assert "Then close." in slide


def test_pptx_finds_the_passage_on_a_later_slide():
    """A pptx passage lives on one slide and the writer does not know which."""
    pkg = _ppkg(_apara(_arun("Nothing relevant here.")),
                _apara(_arun("Bonjour le monde.")))
    out, applied, unresolved = apply_language_parts(pkg, "pptx", {"Bonjour le monde.": "fr"})
    assert applied and not unresolved
    assert 'lang="fr"' in _slide(out, 2)
    assert "lang=" not in _slide(out, 1)               # slide 1 untouched


def test_pptx_splits_a_run_that_only_partly_covers_the_passage():
    pkg = _ppkg(_apara(_arun("English lead in. Bonjour le monde. English tail.")))
    out, _, _ = apply_language_parts(pkg, "pptx", {"Bonjour le monde.": "fr"})
    slide = _slide(out)
    assert slide.count('lang="fr"') == 1
    import re
    assert "".join(re.findall(r"<a:t>([^<]*)</a:t>", slide)) == \
        "English lead in. Bonjour le monde. English tail."


def test_empty_values_are_a_no_op():
    pkg = _pkg(_para(_run("Anything at all.")))
    for fn in (apply_sensory_rewrite, apply_language_parts):
        out, applied, unresolved = fn(pkg, "docx", {"Anything at all.": "   "})
        assert not applied and not unresolved and out == pkg


def test_output_is_still_a_readable_word_document():
    """The write must not corrupt the package -- python-docx has to reopen it."""
    docx = __import__("importlib").util.find_spec("docx")
    if docx is None:                                    # pragma: no cover
        import pytest
        pytest.skip("python-docx not installed")
    from docx import Document
    d = Document()
    d.add_paragraph("Press the round red button to submit.")
    buf = io.BytesIO()
    d.save(buf)
    out, applied, _ = apply_sensory_rewrite(
        buf.getvalue(), "docx", {"Press the round red button to submit.": "Press Submit."})
    assert applied
    reopened = Document(io.BytesIO(out))
    assert "Press Submit." in "\n".join(p.text for p in reopened.paragraphs)
