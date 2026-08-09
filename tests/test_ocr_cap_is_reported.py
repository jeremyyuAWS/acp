"""Exceeding the OCR image cap is reported, not silent.

THE DEFECT. `_MAX_IMAGES` (ACP_OCR_MAX_IMAGES, default 30) bounds how many embedded images are
OCR'd per document. The cap is right — OCR costs ~0.1s per image on a synthetic fixture and more
on a real scanned page, so an unbounded pass over a 500-image deck is a scan nobody waits for.
What was wrong is that exceeding it was SILENT: a 35-image document produced exactly 30
image-of-text findings and nothing said the other five were never looked at. That output is
indistinguishable from a document whose last five images are clean.

RAISING THE CAP IS NOT THE FIX, and this file exists partly to record why. A higher number still
truncates silently at the new number, and choosing a new default is a scan-time decision for
every customer taken on evidence from synthetic images. The knob already exists; what was missing
is any way for an operator to know they need it.

ADVISORY, NOT BLOCKING. The honest claim is not "these images fail" — nobody read them — it is
"this criterion was not fully checked on this document". A blocking finding would assert a defect
there is no evidence for; silence asserts conformance there is no evidence for either.
"""
import io
import sys
import zipfile
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "api"))

import ocr  # noqa: E402

pytestmark = pytest.mark.skipif(not ocr.is_available(), reason="tesseract not installed")


def _docx_with_images(path: Path, n: int, *, text: bool) -> Path:
    """A .docx holding `n` distinct images.

    DISTINCT bytes, always. An earlier fixture drew one of four images repeatedly and the OOXML
    package deduplicated 35 references down to 4 media parts — the OCR pass walks parts, not
    references, so it saw four images and the fixture appeared to prove a cap it never reached.
    """
    from docx import Document
    from docx.shared import Inches
    from PIL import Image, ImageDraw

    doc = Document()
    doc.core_properties.title = "Cap"
    for i in range(n):
        im = Image.new("RGB", (900, 200), "white")
        d = ImageDraw.Draw(im)
        if text:
            d.text((20, 40), f"Notice number {i:03d} regarding your enrollment", fill="black")
            d.text((20, 110), "please contact the benefits office before the deadline",
                   fill="black")
            im = im.resize((1800, 400))
        else:
            for b in range(6):                      # non-uniform, textless
                d.rectangle([20 + b * 140, 190 - ((i + b) * 17 % 150), 130 + b * 140, 190],
                            fill=(30 + b * 12, 80 + (i * 9) % 120, 160 - b * 8))
        buf = io.BytesIO()
        im.save(buf, format="PNG")
        doc.add_paragraph().add_run().add_picture(io.BytesIO(buf.getvalue()), width=Inches(4))
    doc.save(path)
    return path


def test_the_fixture_really_holds_distinct_images(tmp_path):
    """The premise. If the package deduplicated them this file would be testing nothing."""
    p = _docx_with_images(tmp_path / "d.docx", 8, text=True)
    with zipfile.ZipFile(p) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    assert len(media) == 8, f"expected 8 distinct media parts, got {len(media)}"


def test_over_the_cap_is_reported(tmp_path, monkeypatch):
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 3)
    p = _docx_with_images(tmp_path / "over.docx", 5, text=True)
    caps = [f for f in ocr.images_of_text(p, ".docx") if f["ruleId"] == "OCR_IMAGE_CAP_REACHED"]
    assert len(caps) == 1
    d = caps[0]["detail"]
    assert "5 images" in d and "first 3" in d and "2 were not examined" in d


def test_the_notice_is_advisory_not_blocking(tmp_path, monkeypatch):
    """REVIEW, so _rule_outcome routes it to a human rather than failing the document for images
    nobody read. A blocking finding here would assert a defect we have no evidence for."""
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 3)
    p = _docx_with_images(tmp_path / "over.docx", 5, text=True)
    cap = next(f for f in ocr.images_of_text(p, ".docx")
               if f["ruleId"] == "OCR_IMAGE_CAP_REACHED")
    assert cap["severity"] == "REVIEW"


def test_it_names_the_knob(tmp_path, monkeypatch):
    """An operator who cannot act on a limit is only being made anxious by it."""
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 3)
    p = _docx_with_images(tmp_path / "over.docx", 5, text=True)
    cap = next(f for f in ocr.images_of_text(p, ".docx")
               if f["ruleId"] == "OCR_IMAGE_CAP_REACHED")
    assert "ACP_OCR_MAX_IMAGES" in cap["detail"]


def test_a_CLEAN_document_over_the_cap_still_reports_it(tmp_path, monkeypatch):
    """THE case that matters. With textless images the scan finds no images-of-text at all, so
    before this the output was empty — identical to a fully-checked clean document. This is the
    exact document where silence reads as conformance."""
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 3)
    p = _docx_with_images(tmp_path / "clean.docx", 5, text=False)
    out = ocr.images_of_text(p, ".docx")
    assert [f["ruleId"] for f in out] == ["OCR_IMAGE_CAP_REACHED"], (
        "a clean over-cap document must still disclose that it was not fully checked")


def test_under_the_cap_says_nothing(tmp_path, monkeypatch):
    """A notice on a document that WAS fully checked is noise, and noise is how a real one stops
    being read."""
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 30)
    p = _docx_with_images(tmp_path / "under.docx", 4, text=True)
    assert not [f for f in ocr.images_of_text(p, ".docx")
                if f["ruleId"] == "OCR_IMAGE_CAP_REACHED"]


def test_exactly_at_the_cap_says_nothing(tmp_path, monkeypatch):
    """The off-by-one. `total > len(images)` — a document with exactly _MAX_IMAGES was fully
    examined and must not be told otherwise."""
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 4)
    p = _docx_with_images(tmp_path / "exact.docx", 4, text=True)
    assert not [f for f in ocr.images_of_text(p, ".docx")
                if f["ruleId"] == "OCR_IMAGE_CAP_REACHED"]


def test_the_notice_is_not_duplicated_by_the_AAA_rule(tmp_path, monkeypatch):
    """Both rules walk the same capped list. Reporting from each would show the reviewer the same
    truncation twice under two criteria; reporting only from 1.4.9 would hide it at an AA target,
    where 1.4.9 is out of scope — which is the target most scans run."""
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 3)
    p = _docx_with_images(tmp_path / "over.docx", 5, text=True)
    strict = ocr.images_of_text_no_exception(p, ".docx")
    assert not [f for f in strict if f["ruleId"] == "OCR_IMAGE_CAP_REACHED"]
    assert [f for f in ocr.images_of_text(p, ".docx")
            if f["ruleId"] == "OCR_IMAGE_CAP_REACHED"]


def test_the_cap_still_bounds_the_expensive_work(tmp_path, monkeypatch):
    """Counting the remainder must not decode or OCR it — otherwise the disclosure costs exactly
    what the cap was there to save."""
    monkeypatch.setattr(ocr, "_MAX_IMAGES", 3)
    p = _docx_with_images(tmp_path / "over.docx", 12, text=True)
    images, total = ocr._embedded_images_and_total(p, ".docx")
    assert len(images) == 3, "the cap must still bound what is materialised for OCR"
    assert total == 12, "the remainder must still be counted"
