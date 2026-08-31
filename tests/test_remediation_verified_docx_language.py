"""The 3.1.2 docx language-of-parts lane, proved end to end (WCAG 3.1.2 Language of Parts).

The fourth lane to meet the REMEDIATION-VERIFIED bar, asserting the same five properties: the
original document trips the finding, an approval changes the document, a REAL re-scan verifies
it, unrelated content survives, and a broken engine earns no credit. Nothing but the blob store
is patched — `handlers._apply_approved_values` runs the production seam.

THIS LANE DRIVES THE PROPOSER TOO, unlike the 1.1.1 pptx proof. `propose_language_parts` is
deterministic (langdetect over the extracted text, no model), so the test can start from the
document and take the reviewer's value from the draft ACP actually offers rather than authoring
one. That closes the last gap the alt-text proof had to leave open: here the locator crossing
between proposer and writer is exercised, not assumed.

WHAT THE FIX IS, and why it is unusually clean to verify. Marking a passage's runs with
`w:lang` changes NO VISIBLE TEXT — the words a sighted reader sees are byte-identical before and
after. So "unrelated content survives" can be asserted much more strongly than on the other
lanes: the whole extracted text must be unchanged, not merely the parts nobody meant to touch.
What changed is the machine-readable language, which is exactly what 3.1.2 is about.

WHAT THIS CLAIMS: the runs carrying the foreign passage now declare `fr`, the document survived,
and ACP's own criterion stops firing. NOT that the code is the right one for the passage —
langdetect's answer is a probabilistic guess a reviewer approves, which is why the lane is
`assisted` and never `auto`. A wrong-but-approved code would clear the criterion here exactly as
a right one does, and no detector can tell them apart.
"""
from __future__ import annotations

import io
import sys
import tempfile
import zipfile
from pathlib import Path

import pytest

ACP = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ACP / "api"))

pytest.importorskip("docx")
pytest.importorskip("langdetect")

FILE = "policy.docx"
SID = "rv-docx-312"
DOC = "word/document.xml"

TITLE = "Accessibility policy"
EN = "The board reviewed the programme and approved the following statement."
FR = ("Les organismes publics doivent rendre leurs documents accessibles à toutes les personnes "
      "en situation de handicap, conformément à la réglementation européenne en vigueur.")
TAIL = "Unrelated closing paragraph that must survive the write."


def _doc(*, mark: str | None = None) -> bytes:
    """An English document with one French passage. `mark` pre-marks it, for the clean control."""
    from docx import Document
    from docx.oxml.ns import qn

    d = Document()
    d.add_heading(TITLE, level=1)
    d.add_paragraph(EN)
    p = d.add_paragraph()
    run = p.add_run(FR)
    if mark:
        rpr = run._element.get_or_add_rPr()
        rpr.append(rpr.makeelement(qn("w:lang"), {qn("w:val"): mark}))
    d.add_paragraph(TAIL)

    out = Path(tempfile.mkdtemp()) / FILE
    d.save(out)
    return out.read_bytes()


def _assess(data: bytes) -> set[str]:
    from assessment_policy import _extract_sc
    from scanner import analyse_and_assess
    with tempfile.TemporaryDirectory() as dd:
        (Path(dd) / FILE).write_bytes(data)
        fd, _ = analyse_and_assess(Path(dd), FILE, detect_pii=False)
    return {sc for i in (fd or {}).get("issues", []) if (sc := _extract_sc(i.get("wcag", "")))}


def _spill(data: bytes) -> Path:
    p = Path(tempfile.mkdtemp()) / FILE
    p.write_bytes(data)
    return p


def _text(data: bytes) -> str:
    """The document's extracted text — the words a reader sees, through ACP's own extractor."""
    import pii
    return " ".join((pii.extract_text(_spill(data)) or "").split())


def _document_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read(DOC).decode("utf-8")


def _proposals(data: bytes) -> list[dict]:
    """What ACP actually offers a reviewer for this document."""
    import pii
    from proposals import propose_language_parts
    return propose_language_parts(pii.extract_text(_spill(data)))


class _Blob:
    def __init__(self, data: bytes):
        self.data, self.uploads = data, []

    def download_remediated(self, owner, sid, f):
        return self.data

    def upload_remediated(self, owner, sid, f, data, mime):
        self.data = data
        self.uploads.append((f, mime))
        return "http://b/2"


@pytest.fixture()
def store(monkeypatch):
    import store as store_mod
    monkeypatch.setattr(store_mod, "_SQLITE_PATH", Path(tempfile.mkdtemp()) / "rv.db")
    return store_mod.Store()


def _seed(store, props: list[dict], values: list[str]) -> int:
    store.init_scan_run(SID, "drive", 1, "2026-08-31T00:00:00Z", "rubric", "hash")
    store.save_file_result(SID, {
        "file": FILE, "engine": "office", "status": "pass", "score": 60, "compliant": 0,
        "skipped_rules": 0, "drive_file_id": "d1",
        "issues": [{"ruleId": "LANG_PARTS_UNMARKED", "wcag": "3.1.2 Language of Parts",
                    "severity": "SERIOUS"}],
    }, "2026-08-31T00:00:00Z")
    store.record_remediation(SID, FILE, drive_write_url="http://d/1", blob_url="http://b/1")
    item_id = store.enqueue_proposals(SID, FILE, "3.1.2", [
        {k: p.get(k) for k in ("locator", "before", "proposed_value", "rationale", "source")}
        for p in props], rule_name="Language of Parts")
    store.update_hitl_item(item_id, "approved", None, None)
    store.approve_proposal_values(item_id, values)
    return item_id


def _run_lane(monkeypatch, store, blob):
    import core
    import handlers
    monkeypatch.setattr(core, "store", store)
    monkeypatch.setitem(sys.modules, "blob", blob)
    handlers._apply_approved_values({"scan_id": SID, "file": FILE}, {})


@pytest.fixture(scope="module")
def unmarked() -> bytes:
    return _doc()


# ── 1. the finding and its proposal ──────────────────────────────────────────

def test_a_real_assessment_reports_3_1_2_on_an_unmarked_passage(unmarked):
    assert "3.1.2" in _assess(unmarked)


def test_the_same_document_with_the_passage_marked_is_not_flagged():
    """The control. It also proves the fix is the RIGHT one: what clears the criterion is the
    language mark, established here independently of the writer."""
    assert "3.1.2" not in _assess(_doc(mark="fr-FR"))


def test_the_proposer_offers_a_language_code_for_the_passage(unmarked):
    props = _proposals(unmarked)
    assert props, "a 3.1.2 finding with no proposal is a finding a reviewer cannot act on"
    assert props[0]["proposed_value"] == "fr"
    assert props[0]["locator"] and props[0]["locator"] in FR, (
        "the locator is not a prefix of the passage it describes, so the writer cannot find it")


def test_an_english_only_document_is_not_proposed_for():
    """Self-gating, asserted: the proposer must not invent a foreign span."""
    from docx import Document
    d = Document()
    d.add_heading(TITLE, level=1)
    d.add_paragraph(EN)
    d.add_paragraph(TAIL)
    out = Path(tempfile.mkdtemp()) / FILE
    d.save(out)
    assert _proposals(out.read_bytes()) == []


# ── 2. approval → write → re-scan → credit, through the real path ────────────

@pytest.fixture()
def applied(store, monkeypatch, unmarked):
    """The whole chain from the document: propose, approve the draft unedited, run the lane."""
    props = _proposals(unmarked)
    blob = _Blob(unmarked)
    _seed(store, props, [p["proposed_value"] for p in props])
    _run_lane(monkeypatch, store, blob)
    return blob, store


def test_the_passage_runs_now_declare_the_language(applied):
    blob, _ = applied
    xml = _document_xml(blob.data)
    assert 'w:lang' in xml and 'w:val="fr"' in xml, (
        "no w:lang mark reached the document")


def test_not_one_visible_character_changed(applied):
    """Stronger than "unrelated content survives", and available on this lane alone: a language
    mark is metadata, so the whole extracted text must be identical — the fix is invisible to a
    sighted reader by construction, and any difference means the writer touched prose."""
    blob, _ = applied
    assert _text(blob.data) == _text(_doc())


def test_the_document_still_opens_and_keeps_its_structure(applied):
    from docx import Document
    blob, _ = applied
    doc = Document(str(_spill(blob.data)))
    paras = [p.text for p in doc.paragraphs]
    assert TITLE in paras and EN in paras and FR in paras and TAIL in paras


def test_a_second_real_assessment_no_longer_reports_3_1_2(applied):
    blob, _ = applied
    assert "3.1.2" not in _assess(blob.data)


def test_the_row_is_credited_and_the_copy_is_stored(applied):
    blob, store = applied
    assert store.count_unapplied_approved_values(SID, FILE) == 0
    assert blob.uploads


# ── 3. where the lane must NOT credit ────────────────────────────────────────

def test_a_value_that_is_not_a_language_code_is_refused(store, monkeypatch, unmarked):
    """`apply_language_parts` writes into the attribute assistive technology reads as a
    language. A reviewer typing prose there must be reported, never written — a bogus code is
    worse than no mark, because it makes a screen reader switch voice to nothing sensible.
    """
    props = _proposals(unmarked)
    blob = _Blob(unmarked)
    _seed(store, props, ["French please"])
    _run_lane(monkeypatch, store, blob)

    assert "3.1.2" in _assess(blob.data), "the criterion cleared on a value that is not a code"
    assert store.count_unapplied_approved_values(SID, FILE) == 1
    assert not blob.uploads
    assert 'w:val="French please"' not in _document_xml(blob.data)


def test_an_already_marked_document_is_left_alone(store, monkeypatch):
    marked = _doc(mark="fr-FR")
    assert "3.1.2" not in _assess(marked)
    blob = _Blob(marked)
    _seed(store, [], [])
    _run_lane(monkeypatch, store, blob)
    assert blob.data == marked and not blob.uploads


# ── 4. a broken engine earns nothing ─────────────────────────────────────────

@pytest.mark.parametrize("name,script,timeout", [
    ("cannot be launched", None, None),
    ("exits non-zero", "#!/bin/sh\necho boom >&2\nexit 9\n", None),
    ("hangs past the timeout", "#!/bin/sh\nsleep 30\n", "2"),
])
def test_a_broken_office_analyser_never_credits_this_lane_either(monkeypatch, unmarked,
                                                                 name, script, timeout):
    """Re-asserted per lane rather than assumed to inherit: the fail-open lived in one shared
    seam, so a regression there takes every lane at once. 3.1.2 comes from textchecks, which
    runs after the .NET call in its own try/except — so the residual is a real set even with no
    analyser at all."""
    import stat as _stat

    import scanner
    if script is None:
        monkeypatch.setattr(scanner, "DOTNET", "/nonexistent/dotnet", raising=False)
    else:
        fake = Path(tempfile.mkdtemp()) / "dotnet"
        fake.write_text(script)
        fake.chmod(fake.stat().st_mode | _stat.S_IEXEC)
        monkeypatch.setattr(scanner, "DOTNET", str(fake), raising=False)
    if timeout:
        monkeypatch.setenv("ACP_OFFICE_CLI_TIMEOUT", timeout)

    from proposals import verify_residual_scs
    residual = verify_residual_scs(unmarked, FILE)
    assert residual is not None, (
        f"an office CLI that {name} made the re-scan return None — every approved value on this "
        f"lane would be credited on a scan that never happened")
    assert "3.1.2" in residual
