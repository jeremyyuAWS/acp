"""A heading with no text is reported (2.4.6).

THE GAP. The outline walk in `docx_checks` reads pStyle refs and never the heading's content, so
an empty Heading paragraph passed through every check ACP had: it is in the outline (no
pseudo-heading finding), it breaks no level sequence (no skip finding), and it has no runs to
fail contrast on. A `Heading2` with no text produced ZERO findings — verified by reading
word/document.xml on a fixture that had one.

WHY IT MATTERS MORE THAN IT LOOKS. Screen readers offer a heading list as the primary way to
navigate a long document. An empty entry is an announcement of nothing: the user is told a
section exists, cannot tell what it is, and cannot tell whether they have missed content. On a
25-page benefits handbook that is a navigation dead end.

THE EXEMPTIONS ARE THE INTERESTING PART. A heading whose content is an IMAGE is not textless by
mistake — its problem, if it has one, is 1.1.1's, and reporting it here would report one defect
twice under two criteria. That is the same reasoning `_link_purpose_finding` already applies to a
hyperlink wrapping a drawing.
"""
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "api"))

import office_structure as osx  # noqa: E402

RULE = "DOCX_HEADING_EMPTY"


def _docx(tmp_path, build) -> Path:
    from docx import Document
    doc = Document()
    doc.core_properties.title = "T"
    build(doc)
    p = tmp_path / "d.docx"
    doc.save(p)
    return p


def _empties(path: Path) -> list[dict]:
    return [f for f in osx.docx_checks(path) if f["ruleId"] == RULE]


# ── the defect ────────────────────────────────────────────────────────────────

def test_an_empty_heading_is_reported(tmp_path):
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_heading("", level=2)
        d.add_paragraph("Your coverage begins on March 1.")
    out = _empties(_docx(tmp_path, build))
    assert len(out) == 1
    assert out[0]["wcag"].startswith("2.4.6")


def test_the_finding_names_the_level_and_says_why(tmp_path):
    """A reviewer needs to find it in the document and know what it costs a reader."""
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_heading("", level=3)
    detail = _empties(_docx(tmp_path, build))[0]["detail"]
    assert "H3" in detail
    assert "heading list" in detail


def test_a_whitespace_only_heading_counts_as_empty(tmp_path):
    """Three spaces is a bug, not a deliberate blank — and a screen reader announces neither."""
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_heading("   ", level=2)
    assert len(_empties(_docx(tmp_path, build))) == 1


def test_every_empty_heading_is_reported_not_just_the_first(tmp_path):
    """Same reasoning as DOCX_HEADING_SKIP: stopping at the first understates the work and leaves
    a document looking one edit from clean when it is several."""
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_heading("", level=2)
        d.add_paragraph("text")
        d.add_heading("", level=2)
    assert len(_empties(_docx(tmp_path, build))) == 2


# ── what must NOT fire ────────────────────────────────────────────────────────

def test_a_normal_document_is_silent(tmp_path):
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_paragraph("Your coverage begins on March 1.")
        d.add_heading("Plan comparison", level=2)
        d.add_paragraph("Bronze and Silver.")
    assert _empties(_docx(tmp_path, build)) == []


def test_an_empty_BODY_paragraph_is_not_a_heading(tmp_path):
    """Blank spacer paragraphs are everywhere in real documents. Flagging them would bury the
    real finding under noise on every file — the fastest way to make a detector ignored."""
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_paragraph("")
        d.add_paragraph("   ")
        d.add_paragraph("Real text.")
    assert _empties(_docx(tmp_path, build)) == []


def test_a_heading_containing_only_an_IMAGE_is_exempt(tmp_path):
    """Not textless by mistake. Its problem, if any, is 1.1.1's — reporting it here would report
    one defect twice under two criteria."""
    import io

    from docx.shared import Inches
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (300, 80), (40, 90, 160)).save(buf, format="PNG")

    def build(d):
        d.add_heading("Notice", level=1)
        h = d.add_heading("", level=2)
        h.add_run().add_picture(io.BytesIO(buf.getvalue()), width=Inches(2))
    assert _empties(_docx(tmp_path, build)) == []


def test_a_heading_with_text_AND_an_image_is_silent(tmp_path):
    import io

    from docx.shared import Inches
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (300, 80), (40, 90, 160)).save(buf, format="PNG")

    def build(d):
        h = d.add_heading("Plan comparison", level=1)
        h.add_run().add_picture(io.BytesIO(buf.getvalue()), width=Inches(2))
    assert _empties(_docx(tmp_path, build)) == []


def test_it_does_not_disturb_the_level_skip_check(tmp_path):
    """The two 2.4.6 checks share a document and must not interfere: an empty heading still
    participates in the outline, so an H1 → (empty H2) → H3 sequence has no gap."""
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_heading("", level=2)
        d.add_heading("Deductibles", level=3)
    out = osx.docx_checks(_docx(tmp_path, build))
    assert [f["ruleId"] for f in out if f["ruleId"] == "DOCX_HEADING_SKIP"] == []
    assert len([f for f in out if f["ruleId"] == RULE]) == 1


@pytest.mark.parametrize("level", [1, 2, 3, 4, 5, 6])
def test_every_heading_level_is_covered(tmp_path, level):
    def build(d):
        d.add_heading("Notice", level=1)
        d.add_heading("", level=level)
    # An empty H1 immediately after a real H1 is still one empty heading.
    assert len(_empties(_docx(tmp_path, build))) == 1
