"""One remediation path proved end to end: reviewer-approved link text on a .docx (WCAG 2.4.4).

WHY THIS FILE EXISTS. `scripts/gen_capability_levels.py` reports REMEDIATION-VERIFIED as 0 of
17, and the bar it sets is deliberately higher than "the writer wrote something": a lane counts
only when a test drives an approved value through the real path and RE-OPENS the saved document
to see the criterion gone. Nothing met that bar, and the reason is uniform across the existing
apply tests — every one of them monkeypatches the re-scan:

    tests/test_apply_approved_values.py:79   monkeypatch.setattr(handlers, "_verify_residual_scs", lambda b, f: residual)
    tests/test_alt_locator_rid_writeback.py:329                    "                            "  lambda b, f: set()

Supplying `residual=set()` asserts what the lane does GIVEN a clean re-scan. It cannot tell you
whether a re-scan of these bytes would be clean, which is the only question a customer
downloading the file is asking. This module never patches it. Everything below runs the
production seam — `handlers._apply_approved_values` → `proposals.verify_residual_scs` →
`scanner.analyse_and_assess` — against a real Word document built by python-docx.

WHAT "VERIFIED" DOES AND DOES NOT MEAN HERE, stated so the entry in REMEDIATION_VERIFIED is not
read as more than it is. Proved: the visible link text a reader sees changes, the destination
does not, the surrounding document survives byte-for-byte where it should, the file still opens,
a second real assessment no longer reports 2.4.4, and an already-correct document is left alone.
Not proved: that a screen reader announces the new text the way a human would want it announced.
2.4.4 is "Link Purpose (In Context)" — context a user can reach programmatically counts, so a
detector keying on the link text alone is a conservative approximation in BOTH directions, and
"click here" inside a sentence naming the destination is not automatically a failure. This lane
verifies that ACP's own criterion stops firing on a document ACP changed. That is a real and
checkable claim; it is not a claim about conformance.

THE ORDER MATTERS AND IS THE POINT. Every stage below reads its input from the stage before it
rather than from a fixture written by hand, so a break anywhere in the chain surfaces as a
failure here rather than as four green unit tests either side of a hole.
"""
from __future__ import annotations

import sys
import tempfile
import zipfile
from pathlib import Path

import pytest

ACP = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ACP / "api"))
sys.path.insert(0, str(ACP / "scripts"))

pytest.importorskip("docx")

FILE = "quarterly-update.docx"
SID = "rv-docx-244"

VAGUE_HREF = "https://example.org/annual-accessibility-report-2026"
GOOD_HREF = "https://example.org/procurement-policy"
VAGUE_TEXT = "click here"
GOOD_TEXT = "procurement accessibility policy"
# The replacement. Deliberately the shape gen_sc_corpus's `f_link_descriptive` fixture uses for
# its ADVERSARIAL control ("2026 accessibility policy", declared to produce NO finding), because
# that fixture is checked against the real analyser in CI — so this text is known to read as
# descriptive to the detector rather than merely looking descriptive to me.
APPROVED_TEXT = "2026 annual accessibility report"


# ── the document ──────────────────────────────────────────────────────────────

def _build(vague: bool) -> bytes:
    """A small but not degenerate Word document: a heading, a paragraph carrying bold and italic
    runs, a table, and TWO hyperlinks — one under test, one already descriptive.

    The second link is not decoration. A writer that rewrote every hyperlink would pass a
    single-link fixture, and this document is what makes `test_the_other_hyperlink_is_untouched`
    able to fail.
    """
    from docx import Document
    from gen_demo_fixtures import _add_hyperlink

    d = Document()
    d.add_heading("Quarterly accessibility update", level=1)
    p = d.add_paragraph("The board reviewed the programme. ")
    p.add_run("Progress was material").bold = True
    p.add_run(" and ")
    p.add_run("risk fell").italic = True
    p.add_run(".")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Quarter"
    t.cell(0, 1).text = "Findings"
    t.cell(1, 0).text = "Q1"
    t.cell(1, 1).text = "128"
    p2 = d.add_paragraph("The full report is available online: ")
    _add_hyperlink(p2, VAGUE_HREF, VAGUE_TEXT if vague else APPROVED_TEXT)
    p3 = d.add_paragraph("See also the ")
    _add_hyperlink(p3, GOOD_HREF, GOOD_TEXT)
    p3.add_run(" for supplier requirements.")

    with tempfile.TemporaryDirectory() as dd:
        out = Path(dd) / FILE
        d.save(out)
        return out.read_bytes()


def _assess(data: bytes) -> set[str]:
    """The SCs a REAL assessment reports for these bytes — the same call the production
    re-verification makes, through the same normaliser the scan traces use."""
    from assessment_policy import _extract_sc
    from scanner import analyse_and_assess
    with tempfile.TemporaryDirectory() as dd:
        (Path(dd) / FILE).write_bytes(data)
        fd, _ = analyse_and_assess(Path(dd), FILE, detect_pii=False)
    return {sc for i in (fd or {}).get("issues", []) if (sc := _extract_sc(i.get("wcag", "")))}


def _document_xml(data: bytes) -> str:
    with zipfile.ZipFile(Path(_spill(data))) as z:
        return z.read("word/document.xml").decode("utf-8")


def _rels_xml(data: bytes) -> str:
    with zipfile.ZipFile(Path(_spill(data))) as z:
        return z.read("word/_rels/document.xml.rels").decode("utf-8")


def _spill(data: bytes) -> str:
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    f.write(data)
    f.close()
    return f.name


# ── the harness ───────────────────────────────────────────────────────────────

class _Blob:
    """Stands in for Azure Blob. The ONLY thing stubbed in this module, and it stores bytes
    verbatim — it decides nothing about whether the criterion cleared."""

    def __init__(self, data: bytes):
        self.data, self.uploads = data, []

    def download_remediated(self, owner, sid, f):
        return self.data

    def upload_remediated(self, owner, sid, f, data, mime):
        self.data = data
        self.uploads.append((f, mime))
        return "http://b/2"


@pytest.fixture()
def store(monkeypatch):
    import store as store_mod
    monkeypatch.setattr(store_mod, "_SQLITE_PATH", Path(tempfile.mkdtemp()) / "rv.db")
    return store_mod.Store()


def _seed(store, proposals: list[dict], *, sc: str = "2.4.4") -> int:
    """A scanned + remediated file with one HITL card carrying `proposals`, exactly as the
    proposal lane enqueues them."""
    store.init_scan_run(SID, "drive", 1, "2026-08-30T00:00:00Z", "rubric", "hash")
    store.save_file_result(SID, {
        "file": FILE, "engine": "office", "status": "pass", "score": 60, "compliant": 0,
        "skipped_rules": 0, "drive_file_id": "drv1",
        "issues": [{"ruleId": "DOCX_LINK_PURPOSE", "wcag": f"{sc} Link Purpose",
                    "severity": "MODERATE"}],
    }, "2026-08-30T00:00:00Z")
    store.record_remediation(SID, FILE, drive_write_url="http://d/1", blob_url="http://b/1")
    return store.enqueue_proposals(SID, FILE, sc, proposals, rule_name="Link Purpose (In Context)")


def _run_lane(monkeypatch, store, blob):
    """Drive the real handler. `_verify_residual_scs` is NOT patched — the credit below is
    decided by an actual re-scan of the actual written bytes."""
    import core
    import handlers
    monkeypatch.setattr(core, "store", store)
    monkeypatch.setitem(sys.modules, "blob", blob)
    handlers._apply_approved_values({"scan_id": SID, "file": FILE}, {})


@pytest.fixture(scope="module")
def original() -> bytes:
    return _build(vague=True)


@pytest.fixture(scope="module")
def already_correct() -> bytes:
    return _build(vague=False)


# ── 1. the assessment ─────────────────────────────────────────────────────────

def test_a_real_assessment_reports_2_4_4_on_the_original(original):
    """The starting condition, established by running the engine rather than by declaring it."""
    assert "2.4.4" in _assess(original), (
        "the fixture is supposed to fail 2.4.4 before remediation; if this fails, every "
        "'the fix cleared it' assertion below is vacuous")


def test_the_descriptive_control_does_not_report_2_4_4(already_correct):
    """The same document with descriptive link text. Without this, a detector that fired on
    every hyperlink would satisfy the test above and nothing would notice."""
    assert "2.4.4" not in _assess(already_correct)


# ── 2. the proposal ───────────────────────────────────────────────────────────

def test_the_finding_carries_a_proposal_a_reviewer_can_act_on(original):
    """A finding a reviewer cannot act on is not remediation. The proposal must name the
    destination it belongs to and offer a concrete replacement."""
    from proposals import propose_link_texts
    props = propose_link_texts(_spill(original), "docx", ai_enabled=False)
    mine = [p for p in props if p.get("locator") == VAGUE_HREF]
    assert mine, f"no proposal for {VAGUE_HREF}; got {[p.get('locator') for p in props]}"
    assert mine[0]["before"] == VAGUE_TEXT
    assert mine[0]["proposed_value"].strip(), "a proposal with no value is a finding with no fix"
    assert mine[0].get("sc") == "2.4.4"


def test_no_proposal_is_raised_for_the_already_descriptive_link(original):
    """The second hyperlink is fine and must be left alone at the PROPOSAL stage — the
    cheapest place to get this wrong is offering reviewers work that isn't needed."""
    from proposals import propose_link_texts
    props = propose_link_texts(_spill(original), "docx", ai_enabled=False)
    assert GOOD_HREF not in {p.get("locator") for p in props}


# ── 3. approval → write → re-scan → credit, through the real path ─────────────

@pytest.fixture()
def applied(store, monkeypatch, original):
    """The whole chain, run once: seed the proposal, approve it, run the lane. Returns the blob
    so the stages below inspect the bytes that were actually saved."""
    from proposals import propose_link_texts
    props = [p for p in propose_link_texts(_spill(original), "docx", ai_enabled=False)
             if p.get("locator") == VAGUE_HREF]
    item_id = _seed(store, [{k: p[k] for k in ("locator", "before", "proposed_value",
                                               "rationale", "source")} for p in props])
    store.update_hitl_item(item_id, "approved", None, None)
    store.approve_proposal_values(item_id, [APPROVED_TEXT])   # the reviewer edits the draft
    blob = _Blob(original)
    _run_lane(monkeypatch, store, blob)
    return blob, item_id, store


def test_the_saved_document_carries_the_approved_text(applied):
    """The intended property actually changed — read back out of the saved package, not from
    the writer's return value."""
    blob, _, _ = applied
    xml = _document_xml(blob.data)
    assert APPROVED_TEXT in xml
    assert VAGUE_TEXT not in xml, "the vague text is still in the document a user will open"


def test_the_hyperlink_target_is_unchanged(applied):
    """Rewriting the label must not move the destination. A 'fix' that silently repointed a
    link would read as a success everywhere else in this file."""
    blob, _, _ = applied
    rels = _rels_xml(blob.data)
    assert VAGUE_HREF in rels
    assert GOOD_HREF in rels


def test_the_other_hyperlink_is_untouched(applied):
    """Scoped write: the descriptive link keeps its text."""
    blob, _, _ = applied
    assert GOOD_TEXT in _document_xml(blob.data)


def test_unrelated_content_and_formatting_survive(applied):
    """Heading, both emphasised runs and the table are still there. This is the assertion that
    would catch a writer that round-tripped the document through a lossy rebuild."""
    blob, _, _ = applied
    xml = _document_xml(blob.data)
    for fragment in ("Quarterly accessibility update", "Progress was material", "risk fell",
                     "Quarter", "Findings", "128"):
        assert fragment in xml, f"{fragment!r} did not survive the write"
    assert "<w:b/>" in xml and "<w:i/>" in xml, "bold/italic run properties were dropped"


def test_the_saved_document_still_opens(applied):
    """Opens as a Word document through a library that was not involved in writing it, and the
    package is a valid zip with an intact content-type map."""
    from docx import Document
    blob, _, _ = applied
    path = _spill(blob.data)
    assert zipfile.ZipFile(path).testzip() is None
    doc = Document(path)
    assert any("Quarterly accessibility update" in p.text for p in doc.paragraphs)


def test_a_second_real_assessment_no_longer_reports_2_4_4(applied):
    """THE claim. Not "the writer wrote": a fresh assessment of the saved bytes, through the
    same entry point the first one used, no longer reports the criterion."""
    blob, _, _ = applied
    assert "2.4.4" not in _assess(blob.data)


def test_the_row_is_credited_and_the_file_certifies(applied):
    """The bookkeeping the reviewer sees — and it is downstream of the re-scan above, not of
    the approval: `_apply_one_value_kind` credits only when the residual no longer holds the
    criterion."""
    blob, item_id, store = applied
    assert store.count_unapplied_approved_values(SID, FILE) == 0
    assert blob.uploads, "the corrected copy was never stored"


# ── 4. the negative controls — where a lane must NOT credit ───────────────────

def test_an_approved_value_that_does_not_clear_the_criterion_is_not_credited(
        store, monkeypatch, original):
    """A reviewer approves text that is ITSELF vague. The write succeeds; the re-scan still
    reports 2.4.4; nothing may be credited and nothing may be published.

    This is the test that distinguishes a real re-scan from `residual=set()`. Under the stub it
    would pass by construction — the stub says "clean" whatever was written.
    """
    from proposals import propose_link_texts
    props = [p for p in propose_link_texts(_spill(original), "docx", ai_enabled=False)
             if p.get("locator") == VAGUE_HREF]
    item_id = _seed(store, [{k: p[k] for k in ("locator", "before", "proposed_value",
                                               "rationale", "source")} for p in props])
    store.update_hitl_item(item_id, "approved", None, None)
    store.approve_proposal_values(item_id, ["read more"])     # still vague
    blob = _Blob(original)
    _run_lane(monkeypatch, store, blob)

    assert "2.4.4" in _assess(blob.data), "the fixture no longer fails; the control is vacuous"
    assert store.count_unapplied_approved_values(SID, FILE) == 1, (
        "a value that did not clear the criterion was credited anyway")
    assert store.mark_file_compliant_if_reviewed(SID, FILE) is False
    assert not blob.uploads, "an uncleared write was published as the corrected copy"


def test_an_already_correct_document_is_not_damaged(store, monkeypatch, already_correct):
    """Nothing is proposed for it, so the lane has nothing to write — and the stored copy comes
    back byte-identical. 'We did not make it worse' is a claim a customer is entitled to."""
    from proposals import propose_link_texts
    assert [p for p in propose_link_texts(_spill(already_correct), "docx", ai_enabled=False)
            if p.get("locator") == VAGUE_HREF] == []

    _seed(store, [])                       # a card with no proposals: nothing approved
    blob = _Blob(already_correct)
    _run_lane(monkeypatch, store, blob)
    assert blob.data == already_correct
    assert not blob.uploads


# ── 5. the regression this file was written on top of ─────────────────────────

def test_a_missing_office_analyser_does_not_credit_on_a_rescan_that_never_ran(
        store, monkeypatch, original):
    """`scanner._analyse_office` used to let a CLI LAUNCH failure escape — no dotnet on PATH,
    DOTNET pointing at nothing — while catching every other CLI failure as a result. Chained
    with `proposals.verify_residual_scs`'s `except Exception: return None` and
    `_apply_one_value_kind`'s `cleared = residual is None or …`, that credited every approved
    value and certified the file on a re-scan that never ran.

    Here the analyser is pointed at a path that does not exist, an approved value is written
    that does NOT clear 2.4.4, and the lane must still withhold credit — because the first-party
    checks (office_structure, pure Python) run after the .NET call and still report.
    """
    import scanner
    monkeypatch.setattr(scanner, "DOTNET", "/nonexistent/dotnet", raising=False)

    from proposals import propose_link_texts
    props = [p for p in propose_link_texts(_spill(original), "docx", ai_enabled=False)
             if p.get("locator") == VAGUE_HREF]
    item_id = _seed(store, [{k: p[k] for k in ("locator", "before", "proposed_value",
                                               "rationale", "source")} for p in props])
    store.update_hitl_item(item_id, "approved", None, None)
    # "read more", not "click here to continue". propose_link_texts' predicate is deliberately a
    # SUPERSET of the detector's word list (its docstring says so), and the first draft of this
    # test used a phrase only the proposer calls vague — so the re-scan correctly reported
    # nothing and the test failed for a reason that had nothing to do with the analyser. The
    # control has to be text the DETECTOR still fails, or it proves nothing about crediting.
    store.approve_proposal_values(item_id, ["read more"])                # still vague
    blob = _Blob(original)
    _run_lane(monkeypatch, store, blob)

    assert "2.4.4" in _assess(blob.data), (
        "the written text no longer fails 2.4.4 even to the first-party detector, so this "
        "control cannot distinguish a withheld credit from a cleared one")
    assert store.count_unapplied_approved_values(SID, FILE) == 1, (
        "credit was granted with no working Office analyser — the fail-open is back")
    assert not blob.uploads
