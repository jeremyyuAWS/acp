"""Regression test for the .NET HeadingStructureRule's outline-level signal.

HeadingStructureRule (engine/office-analysers/.../Docx/Rules/HeadingStructureRule.cs)
originally matched headings only by the paragraph's built-in style ID
(Heading1..Heading6). A paragraph can be a heading via w:pPr/w:outlineLvl without
using one of those style IDs -- e.g. a custom-named heading style with an explicit
outline level -- and such a paragraph was invisible to the rule's skip-level check,
producing a false "heading level skipped" finding around it.

The fix adds w:pPr/w:outlineLvl as a co-equal signal (level = val + 1, since OOXML
outline levels are 0-based) alongside the style-ID dictionary lookup.

Shells out to the built .NET CLI; self-skips without the toolchain (same pattern as
tests/test_office_language_rules.py).
"""
from __future__ import annotations

import json
import os
import subprocess
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")
from docx import Document               # noqa: E402
from docx.oxml.ns import qn              # noqa: E402

_ACP = Path(__file__).resolve().parents[1]
_DOTNET = Path(os.environ.get("ACP_DOTNET") or os.path.expanduser("~/.dotnet/dotnet"))
_CLI = Path(os.environ.get("ACP_OFFICE_CLI")
            or _ACP / "spike/dotnet/AcpScan.Cli/bin/Release/net10.0/AcpScan.Cli.dll")

pytestmark = pytest.mark.skipif(
    not _DOTNET.exists() or not _CLI.exists(),
    reason="office analyser CLI or dotnet runtime not built/available",
)


def _add_custom_outline_heading(doc: Document, text: str, outline_val: str) -> None:
    """Add a paragraph in a plain style (not one of the built-in HeadingN style IDs)
    that carries an explicit w:pPr/w:outlineLvl instead -- the custom-heading-style
    case the rule must recognise without a HeadingN style ID."""
    p = doc.add_paragraph(text, style="Normal")
    p_pr = p._p.get_or_add_pPr()
    outline_lvl = p_pr.makeelement(qn("w:outlineLvl"), {qn("w:val"): outline_val})
    p_pr.append(outline_lvl)


def _run_cli(indir: Path, tmp_path: Path) -> dict[str, list[dict]]:
    out = tmp_path / "out.json"
    env = {**os.environ, "DOTNET_ROOT": os.path.expanduser("~/.dotnet"),
           "DOTNET_CLI_TELEMETRY_OPTOUT": "1", "DOTNET_NOLOGO": "1"}
    subprocess.run([str(_DOTNET), str(_CLI), str(indir), str(out)],
                   capture_output=True, text=True, env=env, check=True)
    return {item["file"]: item.get("issues", []) for item in json.loads(out.read_text())}


def _skip_findings(issues: list[dict]) -> list[dict]:
    return [i for i in issues if "skipped" in (i.get("title", "") or "").lower()]


def test_custom_style_outline_level_closes_the_gap(tmp_path):
    # H1 -> custom-style "H2-equivalent" (outlineLvl=1, no built-in style ID) -> H3.
    # If outlineLvl is ignored, the rule only sees H1 -> H3 and wrongly reports a skip.
    doc = Document()
    doc.add_heading("Overview", level=1)
    _add_custom_outline_heading(doc, "Custom-styled section heading", "1")
    doc.add_heading("Subsection", level=3)
    src = tmp_path / "custom_outline.docx"
    doc.save(src)

    res = _run_cli(tmp_path, tmp_path)
    assert _skip_findings(res["custom_outline.docx"]) == [], (
        "outlineLvl-based heading should close the H1->H3 gap, not leave it looking skipped"
    )


def test_custom_style_outline_level_still_flags_a_real_skip(tmp_path):
    # H1 -> custom-style "H4-equivalent" (outlineLvl=3) with nothing at H2/H3 in between
    # is a genuine skip, and must still be caught now that outlineLvl feeds the rule.
    doc = Document()
    doc.add_heading("Overview", level=1)
    _add_custom_outline_heading(doc, "Custom-styled deep heading", "3")
    src = tmp_path / "custom_outline_skip.docx"
    doc.save(src)

    res = _run_cli(tmp_path, tmp_path)
    assert _skip_findings(res["custom_outline_skip.docx"]), (
        "a custom-styled heading with outlineLvl should still trip a genuine level skip"
    )
